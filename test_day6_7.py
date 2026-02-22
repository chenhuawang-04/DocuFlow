"""
DocuFlow MCP - Day 6-7 Batch Operations Test

Tests all batch operation tools including:
- batch_format_range
- batch_apply_style
- batch_copy_format
- batch_replace_format
"""

import sys
import os
import time

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from docuflow_mcp.core.registry import dispatch_tool, get_all_registered_tools
from docx import Document


def print_section(title):
    """Print a section header"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}\n")


def test_batch_format_range():
    """Test batch_format_range tool"""
    print_section("Test 1: batch_format_range")

    test_file = "E:/Project/DocuFlow/test_batch_range.docx"

    # Create test document with 20 paragraphs
    doc = Document()
    for i in range(20):
        doc.add_paragraph(f"这是第 {i+1} 个段落的测试内容。")
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file} (20个段落)")

    # Test 1: Format paragraphs 5-15 (11 paragraphs)
    print("\n测试: 批量格式化段落 5-15")
    result = dispatch_tool("batch_format_range", {
        "path": test_file,
        "start_index": 5,
        "end_index": 15,
        "font_name": "宋体",
        "font_size": "14pt",
        "bold": True,
        "alignment": "center",
        "line_spacing": 1.5
    })

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     格式化数量: {result['formatted_count']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Format large range (performance test)
    print("\n测试: 批量格式化段落 0-19 (性能测试)")
    start_time = time.time()

    result = dispatch_tool("batch_format_range", {
        "path": test_file,
        "start_index": 0,
        "end_index": 19,
        "font_name": "微软雅黑",
        "font_size": "12pt",
        "alignment": "justify"
    })

    elapsed = (time.time() - start_time) * 1000

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     性能: {elapsed:.2f}ms for 20 paragraphs")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 3: Invalid range
    print("\n测试: 无效索引范围")
    result = dispatch_tool("batch_format_range", {
        "path": test_file,
        "start_index": 0,
        "end_index": 100  # Out of range
    })

    if not result.get("success"):
        print(f"[OK] 正确检测到错误: {result.get('error')}")
    else:
        print(f"[FAIL] 应该返回错误但返回成功")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] batch_format_range 测试通过")
    return True


def test_batch_apply_style():
    """Test batch_apply_style tool"""
    print_section("Test 2: batch_apply_style")

    test_file = "E:/Project/DocuFlow/test_batch_style.docx"

    # Create test document
    doc = Document()
    for i in range(10):
        doc.add_paragraph(f"段落 {i+1}")
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file} (10个段落)")

    # Test 1: Apply Heading 1 style to multiple paragraphs
    print("\n测试: 批量应用 Heading 1 样式")
    result = dispatch_tool("batch_apply_style", {
        "path": test_file,
        "paragraph_indices": [0, 2, 4, 6, 8],
        "style_name": "Heading 1"
    })

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     应用数量: {result['applied_count']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Verify styles were applied
    doc = Document(test_file)
    correct_count = sum(1 for idx in [0, 2, 4, 6, 8] if doc.paragraphs[idx].style.name == "Heading 1")
    if correct_count == 5:
        print(f"[OK] 验证: 5个段落的样式已正确应用")
    else:
        print(f"[FAIL] 验证失败: 只有 {correct_count}/5 个段落样式正确")
        return False

    # Test 2: Invalid style name
    print("\n测试: 无效样式名称")
    result = dispatch_tool("batch_apply_style", {
        "path": test_file,
        "paragraph_indices": [0, 1],
        "style_name": "NonExistentStyle"
    })

    if not result.get("success"):
        print(f"[OK] 正确检测到错误: {result.get('error')}")
    else:
        print(f"[FAIL] 应该返回错误但返回成功")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] batch_apply_style 测试通过")
    return True


def test_batch_copy_format():
    """Test batch_copy_format tool (format painter)"""
    print_section("Test 3: batch_copy_format (格式刷)")

    test_file = "E:/Project/DocuFlow/test_batch_copy.docx"

    # Create test document
    doc = Document()

    # Source paragraph with specific formatting
    source_para = doc.add_paragraph("源段落 - 特殊格式")
    source_para.runs[0].font.name = "黑体"
    source_para.runs[0].font.size = Pt(16)
    source_para.runs[0].font.bold = True
    source_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_para.paragraph_format.line_spacing = 2.0

    # Target paragraphs with different formatting
    for i in range(5):
        para = doc.add_paragraph(f"目标段落 {i+1}")
        para.runs[0].font.name = "宋体"
        para.runs[0].font.size = Pt(12)

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test: Copy format from paragraph 0 to paragraphs 1-5
    print("\n测试: 复制段落0的格式到段落1-5")
    result = dispatch_tool("batch_copy_format", {
        "path": test_file,
        "source_index": 0,
        "target_indices": [1, 2, 3, 4, 5]
    })

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     复制数量: {result['copied_count']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Verify formatting was copied
    doc = Document(test_file)
    source_para = doc.paragraphs[0]
    target_para = doc.paragraphs[1]

    checks = []
    checks.append(("字体", source_para.runs[0].font.name == target_para.runs[0].font.name))
    checks.append(("字号", source_para.runs[0].font.size == target_para.runs[0].font.size))
    checks.append(("加粗", source_para.runs[0].font.bold == target_para.runs[0].font.bold))
    checks.append(("对齐", source_para.paragraph_format.alignment == target_para.paragraph_format.alignment))
    checks.append(("行距", source_para.paragraph_format.line_spacing == target_para.paragraph_format.line_spacing))

    all_passed = all(check[1] for check in checks)
    if all_passed:
        print(f"[OK] 验证: 所有格式属性已正确复制")
        for name, _ in checks:
            print(f"     - {name}: 匹配")
    else:
        print(f"[FAIL] 验证失败:")
        for name, passed in checks:
            print(f"     - {name}: {'匹配' if passed else '不匹配'}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] batch_copy_format 测试通过")
    return True


def test_batch_replace_format():
    """Test batch_replace_format tool"""
    print_section("Test 4: batch_replace_format")

    test_file = "E:/Project/DocuFlow/test_batch_replace.docx"

    # Create test document with mixed styles
    doc = Document()
    doc.add_paragraph("标题1", style="Heading 1")
    doc.add_paragraph("正文内容1", style="Normal")
    doc.add_paragraph("标题2", style="Heading 1")
    doc.add_paragraph("正文内容2", style="Normal")
    doc.add_paragraph("标题3", style="Heading 1")
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file} (3个Heading 1, 2个Normal)")

    # Test: Replace all Heading 1 paragraphs with new formatting
    print("\n测试: 查找并替换所有 Heading 1 样式的格式")
    result = dispatch_tool("batch_replace_format", {
        "path": test_file,
        "find_style": "Heading 1",
        "replace_options": {
            "font_name": "黑体",
            "font_size": "18pt",
            "font_color": "#FF0000",
            "bold": True,
            "alignment": "center"
        }
    })

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     替换数量: {result['replaced_count']}")
        print(f"     匹配索引: {result['matched_indices']}")

        if result['replaced_count'] != 3:
            print(f"[FAIL] 期望替换3个段落，实际替换 {result['replaced_count']} 个")
            return False
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Search for style with no matches
    print("\n测试: 搜索不存在的样式段落")
    result = dispatch_tool("batch_replace_format", {
        "path": test_file,
        "find_style": "Heading 2",
        "replace_options": {"bold": True}
    })

    if result.get("success") and result['replaced_count'] == 0:
        print(f"[OK] 正确处理: {result['message']}")
    else:
        print(f"[FAIL] 处理错误")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] batch_replace_format 测试通过")
    return True


def test_performance_100_paragraphs():
    """Test performance with 100 paragraphs"""
    print_section("Test 5: 性能测试 (100段落)")

    test_file = "E:/Project/DocuFlow/test_performance.docx"

    # Create document with 100 paragraphs
    print("创建包含100个段落的文档...")
    doc = Document()
    for i in range(100):
        doc.add_paragraph(f"这是第 {i+1} 个段落的内容。Lorem ipsum dolor sit amet.")
    doc.save(test_file)
    print(f"[OK] 文档已创建")

    # Test 1: Format all 100 paragraphs
    print("\n测试1: 批量格式化所有100个段落")
    start_time = time.time()

    result = dispatch_tool("batch_format_range", {
        "path": test_file,
        "start_index": 0,
        "end_index": 99,
        "font_name": "宋体",
        "font_size": "12pt",
        "alignment": "justify",
        "line_spacing": 1.5
    })

    elapsed = (time.time() - start_time) * 1000

    if result.get("success"):
        print(f"[OK] 完成: {result['formatted_count']} 段落")
        print(f"     性能: {elapsed:.2f}ms")
        print(f"     平均: {elapsed/100:.2f}ms/段落")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        os.remove(test_file)
        return False

    # Test 2: Apply style to 50 paragraphs (every other paragraph)
    print("\n测试2: 批量应用样式到50个段落")
    indices = list(range(0, 100, 2))  # 0, 2, 4, ..., 98

    start_time = time.time()

    result = dispatch_tool("batch_apply_style", {
        "path": test_file,
        "paragraph_indices": indices,
        "style_name": "Heading 1"
    })

    elapsed = (time.time() - start_time) * 1000

    if result.get("success"):
        print(f"[OK] 完成: {result['applied_count']} 段落")
        print(f"     性能: {elapsed:.2f}ms")
        print(f"     平均: {elapsed/50:.2f}ms/段落")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        os.remove(test_file)
        return False

    # Test 3: Copy format to 49 paragraphs
    print("\n测试3: 格式刷复制到49个段落")
    target_indices = list(range(1, 100, 2))  # 1, 3, 5, ..., 99

    start_time = time.time()

    result = dispatch_tool("batch_copy_format", {
        "path": test_file,
        "source_index": 0,
        "target_indices": target_indices
    })

    elapsed = (time.time() - start_time) * 1000

    if result.get("success"):
        print(f"[OK] 完成: {result['copied_count']} 段落")
        print(f"     性能: {elapsed:.2f}ms")
        print(f"     平均: {elapsed/49:.2f}ms/段落")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        os.remove(test_file)
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] 性能测试通过")
    return True


def test_tool_registration():
    """Test that all batch tools are registered"""
    print_section("Test 6: 工具注册验证")

    tools = get_all_registered_tools()
    batch_tools = [
        "batch_format_range",
        "batch_apply_style",
        "batch_copy_format",
        "batch_replace_format"
    ]

    print(f"总工具数: {len(tools)}")
    print("\n检查批量操作工具:")

    all_found = True
    for tool in batch_tools:
        if tool in tools:
            print(f"  [OK] {tool}")
        else:
            print(f"  [FAIL] {tool} - 未找到")
            all_found = False

    if all_found:
        print(f"\n[OK] 所有4个批量操作工具已正确注册")
        return True
    else:
        print(f"\n[FAIL] 部分工具未注册")
        return False


def main():
    """Run all tests"""
    # Import necessary modules for testing
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Make them available globally for test functions
    globals()['Pt'] = Pt
    globals()['WD_ALIGN_PARAGRAPH'] = WD_ALIGN_PARAGRAPH

    print("\n" + "="*60)
    print("  DocuFlow MCP - Day 6-7 批量操作测试")
    print("="*60)

    # Import batch module to trigger registration
    from docuflow_mcp.extensions import batch

    tests = [
        ("工具注册验证", test_tool_registration),
        ("batch_format_range", test_batch_format_range),
        ("batch_apply_style", test_batch_apply_style),
        ("batch_copy_format", test_batch_copy_format),
        ("batch_replace_format", test_batch_replace_format),
        ("性能测试", test_performance_100_paragraphs)
    ]

    results = []
    total_start = time.time()

    for name, test_func in tests:
        try:
            result = test_func()
            results.append((name, result))
        except Exception as e:
            print(f"\n[ERROR] 测试 '{name}' 出现异常: {str(e)}")
            import traceback
            traceback.print_exc()
            results.append((name, False))

    total_elapsed = time.time() - total_start

    # Print summary
    print_section("测试总结")

    passed = sum(1 for _, result in results if result)
    total = len(results)

    print(f"总测试数: {total}")
    print(f"通过: {passed}")
    print(f"失败: {total - passed}")
    print(f"总耗时: {total_elapsed:.2f}秒")
    print()

    for name, result in results:
        status = "[OK]" if result else "[FAIL]"
        print(f"  {status} {name}")

    print()
    if passed == total:
        print("="*60)
        print("  所有测试通过！Day 6-7 批量操作功能正常工作")
        print("="*60)
        return 0
    else:
        print("="*60)
        print(f"  {total - passed} 个测试失败")
        print("="*60)
        return 1


if __name__ == "__main__":
    exit(main())
