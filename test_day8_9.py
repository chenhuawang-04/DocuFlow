"""
DocuFlow MCP - Day 8-9 Format Validation Test

Tests all format validation tools including:
- validate_format
- validate_auto_fix
- validate_generate_report
- validate_check_consistency
"""

import sys
import os
import time

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from docuflow_mcp.core.registry import dispatch_tool, get_all_registered_tools
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def print_section(title):
    """Print a section header"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}\n")


def test_validate_format():
    """Test validate_format tool"""
    print_section("Test 1: validate_format")

    test_file = "E:/Project/DocuFlow/test_validate.docx"

    # Create test document with non-compliant format
    doc = Document()

    # Set wrong margins
    section = doc.sections[0]
    section.top_margin = Cm(3)  # Should be 2cm for MBA
    section.left_margin = Cm(2)  # Should be 3.5cm for MBA

    # Add paragraphs with wrong formatting
    para = doc.add_paragraph("正文段落")
    para.runs[0].font.name = "Arial"  # Should be 宋体
    para.runs[0].font.size = Pt(14)  # Should be 12pt

    para2 = doc.add_heading("标题1", level=1)
    para2.runs[0].font.name = "宋体"  # Should be 黑体

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test 1: Validate with preset rules
    print("\n测试1: 使用预设规则验证 (mba_thesis)")
    result = dispatch_tool("validate_format", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })

    if result.get("success"):
        print(f"[OK] 验证完成")
        print(f"     符合规范: {result['compliant']}")
        print(f"     总问题数: {result['total_issues']}")
        print(f"     错误: {result['errors']}, 警告: {result['warnings']}, 信息: {result['info']}")

        if result['total_issues'] > 0:
            print(f"     发现的问题:")
            for issue in result['issues'][:3]:  # Show first 3 issues
                print(f"       - {issue.get('type')}: {issue.get('expected', '')} vs {issue.get('actual', '')}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Validate with custom rules
    print("\n测试2: 使用自定义规则验证")
    custom_rules = {
        "page_setup": {
            "margins": {
                "top": "2cm",
                "tolerance": "0.1cm"
            }
        },
        "styles": {
            "Normal": {
                "font": {
                    "name": "宋体",
                    "size": "12pt"
                }
            }
        }
    }

    result = dispatch_tool("validate_format", {
        "path": test_file,
        "rules": custom_rules
    })

    if result.get("success"):
        print(f"[OK] 验证完成")
        print(f"     总问题数: {result['total_issues']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] validate_format 测试通过")
    return True


def test_validate_auto_fix():
    """Test validate_auto_fix tool"""
    print_section("Test 2: validate_auto_fix")

    test_file = "E:/Project/DocuFlow/test_autofix.docx"

    # Create test document with issues
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3)  # Wrong
    section.left_margin = Cm(2)  # Wrong

    para = doc.add_paragraph("测试段落")
    para.runs[0].font.name = "Arial"  # Wrong
    para.runs[0].font.size = Pt(14)  # Wrong

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Validate before fix
    print("\n验证修复前的问题:")
    result = dispatch_tool("validate_format", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })
    issues_before = result.get('total_issues', 0)
    print(f"[INFO] 修复前问题数: {issues_before}")

    # Auto fix
    print("\n测试: 自动修正格式问题")
    result = dispatch_tool("validate_auto_fix", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })

    if result.get("success"):
        print(f"[OK] {result['message']}")
        print(f"     修正数量: {result['fixed_count']}")
        if result['fixes_applied']:
            print(f"     应用的修正: {', '.join(result['fixes_applied'][:5])}...")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Validate after fix
    print("\n验证修复后:")
    result = dispatch_tool("validate_format", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })
    issues_after = result.get('total_issues', 0)
    print(f"[INFO] 修复后问题数: {issues_after}")

    if issues_after < issues_before:
        print(f"[OK] 问题数已减少: {issues_before} -> {issues_after}")
    else:
        print(f"[FAIL] 问题数未减少")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] validate_auto_fix 测试通过")
    return True


def test_validate_generate_report():
    """Test validate_generate_report tool"""
    print_section("Test 3: validate_generate_report")

    test_file = "E:/Project/DocuFlow/test_report.docx"
    report_file = "E:/Project/DocuFlow/test_report.txt"

    # Create test document
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3)  # Wrong

    para = doc.add_paragraph("测试内容")
    para.runs[0].font.name = "Arial"  # Wrong

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test 1: Generate report without saving
    print("\n测试1: 生成报告（不保存）")
    result = dispatch_tool("validate_generate_report", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })

    if result.get("success"):
        print(f"[OK] 报告已生成")
        print(f"     符合规范: {result['compliant']}")
        print(f"     总问题数: {result['total_issues']}")
        print(f"\n报告预览 (前5行):")
        report_lines = result['report'].split('\n')[:5]
        for line in report_lines:
            print(f"     {line}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Generate and save report
    print("\n测试2: 生成报告并保存")
    result = dispatch_tool("validate_generate_report", {
        "path": test_file,
        "preset_rules": "mba_thesis",
        "output_path": report_file
    })

    if result.get("success") and os.path.exists(report_file):
        print(f"[OK] 报告已保存到: {report_file}")
        print(f"     文件大小: {os.path.getsize(report_file)} 字节")
    else:
        print(f"[FAIL] 报告保存失败")
        return False

    # Clean up
    os.remove(test_file)
    os.remove(report_file)
    print("\n[OK] validate_generate_report 测试通过")
    return True


def test_validate_check_consistency():
    """Test validate_check_consistency tool"""
    print_section("Test 4: validate_check_consistency")

    test_file = "E:/Project/DocuFlow/test_consistency.docx"

    # Create document with inconsistent formatting
    doc = Document()

    # Most paragraphs use 宋体 12pt
    for i in range(15):
        para = doc.add_paragraph(f"正常段落 {i+1}")
        para.runs[0].font.name = "宋体"
        para.runs[0].font.size = Pt(12)

    # Few paragraphs use different fonts (inconsistent)
    para = doc.add_paragraph("异常段落1")
    para.runs[0].font.name = "Arial"
    para.runs[0].font.size = Pt(14)

    para = doc.add_paragraph("异常段落2")
    para.runs[0].font.name = "黑体"
    para.runs[0].font.size = Pt(16)

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file} (17段落)")

    # Test: Check consistency
    print("\n测试: 检查格式一致性")
    result = dispatch_tool("validate_check_consistency", {
        "path": test_file
    })

    if result.get("success"):
        print(f"[OK] 一致性检查完成")
        print(f"     格式一致: {result['consistent']}")
        print(f"     发现问题: {result['total_issues']}")

        if result.get('statistics'):
            stats = result['statistics']
            print(f"\n     统计信息:")
            print(f"       字体使用: {stats.get('font_usage', {})}")
            print(f"       字号使用: {stats.get('font_size_usage', {})}")

        if result.get('issues'):
            print(f"\n     一致性问题:")
            for issue in result['issues']:
                print(f"       - {issue.get('type')}: {issue.get('message')}")
                if issue.get('details'):
                    details = issue['details']
                    if isinstance(details, dict):
                        for k, v in list(details.items())[:3]:
                            print(f"         · {k}: {v}次")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] validate_check_consistency 测试通过")
    return True


def test_tool_registration():
    """Test that all validation tools are registered"""
    print_section("Test 5: 工具注册验证")

    tools = get_all_registered_tools()
    validation_tools = [
        "validate_format",
        "validate_auto_fix",
        "validate_generate_report",
        "validate_check_consistency"
    ]

    print(f"总工具数: {len(tools)}")
    print("\n检查验证工具:")

    all_found = True
    for tool in validation_tools:
        if tool in tools:
            print(f"  [OK] {tool}")
        else:
            print(f"  [FAIL] {tool} - 未找到")
            all_found = False

    if all_found:
        print(f"\n[OK] 所有4个验证工具已正确注册")
        return True
    else:
        print(f"\n[FAIL] 部分工具未注册")
        return False


def test_integration():
    """Integration test: Full workflow"""
    print_section("Test 6: 集成测试 - 完整工作流")

    test_file = "E:/Project/DocuFlow/test_integration.docx"

    # Step 1: Create document from MBA template
    print("步骤1: 从MBA模板创建文档")
    result = dispatch_tool("template_create_from_preset", {
        "preset_name": "mba_thesis",
        "output_path": test_file,
        "title": "集成测试文档"
    })

    if not result.get("success"):
        print(f"[FAIL] 模板创建失败: {result.get('error')}")
        return False
    print(f"[OK] 文档已创建")

    # Step 2: Add some content
    print("\n步骤2: 添加内容")
    doc = Document(test_file)
    doc.add_paragraph("这是第一段正文内容。")
    doc.add_paragraph("这是第二段正文内容。")
    doc.add_heading("第一章", level=1)
    doc.add_paragraph("章节内容。")
    doc.save(test_file)
    print(f"[OK] 内容已添加")

    # Step 3: Validate format
    print("\n步骤3: 验证格式")
    result = dispatch_tool("validate_format", {
        "path": test_file,
        "preset_rules": "mba_thesis"
    })

    if result.get("success"):
        print(f"[OK] 验证完成: 符合={result['compliant']}, 问题={result['total_issues']}")

        # Step 4: Auto fix if needed
        if not result['compliant']:
            print("\n步骤4: 自动修正格式")
            fix_result = dispatch_tool("validate_auto_fix", {
                "path": test_file,
                "preset_rules": "mba_thesis"
            })
            print(f"[OK] 修正完成: {fix_result.get('fixed_count', 0)} 个问题")

            # Step 5: Re-validate
            print("\n步骤5: 重新验证")
            result = dispatch_tool("validate_format", {
                "path": test_file,
                "preset_rules": "mba_thesis"
            })
            print(f"[OK] 重新验证: 符合={result['compliant']}, 问题={result['total_issues']}")

        # Step 6: Check consistency
        print("\n步骤6: 检查一致性")
        result = dispatch_tool("validate_check_consistency", {
            "path": test_file
        })
        print(f"[OK] 一致性: {result['consistent']}, 问题={result['total_issues']}")

    else:
        print(f"[FAIL] 验证失败")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] 集成测试通过")
    return True


def main():
    """Run all tests"""
    print("\n" + "="*60)
    print("  DocuFlow MCP - Day 8-9 格式验证测试")
    print("="*60)

    # Import validator module to trigger registration
    from docuflow_mcp.extensions import validator

    tests = [
        ("工具注册验证", test_tool_registration),
        ("validate_format", test_validate_format),
        ("validate_auto_fix", test_validate_auto_fix),
        ("validate_generate_report", test_validate_generate_report),
        ("validate_check_consistency", test_validate_check_consistency),
        ("集成测试", test_integration)
    ]

    results = []
    total_start = time.time()

    for name, test_func in tests:
        try:
            result = test_func()
            results.append((name, result))
        except Exception as e:
            print(f"\n[ERROR] 测试 '{name}' 出现异常: {str(e)}")
            import traceback
            traceback.print_exc()
            results.append((name, False))

    total_elapsed = time.time() - total_start

    # Print summary
    print_section("测试总结")

    passed = sum(1 for _, result in results if result)
    total = len(results)

    print(f"总测试数: {total}")
    print(f"通过: {passed}")
    print(f"失败: {total - passed}")
    print(f"总耗时: {total_elapsed:.2f}秒")
    print()

    for name, result in results:
        status = "[OK]" if result else "[FAIL]"
        print(f"  {status} {name}")

    print()
    if passed == total:
        print("="*60)
        print("  所有测试通过！Day 8-9 格式验证功能正常工作")
        print("="*60)
        return 0
    else:
        print("="*60)
        print(f"  {total - passed} 个测试失败")
        print("="*60)
        return 1


if __name__ == "__main__":
    exit(main())
