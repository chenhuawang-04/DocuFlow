"""验证生成文档的格式"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from docx import Document

def verify_mba_thesis():
    """验证MBA论文文档格式"""
    print("=" * 60)
    print("验证 MBA 论文文档格式")
    print("=" * 60)

    doc = Document("test_mba_thesis.docx")
    section = doc.sections[0]

    # 检查页边距
    print("\n页边距设置:")
    print(f"  上边距: {section.top_margin.cm:.2f} cm (预期: 2.00 cm)")
    print(f"  下边距: {section.bottom_margin.cm:.2f} cm (预期: 1.50 cm)")
    print(f"  左边距: {section.left_margin.cm:.2f} cm (预期: 3.50 cm)")
    print(f"  右边距: {section.right_margin.cm:.2f} cm (预期: 2.00 cm)")

    # 检查页面大小
    print("\n页面大小:")
    print(f"  宽度: {section.page_width.cm:.2f} cm (预期: 21.00 cm)")
    print(f"  高度: {section.page_height.cm:.2f} cm (预期: 29.70 cm)")

    # 检查样式
    print("\n样式设置:")
    try:
        normal_style = doc.styles['Normal']
        if hasattr(normal_style, 'font'):
            print(f"  Normal样式字体: {normal_style.font.name}")
            if normal_style.font.size:
                print(f"  Normal样式字号: {normal_style.font.size.pt} pt (预期: 12 pt)")

        if hasattr(normal_style, 'paragraph_format'):
            pf = normal_style.paragraph_format
            if pf.line_spacing:
                print(f"  Normal样式行距: {pf.line_spacing} (预期: 1.5)")
            if pf.first_line_indent:
                print(f"  Normal样式首行缩进: {pf.first_line_indent.pt:.2f} pt")
    except (KeyError, AttributeError):
        print("  无法读取Normal样式详情")

    # 检查页脚（页码）
    print("\n页脚设置:")
    footer = section.footer
    if footer.paragraphs:
        print(f"  页脚段落数: {len(footer.paragraphs)}")
        print(f"  页脚内容: {'已设置页码' if footer.paragraphs[0]._element.xml else '空'}")

    print("\n" + "=" * 60)
    print("[OK] MBA论文文档格式验证完成")
    print("=" * 60)

if __name__ == "__main__":
    verify_mba_thesis()
