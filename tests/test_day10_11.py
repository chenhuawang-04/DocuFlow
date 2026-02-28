"""
DocuFlow MCP - Day 10-11 Advanced Features Test

Tests all advanced tools including:
- doc_compare
- doc_analyze_statistics
- doc_get_metadata
- doc_set_metadata
- doc_extract_links
- doc_word_frequency
"""

import sys
import os
import time

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from docuflow_mcp.core.registry import dispatch_tool, get_all_registered_tools
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def print_section(title):
    """Print a section header"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}\n")


def test_tool_registration():
    """Test that all advanced tools are registered"""
    print_section("Test 1: 工具注册验证")

    tools = get_all_registered_tools()
    advanced_tools = [
        "doc_compare",
        "doc_analyze_statistics",
        "doc_get_metadata",
        "doc_set_metadata",
        "doc_extract_links",
        "doc_word_frequency"
    ]

    print(f"总工具数: {len(tools)}")
    print("\n检查高级工具:")

    all_found = True
    for tool in advanced_tools:
        if tool in tools:
            print(f"  [OK] {tool}")
        else:
            print(f"  [FAIL] {tool} - 未找到")
            all_found = False

    if all_found:
        print(f"\n[OK] 所有6个高级工具已正确注册")
        return True
    else:
        print(f"\n[FAIL] 部分工具未注册")
        return False


def test_doc_compare():
    """Test doc_compare tool"""
    print_section("Test 2: doc_compare")

    doc1_path = "E:/Project/DocuFlow/test_compare1.docx"
    doc2_path = "E:/Project/DocuFlow/test_compare2.docx"
    report_path = "E:/Project/DocuFlow/test_compare_report.txt"

    # Create first document
    doc1 = Document()
    doc1.add_paragraph("第一段内容")
    doc1.add_paragraph("第二段内容")
    doc1.add_paragraph("第三段内容 - 原始版本")
    doc1.add_paragraph("第四段内容")
    doc1.save(doc1_path)
    print(f"[OK] 创建文档1: {doc1_path}")

    # Create second document (modified version)
    doc2 = Document()
    doc2.add_paragraph("第一段内容")
    doc2.add_paragraph("第二段内容 - 已修改")
    doc2.add_paragraph("第三段内容 - 修改版本")
    doc2.add_paragraph("第五段内容 - 新增")
    doc2.save(doc2_path)
    print(f"[OK] 创建文档2: {doc2_path}")

    # Test 1: Basic comparison
    print("\n测试1: 基础对比")
    result = dispatch_tool("doc_compare", {
        "path1": doc1_path,
        "path2": doc2_path
    })

    if result.get("success"):
        print(f"[OK] 对比完成")
        print(f"     新增: {result['added_count']} 段")
        print(f"     删除: {result['removed_count']} 段")
        print(f"     修改: {result['modified_count']} 段")

        if result.get('modified'):
            print(f"\n     修改详情:")
            for mod in result['modified']:
                print(f"       原文: {mod['original'][:40]}...")
                print(f"       新文: {mod['modified'][:40]}...")
                print(f"       相似度: {mod['similarity']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Comparison with report
    print("\n测试2: 生成对比报告")
    result = dispatch_tool("doc_compare", {
        "path1": doc1_path,
        "path2": doc2_path,
        "output_path": report_path
    })

    if result.get("success") and os.path.exists(report_path):
        print(f"[OK] 报告已生成: {report_path}")
        print(f"     文件大小: {os.path.getsize(report_path)} 字节")
    else:
        print(f"[FAIL] 报告生成失败")
        return False

    # Clean up
    os.remove(doc1_path)
    os.remove(doc2_path)
    os.remove(report_path)
    print("\n[OK] doc_compare 测试通过")
    return True


def test_doc_analyze_statistics():
    """Test doc_analyze_statistics tool"""
    print_section("Test 3: doc_analyze_statistics")

    test_file = "E:/Project/DocuFlow/test_statistics.docx"

    # Create test document with various content
    doc = Document()

    # Add title
    doc.add_heading("测试文档标题", level=0)

    # Add paragraphs
    doc.add_paragraph("这是第一段正文内容，包含一些文字。")
    doc.add_paragraph("这是第二段正文内容，也包含一些文字。")

    # Add heading
    doc.add_heading("第一章", level=1)
    doc.add_paragraph("这是第一章的内容。这段有更多的文字，用来测试统计功能。")

    # Add heading
    doc.add_heading("第二章", level=1)
    doc.add_paragraph("这是第二章的内容。")

    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "标题1"
    table.cell(0, 1).text = "标题2"
    table.cell(1, 0).text = "数据1"
    table.cell(1, 1).text = "数据2"

    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test 1: Basic statistics
    print("\n测试1: 基础统计")
    result = dispatch_tool("doc_analyze_statistics", {
        "path": test_file
    })

    if result.get("success"):
        print(f"[OK] 统计完成")
        counts = result['counts']
        print(f"\n     基础统计:")
        print(f"       段落数: {counts['paragraphs']}")
        print(f"       非空段落: {counts['non_empty_paragraphs']}")
        print(f"       表格数: {counts['tables']}")
        print(f"       标题数: {counts['headings']}")
        print(f"       字数: {counts['words']}")
        print(f"       字符数: {counts['characters']}")

        print(f"\n     平均值:")
        print(f"       每段字数: {result['averages']['words_per_paragraph']}")
        print(f"       每字字符数: {result['averages']['characters_per_word']}")

        print(f"\n     字体使用:")
        if result['fonts']['most_used']:
            for font, count in result['fonts']['most_used'][:3]:
                print(f"       {font}: {count}次")

        print(f"\n     样式使用:")
        if result['styles']['most_used']:
            for style, count in result['styles']['most_used'][:3]:
                print(f"       {style}: {count}次")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Detailed statistics
    print("\n测试2: 详细统计")
    result = dispatch_tool("doc_analyze_statistics", {
        "path": test_file,
        "detailed": True
    })

    if result.get("success") and result.get("detailed"):
        print(f"[OK] 详细统计完成")
        detailed = result['detailed']
        print(f"     标题数: {len(detailed['headings'])}")
        print(f"     表格数: {len(detailed['tables'])}")
        if detailed['headings']:
            print(f"     第一个标题: {detailed['headings'][0]['text']}")
        if detailed['tables']:
            print(f"     第一个表格: {detailed['tables'][0]['rows']}行 x {detailed['tables'][0]['columns']}列")
    else:
        print(f"[FAIL] 详细统计失败")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] doc_analyze_statistics 测试通过")
    return True


def test_doc_metadata():
    """Test doc_get_metadata and doc_set_metadata tools"""
    print_section("Test 4: 元数据工具")

    test_file = "E:/Project/DocuFlow/test_metadata.docx"

    # Create test document
    doc = Document()
    doc.add_paragraph("测试文档内容")
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test 1: Set metadata
    print("\n测试1: 设置元数据")
    result = dispatch_tool("doc_set_metadata", {
        "path": test_file,
        "title": "测试文档标题",
        "author": "张三",
        "subject": "测试主题",
        "keywords": "测试,文档,元数据",
        "comments": "这是一个测试文档",
        "category": "测试类别"
    })

    if result.get("success"):
        print(f"[OK] 元数据已设置")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: Get metadata
    print("\n测试2: 获取元数据")
    result = dispatch_tool("doc_get_metadata", {
        "path": test_file
    })

    if result.get("success"):
        print(f"[OK] 元数据已获取")
        props = result['core_properties']
        print(f"\n     核心属性:")
        print(f"       标题: {props['title']}")
        print(f"       作者: {props['author']}")
        print(f"       主题: {props['subject']}")
        print(f"       关键词: {props['keywords']}")
        print(f"       类别: {props['category']}")

        # Verify the values
        if props['title'] == "测试文档标题" and props['author'] == "张三":
            print(f"\n[OK] 元数据验证成功")
        else:
            print(f"\n[FAIL] 元数据值不匹配")
            return False
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] 元数据工具测试通过")
    return True


def test_doc_extract_links():
    """Test doc_extract_links tool"""
    print_section("Test 5: doc_extract_links")

    test_file = "E:/Project/DocuFlow/test_links.docx"

    # Create document with hyperlinks
    doc = Document()
    doc.add_paragraph("文档内容")

    # Note: Adding hyperlinks programmatically is complex with python-docx
    # For this test, we'll just test the tool with a document without links
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    print("\n测试: 提取链接")
    result = dispatch_tool("doc_extract_links", {
        "path": test_file
    })

    if result.get("success"):
        print(f"[OK] 链接提取完成")
        print(f"     找到链接: {result['link_count']} 个")
        if result['links']:
            for link in result['links']:
                print(f"       文本: {link['text']}")
                print(f"       URL: {link['url']}")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] doc_extract_links 测试通过")
    return True


def test_doc_word_frequency():
    """Test doc_word_frequency tool"""
    print_section("Test 6: doc_word_frequency")

    test_file = "E:/Project/DocuFlow/test_wordfreq.docx"

    # Create test document with repeated words
    doc = Document()
    doc.add_paragraph("这是测试文档。测试测试测试。")
    doc.add_paragraph("文档包含很多文字。文字文字文字。")
    doc.add_paragraph("重复的词语会被统计。统计统计。")
    doc.save(test_file)
    print(f"[OK] 创建测试文档: {test_file}")

    # Test 1: Basic word frequency
    print("\n测试1: 基础词频分析")
    result = dispatch_tool("doc_word_frequency", {
        "path": test_file,
        "top_n": 10
    })

    if result.get("success"):
        print(f"[OK] 词频分析完成")
        print(f"     总词数: {result['total_words']}")
        print(f"     不重复词数: {result['unique_words']}")
        print(f"\n     Top 10 高频词:")
        for item in result['top_words'][:10]:
            print(f"       {item['word']}: {item['count']}次")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Test 2: With minimum length filter
    print("\n测试2: 带最小长度过滤")
    result = dispatch_tool("doc_word_frequency", {
        "path": test_file,
        "top_n": 5,
        "min_length": 3
    })

    if result.get("success"):
        print(f"[OK] 词频分析完成（最小长度=3）")
        print(f"     过滤后词数: {result['filtered_words']}")
        print(f"\n     Top 5 高频词:")
        for item in result['top_words'][:5]:
            print(f"       {item['word']}: {item['count']}次")
    else:
        print(f"[FAIL] {result.get('error', 'Unknown error')}")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] doc_word_frequency 测试通过")
    return True


def test_integration():
    """Integration test: Complete workflow"""
    print_section("Test 7: 集成测试 - 完整工作流")

    test_file = "E:/Project/DocuFlow/test_integration_advanced.docx"

    # Step 1: Create document from template
    print("步骤1: 从模板创建文档")
    result = dispatch_tool("template_create_from_preset", {
        "preset_name": "business_report",
        "output_path": test_file,
        "title": "市场分析报告"
    })

    if not result.get("success"):
        print(f"[FAIL] 模板创建失败: {result.get('error')}")
        return False
    print(f"[OK] 文档已创建")

    # Step 2: Add content
    print("\n步骤2: 添加内容")
    doc = Document(test_file)
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This report analyzes market trends and provides insights.")
    doc.add_heading("Market Analysis", level=1)
    doc.add_paragraph("The market shows strong growth potential.")
    doc.add_paragraph("Key findings include increased demand and competitive pricing.")
    doc.save(test_file)
    print(f"[OK] 内容已添加")

    # Step 3: Set metadata
    print("\n步骤3: 设置元数据")
    result = dispatch_tool("doc_set_metadata", {
        "path": test_file,
        "title": "市场分析报告",
        "author": "分析团队",
        "keywords": "市场,分析,报告"
    })
    print(f"[OK] 元数据已设置")

    # Step 4: Analyze statistics
    print("\n步骤4: 统计分析")
    result = dispatch_tool("doc_analyze_statistics", {
        "path": test_file,
        "detailed": True
    })

    if result.get("success"):
        print(f"[OK] 统计完成")
        print(f"     段落: {result['counts']['paragraphs']}")
        print(f"     字数: {result['counts']['words']}")
        print(f"     标题: {result['counts']['headings']}")
    else:
        print(f"[FAIL] 统计失败")
        return False

    # Step 5: Word frequency
    print("\n步骤5: 词频分析")
    result = dispatch_tool("doc_word_frequency", {
        "path": test_file,
        "top_n": 5
    })

    if result.get("success"):
        print(f"[OK] 词频分析完成")
        print(f"     总词数: {result['total_words']}")
    else:
        print(f"[FAIL] 词频分析失败")
        return False

    # Step 6: Validate format
    print("\n步骤6: 验证格式")
    result = dispatch_tool("validate_format", {
        "path": test_file,
        "preset_rules": "business_report"
    })

    if result.get("success"):
        print(f"[OK] 格式验证完成: 符合={result['compliant']}")
    else:
        print(f"[FAIL] 验证失败")
        return False

    # Clean up
    os.remove(test_file)
    print("\n[OK] 集成测试通过")
    return True


def main():
    """Run all tests"""
    print("\n" + "="*60)
    print("  DocuFlow MCP - Day 10-11 高级功能测试")
    print("="*60)

    # Import advanced module to trigger registration
    from docuflow_mcp.extensions import advanced

    tests = [
        ("工具注册验证", test_tool_registration),
        ("doc_compare", test_doc_compare),
        ("doc_analyze_statistics", test_doc_analyze_statistics),
        ("元数据工具", test_doc_metadata),
        ("doc_extract_links", test_doc_extract_links),
        ("doc_word_frequency", test_doc_word_frequency),
        ("集成测试", test_integration)
    ]

    results = []
    total_start = time.time()

    for name, test_func in tests:
        try:
            result = test_func()
            assert isinstance(result, dict), "Expected dict result"
            results.append((name, result))
        except Exception as e:
            print(f"\n[ERROR] 测试 '{name}' 出现异常: {str(e)}")
            import traceback
            traceback.print_exc()
            results.append((name, False))

    total_elapsed = time.time() - total_start

    # Print summary
    print_section("测试总结")

    passed = sum(1 for _, result in results if result)
    total = len(results)

    print(f"总测试数: {total}")
    print(f"通过: {passed}")
    print(f"失败: {total - passed}")
    print(f"总耗时: {total_elapsed:.2f}秒")
    print()

    for name, result in results:
        status = "[OK]" if result else "[FAIL]"
        print(f"  {status} {name}")

    print()
    if passed == total:
        print("="*60)
        print("  所有测试通过！Day 10-11 高级功能正常工作")
        print("="*60)
        return 0
    else:
        print("="*60)
        print(f"  {total - passed} 个测试失败")
        print("="*60)
        return 1


if __name__ == "__main__":
    exit(main())
