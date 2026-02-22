"""
DocuFlow MCP - 文档操作核心模块

提供 Word 文档的所有操作功能
"""

import os
import json
import copy
import re
from typing import Optional, List, Dict, Any, Union, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from docuflow_mcp.utils.formatters import apply_font_format, apply_paragraph_format
from docuflow_mcp.core.registry import register_tool


# ============================================================
# 辅助函数
# ============================================================

def parse_color(color_str: str) -> RGBColor:
    """解析颜色字符串，支持 hex (#FF0000) 或 rgb(255,0,0) 格式"""
    if not color_str:
        return None

    color_str = color_str.strip()

    # Hex 格式
    if color_str.startswith('#'):
        hex_color = color_str[1:]
        if len(hex_color) == 3:
            hex_color = ''.join([c*2 for c in hex_color])
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)

    # RGB 格式
    if color_str.startswith('rgb'):
        match = re.match(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
        if match:
            r, g, b = map(int, match.groups())
            return RGBColor(r, g, b)

    # 预定义颜色
    colors = {
        'red': RGBColor(255, 0, 0),
        'green': RGBColor(0, 128, 0),
        'blue': RGBColor(0, 0, 255),
        'black': RGBColor(0, 0, 0),
        'white': RGBColor(255, 255, 255),
        'yellow': RGBColor(255, 255, 0),
        'orange': RGBColor(255, 165, 0),
        'purple': RGBColor(128, 0, 128),
        'gray': RGBColor(128, 128, 128),
        'grey': RGBColor(128, 128, 128),
        'pink': RGBColor(255, 192, 203),
        'brown': RGBColor(165, 42, 42),
        'cyan': RGBColor(0, 255, 255),
        'magenta': RGBColor(255, 0, 255),
    }

    return colors.get(color_str.lower())


def parse_size(size_str: str) -> Union[Pt, Inches, Cm, Mm]:
    """解析尺寸字符串，支持 pt, in, cm, mm 单位"""
    if isinstance(size_str, (int, float)):
        return Pt(size_str)

    size_str = str(size_str).strip().lower()

    match = re.match(r'([\d.]+)\s*(pt|in|inch|inches|cm|mm)?', size_str)
    if match:
        value = float(match.group(1))
        unit = match.group(2) or 'pt'

        if unit == 'pt':
            return Pt(value)
        elif unit in ('in', 'inch', 'inches'):
            return Inches(value)
        elif unit == 'cm':
            return Cm(value)
        elif unit == 'mm':
            return Mm(value)

    return Pt(float(size_str))


def get_alignment(align_str: str) -> WD_ALIGN_PARAGRAPH:
    """获取段落对齐方式"""
    alignments = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return alignments.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)


# ============================================================
# 文档级操作
# ============================================================

class DocumentOperations:
    """文档级操作类"""

    @register_tool("doc_create", required_params=['path'], optional_params=['title', 'template', 'preset_template'])
    @staticmethod
    def create(path: str, title: Optional[str] = None, template: Optional[str] = None, preset_template: Optional[str] = None) -> Dict[str, Any]:
        """
        创建新的 Word 文档

        Args:
            path: 文档保存路径
            title: 可选的文档标题
            template: 可选的模板文件路径
            preset_template: 可选的预设模板名称
        """
        # If preset_template is specified, use TemplateManager
        if preset_template:
            from docuflow_mcp.extensions.templates import TemplateManager
            return TemplateManager.create_from_preset(preset_template, path, title)

        if template and os.path.exists(template):
            doc = Document(template)
        else:
            doc = Document()

        if title:
            doc.add_heading(title, level=0)

        # 确保目录存在
        dir_path = os.path.dirname(path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        doc.save(path)

        return {
            "success": True,
            "message": f"文档已创建: {path}",
            "path": path
        }

    @register_tool("doc_read", required_params=['path'], optional_params=['include_formatting'])
    @staticmethod
    def read(path: str, include_formatting: bool = False) -> Dict[str, Any]:
        """
        读取文档内容

        Args:
            path: 文档路径
            include_formatting: 是否包含格式信息
        """
        doc = Document(path)
        content = []

        for i, para in enumerate(doc.paragraphs):
            para_info = {
                "index": i,
                "type": "paragraph",
                "text": para.text,
                "style": para.style.name if para.style else "Normal"
            }

            if include_formatting:
                para_info["alignment"] = str(para.alignment) if para.alignment else "left"
                para_info["runs"] = []
                for run in para.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                    }
                    if run.font.size:
                        run_info["font_size"] = run.font.size.pt
                    if run.font.name:
                        run_info["font_name"] = run.font.name
                    para_info["runs"].append(run_info)

            content.append(para_info)

        # 读取表格
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            content.append({
                "index": i,
                "type": "table",
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data
            })

        return {
            "success": True,
            "path": path,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "content": content
        }

    @register_tool("doc_info", required_params=['path'], optional_params=[])
    @staticmethod
    def get_info(path: str) -> Dict[str, Any]:
        """获取文档基本信息"""
        doc = Document(path)

        # 核心属性
        core_props = doc.core_properties

        # 统计信息
        char_count = sum(len(p.text) for p in doc.paragraphs)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)

        # 节信息
        sections_info = []
        for i, section in enumerate(doc.sections):
            sections_info.append({
                "index": i,
                "page_width": section.page_width.inches if section.page_width else None,
                "page_height": section.page_height.inches if section.page_height else None,
                "left_margin": section.left_margin.inches if section.left_margin else None,
                "right_margin": section.right_margin.inches if section.right_margin else None,
                "top_margin": section.top_margin.inches if section.top_margin else None,
                "bottom_margin": section.bottom_margin.inches if section.bottom_margin else None,
                "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
            })

        return {
            "success": True,
            "path": path,
            "properties": {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
            },
            "statistics": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "character_count": char_count,
                "word_count": word_count,
            },
            "sections": sections_info
        }

    @register_tool("doc_set_properties", required_params=['path', 'properties'], optional_params=[])
    @staticmethod
    def set_properties(path: str, properties: Dict[str, str]) -> Dict[str, Any]:
        """设置文档属性"""
        doc = Document(path)
        core_props = doc.core_properties

        if 'title' in properties:
            core_props.title = properties['title']
        if 'author' in properties:
            core_props.author = properties['author']
        if 'subject' in properties:
            core_props.subject = properties['subject']
        if 'keywords' in properties:
            core_props.keywords = properties['keywords']
        if 'comments' in properties:
            core_props.comments = properties['comments']
        if 'category' in properties:
            core_props.category = properties['category']

        doc.save(path)

        return {
            "success": True,
            "message": "文档属性已更新",
            "properties": properties
        }

    @register_tool("doc_merge", required_params=['paths', 'output_path'], optional_params=['add_page_break'])
    @staticmethod
    def merge(paths: List[str], output_path: str, add_page_break: bool = True) -> Dict[str, Any]:
        """合并多个文档"""
        if not paths:
            return {"success": False, "error": "没有提供要合并的文档"}

        # 使用第一个文档作为基础
        merged_doc = Document(paths[0])

        for path in paths[1:]:
            if add_page_break:
                merged_doc.add_page_break()

            doc = Document(path)

            for element in doc.element.body:
                merged_doc.element.body.append(copy.deepcopy(element))

        # 确保目录存在
        dir_path = os.path.dirname(output_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        merged_doc.save(output_path)

        return {
            "success": True,
            "message": f"已合并 {len(paths)} 个文档",
            "output_path": output_path,
            "merged_files": paths
        }

    @register_tool("doc_get_styles", required_params=['path'], optional_params=[])
    @staticmethod
    def get_styles(path: str) -> Dict[str, Any]:
        """获取文档中可用的样式列表"""
        doc = Document(path)

        styles = {
            "paragraph": [],
            "character": [],
            "table": [],
            "list": []
        }

        for style in doc.styles:
            style_info = {
                "name": style.name,
                "style_id": style.style_id,
                "builtin": style.builtin,
            }

            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                styles["paragraph"].append(style_info)
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                styles["character"].append(style_info)
            elif style.type == WD_STYLE_TYPE.TABLE:
                styles["table"].append(style_info)
            elif style.type == WD_STYLE_TYPE.LIST:
                styles["list"].append(style_info)

        return {
            "success": True,
            "styles": styles
        }


# ============================================================
# 段落操作
# ============================================================

class ParagraphOperations:
    """段落操作类"""

    @register_tool("paragraph_add", required_params=['path', 'text'], optional_params=['style', 'alignment', 'font_name', 'font_size', 'font_color', 'bold', 'italic', 'underline', 'line_spacing', 'space_before', 'space_after', 'first_line_indent'])
    @staticmethod
    def add(
        path: str,
        text: str,
        style: Optional[str] = None,
        alignment: Optional[str] = None,
        font_name: Optional[str] = None,
        font_size: Optional[str] = None,
        font_color: Optional[str] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
        line_spacing: Optional[float] = None,
        space_before: Optional[str] = None,
        space_after: Optional[str] = None,
        first_line_indent: Optional[str] = None,
        insert_after: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        添加段落

        Args:
            path: 文档路径
            text: 段落文本
            style: 样式名称
            alignment: 对齐方式 (left/center/right/justify)
            font_name: 字体名称
            font_size: 字号 (如 "12pt", "14")
            font_color: 字体颜色 (如 "#FF0000", "red")
            bold: 是否加粗
            italic: 是否斜体
            underline: 是否下划线
            line_spacing: 行距倍数
            space_before: 段前距离
            space_after: 段后距离
            first_line_indent: 首行缩进
            insert_after: 在指定段落索引之后插入
        """
        doc = Document(path)

        # 创建段落
        if insert_after is not None and 0 <= insert_after < len(doc.paragraphs):
            # 在指定位置插入
            target_para = doc.paragraphs[insert_after]
            new_para = target_para.insert_paragraph_before(text)
            # 由于 insert_paragraph_before 是插入在前面，我们需要调整
            # 实际上 python-docx 没有直接的 insert_after，需要用其他方法
            para = doc.add_paragraph(text)
        else:
            para = doc.add_paragraph(text)

        # 应用样式
        if style:
            try:
                para.style = style
            except KeyError:
                pass  # 样式不存在则忽略

        # 设置对齐
        if alignment:
            para.alignment = get_alignment(alignment)

        # 设置段落格式
        pf = para.paragraph_format

        if line_spacing:
            pf.line_spacing = line_spacing

        if space_before:
            pf.space_before = parse_size(space_before)

        if space_after:
            pf.space_after = parse_size(space_after)

        if first_line_indent:
            pf.first_line_indent = parse_size(first_line_indent)

        # 设置字体格式（应用于整个段落）
        if any([font_name, font_size, font_color, bold is not None, italic is not None, underline is not None]):
            for run in para.runs:
                apply_font_format(run, font_name, font_size, font_color, bold, italic, underline)

        doc.save(path)

        return {
            "success": True,
            "message": "段落已添加",
            "paragraph_index": len(doc.paragraphs) - 1
        }

    @register_tool("paragraph_modify", required_params=['path', 'index'], optional_params=['text', 'style', 'alignment', 'font_name', 'font_size', 'font_color', 'bold', 'italic', 'underline'])
    @staticmethod
    def modify(
        path: str,
        index: int,
        text: Optional[str] = None,
        style: Optional[str] = None,
        alignment: Optional[str] = None,
        font_name: Optional[str] = None,
        font_size: Optional[str] = None,
        font_color: Optional[str] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None
    ) -> Dict[str, Any]:
        """修改指定段落"""
        doc = Document(path)

        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {index} 超出范围"}

        para = doc.paragraphs[index]

        # 修改文本
        if text is not None:
            # 清除原有内容
            para.clear()
            run = para.add_run(text)

        # 应用样式
        if style:
            try:
                para.style = style
            except KeyError:
                pass

        # 设置对齐
        if alignment:
            para.alignment = get_alignment(alignment)

        # 设置字体格式
        for run in para.runs:
            apply_font_format(run, font_name, font_size, font_color, bold, italic, underline)

        doc.save(path)

        return {
            "success": True,
            "message": f"段落 {index} 已修改"
        }

    @register_tool("paragraph_delete", required_params=['path', 'index'], optional_params=[])
    @staticmethod
    def delete(path: str, index: int) -> Dict[str, Any]:
        """删除指定段落"""
        doc = Document(path)

        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {index} 超出范围"}

        para = doc.paragraphs[index]
        p = para._element
        p.getparent().remove(p)

        doc.save(path)

        return {
            "success": True,
            "message": f"段落 {index} 已删除"
        }

    @register_tool("paragraph_get", required_params=['path', 'index'], optional_params=[])
    @staticmethod
    def get(path: str, index: int) -> Dict[str, Any]:
        """获取指定段落的详细信息"""
        doc = Document(path)

        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {index} 超出范围"}

        para = doc.paragraphs[index]

        runs_info = []
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline is not None and run.underline,
            }
            if run.font.name:
                run_info["font_name"] = run.font.name
            if run.font.size:
                run_info["font_size_pt"] = run.font.size.pt
            if run.font.color.rgb:
                run_info["font_color"] = str(run.font.color.rgb)
            runs_info.append(run_info)

        return {
            "success": True,
            "index": index,
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": runs_info
        }


# ============================================================
# 标题操作
# ============================================================

class HeadingOperations:
    """标题操作类"""

    @register_tool("heading_add", required_params=['path', 'text', 'level'], optional_params=['alignment'])
    @staticmethod
    def add(
        path: str,
        text: str,
        level: int = 1,
        alignment: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        添加标题

        Args:
            path: 文档路径
            text: 标题文本
            level: 标题级别 (0-9, 0为Title样式)
            alignment: 对齐方式
        """
        doc = Document(path)

        heading = doc.add_heading(text, level=level)

        if alignment:
            heading.alignment = get_alignment(alignment)

        doc.save(path)

        return {
            "success": True,
            "message": f"标题已添加 (级别 {level})",
            "text": text
        }

    @register_tool("heading_get_outline", required_params=['path'], optional_params=[])
    @staticmethod
    def get_outline(path: str) -> Dict[str, Any]:
        """获取文档大纲结构"""
        doc = Document(path)

        outline = []
        heading_styles = ['Title', 'Heading 1', 'Heading 2', 'Heading 3',
                         'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7',
                         'Heading 8', 'Heading 9']

        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            if style_name in heading_styles:
                level = heading_styles.index(style_name)
                outline.append({
                    "index": i,
                    "level": level,
                    "text": para.text,
                    "style": style_name
                })

        return {
            "success": True,
            "outline": outline
        }


# ============================================================
# 表格操作
# ============================================================

class TableOperations:
    """表格操作类"""

    @register_tool("table_add", required_params=['path', 'rows', 'cols'], optional_params=['data', 'style', 'header_row'])
    @staticmethod
    def add(
        path: str,
        rows: int,
        cols: int,
        data: Optional[List[List[str]]] = None,
        style: Optional[str] = None,
        header_row: bool = True
    ) -> Dict[str, Any]:
        """
        添加表格

        Args:
            path: 文档路径
            rows: 行数
            cols: 列数
            data: 表格数据（二维数组）
            style: 表格样式
            header_row: 是否将第一行作为标题行
        """
        doc = Document(path)

        table = doc.add_table(rows=rows, cols=cols)

        # 设置样式
        if style:
            try:
                table.style = style
            except KeyError:
                table.style = 'Table Grid'
        else:
            table.style = 'Table Grid'

        # 填充数据
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = str(cell_text)

        doc.save(path)

        return {
            "success": True,
            "message": f"表格已添加: {rows}行 x {cols}列",
            "table_index": len(doc.tables) - 1
        }

    @register_tool("table_get", required_params=['path', 'table_index'], optional_params=[])
    @staticmethod
    def get(path: str, table_index: int) -> Dict[str, Any]:
        """获取指定表格的内容"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)

        return {
            "success": True,
            "table_index": table_index,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "data": data
        }

    @register_tool("table_set_cell", required_params=['path', 'table_index', 'row', 'col', 'text'], optional_params=['bold', 'alignment', 'vertical_alignment', 'background_color'])
    @staticmethod
    def set_cell(
        path: str,
        table_index: int,
        row: int,
        col: int,
        text: str,
        bold: Optional[bool] = None,
        alignment: Optional[str] = None,
        vertical_alignment: Optional[str] = None,
        background_color: Optional[str] = None
    ) -> Dict[str, Any]:
        """设置单元格内容和格式"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        if row < 0 or row >= len(table.rows):
            return {"success": False, "error": f"行索引 {row} 超出范围"}

        if col < 0 or col >= len(table.columns):
            return {"success": False, "error": f"列索引 {col} 超出范围"}

        cell = table.rows[row].cells[col]
        cell.text = text

        # 设置段落格式
        if cell.paragraphs:
            para = cell.paragraphs[0]

            if alignment:
                para.alignment = get_alignment(alignment)

            if bold is not None:
                for run in para.runs:
                    run.bold = bold

        # 设置垂直对齐
        if vertical_alignment:
            v_alignments = {
                'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
                'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
            }
            cell.vertical_alignment = v_alignments.get(vertical_alignment.lower(), WD_CELL_VERTICAL_ALIGNMENT.TOP)

        # 设置背景色
        if background_color:
            color = parse_color(background_color)
            if color:
                shading_elm = OxmlElement('w:shd')
                # RGBColor 可以直接转换为字符串 "RRGGBB"
                shading_elm.set(qn('w:fill'), str(color))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        doc.save(path)

        return {
            "success": True,
            "message": f"单元格 [{row}][{col}] 已更新"
        }

    @register_tool("table_add_row", required_params=['path', 'table_index'], optional_params=['data'])
    @staticmethod
    def add_row(path: str, table_index: int, data: Optional[List[str]] = None) -> Dict[str, Any]:
        """向表格添加行"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]
        row = table.add_row()

        if data:
            for i, cell_text in enumerate(data):
                if i < len(row.cells):
                    row.cells[i].text = str(cell_text)

        doc.save(path)

        return {
            "success": True,
            "message": "行已添加",
            "new_row_index": len(table.rows) - 1
        }

    @register_tool("table_add_column", required_params=['path', 'table_index'], optional_params=['data'])
    @staticmethod
    def add_column(path: str, table_index: int, data: Optional[List[str]] = None) -> Dict[str, Any]:
        """向表格添加列"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        # python-docx 没有直接的 add_column 方法，需要手动添加
        for i, row in enumerate(table.rows):
            cell = row.add_cell()
            if data and i < len(data):
                cell.text = str(data[i])

        doc.save(path)

        return {
            "success": True,
            "message": "列已添加",
            "new_col_index": len(table.columns) - 1
        }

    @register_tool("table_delete_row", required_params=['path', 'table_index', 'row_index'], optional_params=[])
    @staticmethod
    def delete_row(path: str, table_index: int, row_index: int) -> Dict[str, Any]:
        """删除表格中的行"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        if row_index < 0 or row_index >= len(table.rows):
            return {"success": False, "error": f"行索引 {row_index} 超出范围"}

        row = table.rows[row_index]
        tr = row._tr
        tr.getparent().remove(tr)

        doc.save(path)

        return {
            "success": True,
            "message": f"行 {row_index} 已删除"
        }

    @register_tool("table_merge_cells", required_params=['path', 'table_index', 'start_row', 'start_col', 'end_row', 'end_col'], optional_params=[])
    @staticmethod
    def merge_cells(
        path: str,
        table_index: int,
        start_row: int,
        start_col: int,
        end_row: int,
        end_col: int
    ) -> Dict[str, Any]:
        """合并单元格"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        start_cell = table.rows[start_row].cells[start_col]
        end_cell = table.rows[end_row].cells[end_col]

        start_cell.merge(end_cell)

        doc.save(path)

        return {
            "success": True,
            "message": f"单元格已合并: [{start_row},{start_col}] 到 [{end_row},{end_col}]"
        }

    @register_tool("table_set_column_width", required_params=['path', 'table_index', 'col_index', 'width'], optional_params=[])
    @staticmethod
    def set_column_width(path: str, table_index: int, col_index: int, width: str) -> Dict[str, Any]:
        """设置列宽"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]

        width_val = parse_size(width)

        for row in table.rows:
            if col_index < len(row.cells):
                row.cells[col_index].width = width_val

        doc.save(path)

        return {
            "success": True,
            "message": f"列 {col_index} 宽度已设置为 {width}"
        }

    @register_tool("table_delete", required_params=['path', 'table_index'], optional_params=[])
    @staticmethod
    def delete(path: str, table_index: int) -> Dict[str, Any]:
        """删除整个表格"""
        doc = Document(path)

        if table_index < 0 or table_index >= len(doc.tables):
            return {"success": False, "error": f"表格索引 {table_index} 超出范围"}

        table = doc.tables[table_index]
        tbl = table._tbl
        tbl.getparent().remove(tbl)

        doc.save(path)

        return {
            "success": True,
            "message": f"表格 {table_index} 已删除"
        }


# ============================================================
# 图片操作
# ============================================================

class ImageOperations:
    """图片操作类"""

    @register_tool("image_add", required_params=['path', 'image_path'], optional_params=['width', 'height', 'alignment'])
    @staticmethod
    def add(
        path: str,
        image_path: str,
        width: Optional[str] = None,
        height: Optional[str] = None,
        alignment: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        插入图片

        Args:
            path: 文档路径
            image_path: 图片文件路径
            width: 宽度 (如 "3in", "5cm")
            height: 高度
            alignment: 对齐方式
        """
        doc = Document(path)

        if not os.path.exists(image_path):
            return {"success": False, "error": f"图片文件不存在: {image_path}"}

        # 添加段落并插入图片
        para = doc.add_paragraph()
        run = para.add_run()

        # 设置尺寸
        kwargs = {}
        if width:
            kwargs['width'] = parse_size(width)
        if height:
            kwargs['height'] = parse_size(height)

        run.add_picture(image_path, **kwargs)

        # 设置对齐
        if alignment:
            para.alignment = get_alignment(alignment)

        doc.save(path)

        return {
            "success": True,
            "message": f"图片已插入: {image_path}"
        }

    @register_tool("image_add_to_paragraph", required_params=['path', 'paragraph_index', 'image_path'], optional_params=['width', 'height'])
    @staticmethod
    def add_to_paragraph(
        path: str,
        paragraph_index: int,
        image_path: str,
        width: Optional[str] = None,
        height: Optional[str] = None
    ) -> Dict[str, Any]:
        """在指定段落中插入图片"""
        doc = Document(path)

        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {paragraph_index} 超出范围"}

        if not os.path.exists(image_path):
            return {"success": False, "error": f"图片文件不存在: {image_path}"}

        para = doc.paragraphs[paragraph_index]
        run = para.add_run()

        kwargs = {}
        if width:
            kwargs['width'] = parse_size(width)
        if height:
            kwargs['height'] = parse_size(height)

        run.add_picture(image_path, **kwargs)

        doc.save(path)

        return {
            "success": True,
            "message": f"图片已插入到段落 {paragraph_index}"
        }


# ============================================================
# 列表操作
# ============================================================

class ListOperations:
    """列表操作类"""

    @register_tool("list_add_bullet", required_params=['path', 'items'], optional_params=['level'])
    @staticmethod
    def add_bullet_list(path: str, items: List[str], level: int = 0) -> Dict[str, Any]:
        """添加无序列表"""
        doc = Document(path)

        for item in items:
            para = doc.add_paragraph(item, style='List Bullet')
            if level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * level)

        doc.save(path)

        return {
            "success": True,
            "message": f"无序列表已添加 ({len(items)} 项)"
        }

    @register_tool("list_add_numbered", required_params=['path', 'items'], optional_params=['level'])
    @staticmethod
    def add_numbered_list(path: str, items: List[str], level: int = 0) -> Dict[str, Any]:
        """添加有序列表"""
        doc = Document(path)

        for item in items:
            para = doc.add_paragraph(item, style='List Number')
            if level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * level)

        doc.save(path)

        return {
            "success": True,
            "message": f"有序列表已添加 ({len(items)} 项)"
        }


# ============================================================
# 页面设置
# ============================================================

class PageOperations:
    """页面操作类"""

    @register_tool("page_set_margins", required_params=['path'], optional_params=['top', 'bottom', 'left', 'right', 'section_index'])
    @staticmethod
    def set_margins(
        path: str,
        top: Optional[str] = None,
        bottom: Optional[str] = None,
        left: Optional[str] = None,
        right: Optional[str] = None,
        section_index: int = 0
    ) -> Dict[str, Any]:
        """设置页边距"""
        doc = Document(path)

        if section_index < 0 or section_index >= len(doc.sections):
            return {"success": False, "error": f"节索引 {section_index} 超出范围"}

        section = doc.sections[section_index]

        if top:
            section.top_margin = parse_size(top)
        if bottom:
            section.bottom_margin = parse_size(bottom)
        if left:
            section.left_margin = parse_size(left)
        if right:
            section.right_margin = parse_size(right)

        doc.save(path)

        return {
            "success": True,
            "message": "页边距已设置"
        }

    @register_tool("page_set_size", required_params=['path'], optional_params=['width', 'height', 'orientation', 'section_index'])
    @staticmethod
    def set_page_size(
        path: str,
        width: Optional[str] = None,
        height: Optional[str] = None,
        orientation: Optional[str] = None,
        section_index: int = 0
    ) -> Dict[str, Any]:
        """设置页面大小和方向"""
        doc = Document(path)

        if section_index < 0 or section_index >= len(doc.sections):
            return {"success": False, "error": f"节索引 {section_index} 超出范围"}

        section = doc.sections[section_index]

        if width:
            section.page_width = parse_size(width)
        if height:
            section.page_height = parse_size(height)

        if orientation:
            if orientation.lower() == 'landscape':
                section.orientation = WD_ORIENT.LANDSCAPE
                # 交换宽高
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height
            elif orientation.lower() == 'portrait':
                section.orientation = WD_ORIENT.PORTRAIT

        doc.save(path)

        return {
            "success": True,
            "message": "页面设置已更新"
        }

    @register_tool("page_add_break", required_params=['path'], optional_params=[])
    @staticmethod
    def add_page_break(path: str) -> Dict[str, Any]:
        """添加分页符"""
        doc = Document(path)
        doc.add_page_break()
        doc.save(path)

        return {
            "success": True,
            "message": "分页符已添加"
        }

    @register_tool("page_add_section_break", required_params=['path'], optional_params=['break_type'])
    @staticmethod
    def add_section_break(path: str, break_type: str = "next_page") -> Dict[str, Any]:
        """添加分节符"""
        from docx.enum.section import WD_SECTION

        doc = Document(path)

        break_types = {
            'next_page': WD_SECTION.NEW_PAGE,
            'continuous': WD_SECTION.CONTINUOUS,
            'even_page': WD_SECTION.EVEN_PAGE,
            'odd_page': WD_SECTION.ODD_PAGE,
        }

        # 添加新节
        new_section = doc.add_section(break_types.get(break_type, WD_SECTION.NEW_PAGE))

        doc.save(path)

        return {
            "success": True,
            "message": f"分节符已添加 ({break_type})"
        }


# ============================================================
# 页眉页脚
# ============================================================

class HeaderFooterOperations:
    """页眉页脚操作类"""

    @register_tool("header_set", required_params=['path', 'text'], optional_params=['alignment', 'section_index'])
    @staticmethod
    def set_header(
        path: str,
        text: str,
        alignment: Optional[str] = None,
        section_index: int = 0
    ) -> Dict[str, Any]:
        """设置页眉"""
        doc = Document(path)

        if section_index < 0 or section_index >= len(doc.sections):
            return {"success": False, "error": f"节索引 {section_index} 超出范围"}

        section = doc.sections[section_index]
        header = section.header

        # 清除现有内容
        for para in header.paragraphs:
            para.clear()

        # 添加新内容
        if header.paragraphs:
            para = header.paragraphs[0]
        else:
            para = header.add_paragraph()

        para.text = text

        if alignment:
            para.alignment = get_alignment(alignment)

        doc.save(path)

        return {
            "success": True,
            "message": "页眉已设置"
        }

    @register_tool("footer_set", required_params=['path', 'text'], optional_params=['alignment', 'section_index'])
    @staticmethod
    def set_footer(
        path: str,
        text: str,
        alignment: Optional[str] = None,
        section_index: int = 0
    ) -> Dict[str, Any]:
        """设置页脚"""
        doc = Document(path)

        if section_index < 0 or section_index >= len(doc.sections):
            return {"success": False, "error": f"节索引 {section_index} 超出范围"}

        section = doc.sections[section_index]
        footer = section.footer

        # 清除现有内容
        for para in footer.paragraphs:
            para.clear()

        # 添加新内容
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        para.text = text

        if alignment:
            para.alignment = get_alignment(alignment)

        doc.save(path)

        return {
            "success": True,
            "message": "页脚已设置"
        }

    @register_tool("page_number_add", required_params=['path'], optional_params=['position', 'alignment'])
    @staticmethod
    def add_page_number(path: str, position: str = "footer", alignment: str = "center") -> Dict[str, Any]:
        """添加页码"""
        doc = Document(path)

        for section in doc.sections:
            if position == "footer":
                target = section.footer
            else:
                target = section.header

            # 清除现有内容或获取段落
            if target.paragraphs:
                para = target.paragraphs[0]
                para.clear()
            else:
                para = target.add_paragraph()

            para.alignment = get_alignment(alignment)

            # 添加页码字段
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        doc.save(path)

        return {
            "success": True,
            "message": "页码已添加"
        }


# ============================================================
# 搜索替换
# ============================================================

class SearchOperations:
    """搜索替换操作类"""

    @register_tool("search_find", required_params=['path', 'text'], optional_params=['case_sensitive'])
    @staticmethod
    def find(path: str, text: str, case_sensitive: bool = False) -> Dict[str, Any]:
        """查找文本"""
        doc = Document(path)

        results = []
        search_text = text if case_sensitive else text.lower()

        for i, para in enumerate(doc.paragraphs):
            para_text = para.text if case_sensitive else para.text.lower()

            if search_text in para_text:
                # 查找所有出现位置
                start = 0
                while True:
                    pos = para_text.find(search_text, start)
                    if pos == -1:
                        break
                    results.append({
                        "paragraph_index": i,
                        "position": pos,
                        "context": para.text[max(0, pos-20):pos+len(text)+20]
                    })
                    start = pos + 1

        return {
            "success": True,
            "query": text,
            "count": len(results),
            "results": results
        }

    @register_tool("search_replace", required_params=['path', 'old_text', 'new_text'], optional_params=['case_sensitive', 'replace_all'])
    @staticmethod
    def replace(
        path: str,
        old_text: str,
        new_text: str,
        case_sensitive: bool = False,
        replace_all: bool = True
    ) -> Dict[str, Any]:
        """替换文本"""
        doc = Document(path)

        count = 0

        for para in doc.paragraphs:
            for run in para.runs:
                if case_sensitive:
                    if old_text in run.text:
                        if replace_all:
                            occurrences = run.text.count(old_text)
                            run.text = run.text.replace(old_text, new_text)
                            count += occurrences
                        else:
                            run.text = run.text.replace(old_text, new_text, 1)
                            count += 1
                            if not replace_all:
                                break
                else:
                    # 不区分大小写
                    pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                    if pattern.search(run.text):
                        if replace_all:
                            occurrences = len(pattern.findall(run.text))
                            run.text = pattern.sub(new_text, run.text)
                            count += occurrences
                        else:
                            run.text = pattern.sub(new_text, run.text, count=1)
                            count += 1
                            break

            if count > 0 and not replace_all:
                break

        # 同样处理表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if case_sensitive:
                                if old_text in run.text:
                                    if replace_all:
                                        occurrences = run.text.count(old_text)
                                        run.text = run.text.replace(old_text, new_text)
                                        count += occurrences
                                    else:
                                        run.text = run.text.replace(old_text, new_text, 1)
                                        count += 1
                            else:
                                pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                                if pattern.search(run.text):
                                    if replace_all:
                                        occurrences = len(pattern.findall(run.text))
                                        run.text = pattern.sub(new_text, run.text)
                                        count += occurrences
                                    else:
                                        run.text = pattern.sub(new_text, run.text, count=1)
                                        count += 1

        doc.save(path)

        return {
            "success": True,
            "message": f"已替换 {count} 处",
            "old_text": old_text,
            "new_text": new_text,
            "replacement_count": count
        }


# ============================================================
# 特殊内容
# ============================================================

class SpecialOperations:
    """特殊内容操作类"""

    @register_tool("hyperlink_add", required_params=['path', 'text', 'url'], optional_params=['paragraph_index'])
    @staticmethod
    def add_hyperlink(
        path: str,
        text: str,
        url: str,
        paragraph_index: Optional[int] = None
    ) -> Dict[str, Any]:
        """添加超链接"""
        doc = Document(path)

        if paragraph_index is not None:
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"段落索引 {paragraph_index} 超出范围"}
            para = doc.paragraphs[paragraph_index]
        else:
            para = doc.add_paragraph()

        # 创建超链接
        part = doc.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # 设置超链接样式（蓝色下划线）
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)

        hyperlink.append(new_run)
        para._p.append(hyperlink)

        doc.save(path)

        return {
            "success": True,
            "message": f"超链接已添加: {text} -> {url}"
        }

    @register_tool("toc_add", required_params=['path'], optional_params=[])
    @staticmethod
    def add_table_of_contents(path: str) -> Dict[str, Any]:
        """添加目录（占位符，需要在 Word 中更新）"""
        doc = Document(path)

        para = doc.add_paragraph()
        run = para.add_run()

        # 添加 TOC 字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        doc.save(path)

        return {
            "success": True,
            "message": "目录已添加（请在 Word 中按 F9 更新目录）"
        }

    @register_tool("line_break_add", required_params=['path', 'paragraph_index'], optional_params=[])
    @staticmethod
    def add_line_break(path: str, paragraph_index: int) -> Dict[str, Any]:
        """在段落中添加换行符"""
        doc = Document(path)

        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {paragraph_index} 超出范围"}

        para = doc.paragraphs[paragraph_index]
        run = para.add_run()
        run.add_break(WD_BREAK.LINE)

        doc.save(path)

        return {
            "success": True,
            "message": f"换行符已添加到段落 {paragraph_index}"
        }

    @register_tool("horizontal_line_add", required_params=['path'], optional_params=[])
    @staticmethod
    def add_horizontal_line(path: str) -> Dict[str, Any]:
        """添加水平线"""
        doc = Document(path)

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

        # 添加下边框作为水平线
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.save(path)

        return {
            "success": True,
            "message": "水平线已添加"
        }


# ============================================================
# 导出功能
# ============================================================

class ExportOperations:
    """导出操作类"""

    @register_tool("export_to_text", required_params=['path'], optional_params=['output_path'])
    @staticmethod
    def to_text(path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """导出为纯文本"""
        doc = Document(path)

        text_content = []
        for para in doc.paragraphs:
            text_content.append(para.text)

        full_text = '\n'.join(text_content)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            return {
                "success": True,
                "message": f"已导出到: {output_path}",
                "output_path": output_path
            }
        else:
            return {
                "success": True,
                "text": full_text
            }

    @register_tool("export_to_markdown", required_params=['path'], optional_params=['output_path'])
    @staticmethod
    def to_markdown(path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """导出为 Markdown 格式"""
        doc = Document(path)

        md_content = []
        heading_styles = {
            'Title': '#',
            'Heading 1': '#',
            'Heading 2': '##',
            'Heading 3': '###',
            'Heading 4': '####',
            'Heading 5': '#####',
            'Heading 6': '######',
        }

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            text = para.text

            if not text.strip():
                md_content.append('')
                continue

            # 处理标题
            if style_name in heading_styles:
                prefix = heading_styles[style_name]
                md_content.append(f"{prefix} {text}")
            # 处理列表
            elif 'List Bullet' in style_name:
                md_content.append(f"- {text}")
            elif 'List Number' in style_name:
                md_content.append(f"1. {text}")
            else:
                # 处理格式化文本
                formatted_text = []
                for run in para.runs:
                    run_text = run.text
                    if run.bold and run.italic:
                        run_text = f"***{run_text}***"
                    elif run.bold:
                        run_text = f"**{run_text}**"
                    elif run.italic:
                        run_text = f"*{run_text}*"
                    formatted_text.append(run_text)

                md_content.append(''.join(formatted_text) or text)

        # 处理表格
        for table in doc.tables:
            md_content.append('')
            for i, row in enumerate(table.rows):
                cells = [cell.text.replace('|', '\\|') for cell in row.cells]
                md_content.append('| ' + ' | '.join(cells) + ' |')
                if i == 0:
                    # 添加表头分隔符
                    md_content.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
            md_content.append('')

        full_md = '\n'.join(md_content)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_md)
            return {
                "success": True,
                "message": f"已导出到: {output_path}",
                "output_path": output_path
            }
        else:
            return {
                "success": True,
                "markdown": full_md
            }


class CommentOperations:
    """Word文档批注操作"""

    @register_tool("comment_add",
                   required_params=['path', 'paragraph_index', 'text'],
                   optional_params=['author', 'date'])
    @staticmethod
    def add_comment(path: str,
                    paragraph_index: int,
                    text: str,
                    author: Optional[str] = None,
                    date: Optional[str] = None) -> Dict[str, Any]:
        """
        为指定段落添加批注

        Args:
            path: 文档路径
            paragraph_index: 段落索引（从0开始）
            text: 批注文本
            author: 作者名称
            date: 日期（ISO格式，默认当前时间）

        Returns:
            {success, comment_id, message}
        """
        try:
            from datetime import datetime
            from lxml import etree

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            doc = Document(path)

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"success": False, "error": f"段落索引超出范围 (0-{len(doc.paragraphs)-1})"}

            # 设置默认值
            if not author:
                author = "DocuFlow"
            if not date:
                date = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

            # Word XML 命名空间
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            w_tag = '{%s}' % w_ns

            # 查找或创建 comments part
            comments_part = None
            comments_rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'

            # 检查是否已有 comments part
            for rel in doc.part.rels.values():
                if rel.reltype == comments_rel_type:
                    comments_part = rel.target_part
                    break

            if comments_part is not None:
                # 已有 comments part — 它是 XmlPart 子类，直接操作 _element
                comments_elem = comments_part._element
            else:
                # 创建新的 comments XML
                nsmap_comments = {
                    'w': w_ns,
                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                }
                comments_elem = etree.Element(w_tag + 'comments', nsmap=nsmap_comments)

            # 确定下一个 comment ID
            existing_ids = []
            for comment in comments_elem.findall(w_tag + 'comment'):
                cid = comment.get(w_tag + 'id')
                if cid is not None:
                    existing_ids.append(int(cid))
            next_id = max(existing_ids) + 1 if existing_ids else 0

            # 创建 comment 元素
            comment_elem = etree.SubElement(comments_elem, w_tag + 'comment')
            comment_elem.set(w_tag + 'id', str(next_id))
            comment_elem.set(w_tag + 'author', author)
            comment_elem.set(w_tag + 'date', date)

            # 批注内容段落
            comment_p = etree.SubElement(comment_elem, w_tag + 'p')
            comment_r = etree.SubElement(comment_p, w_tag + 'r')
            comment_t = etree.SubElement(comment_r, w_tag + 't')
            comment_t.text = text

            if comments_part is None:
                # 创建新的 Part 并添加关系
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI

                comments_xml = etree.tostring(comments_elem, xml_declaration=True,
                                               encoding='UTF-8', standalone=True)
                comments_part = Part(
                    partname=PackURI('/word/comments.xml'),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
                    blob=comments_xml,
                    package=doc.part.package
                )
                doc.part.relate_to(comments_part, comments_rel_type)

            # 在段落XML中插入批注标记
            para = doc.paragraphs[paragraph_index]
            para_elem = para._element

            # 添加 commentRangeStart
            range_start = OxmlElement('w:commentRangeStart')
            range_start.set(qn('w:id'), str(next_id))

            # 添加 commentRangeEnd
            range_end = OxmlElement('w:commentRangeEnd')
            range_end.set(qn('w:id'), str(next_id))

            # 添加 commentReference（在一个 run 中）
            ref_run = OxmlElement('w:r')
            ref_rpr = OxmlElement('w:rPr')
            ref_rstyle = OxmlElement('w:rStyle')
            ref_rstyle.set(qn('w:val'), 'CommentReference')
            ref_rpr.append(ref_rstyle)
            ref_run.append(ref_rpr)

            comment_ref = OxmlElement('w:commentReference')
            comment_ref.set(qn('w:id'), str(next_id))
            ref_run.append(comment_ref)

            # 插入到段落中
            para_elem.insert(0, range_start)
            para_elem.append(range_end)
            para_elem.append(ref_run)

            doc.save(path)

            return {
                "success": True,
                "path": path,
                "paragraph_index": paragraph_index,
                "comment_id": next_id,
                "author": author,
                "text": text,
                "message": f"已在第 {paragraph_index} 段添加批注 (ID: {next_id})"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("comment_list",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def list_comments(path: str) -> Dict[str, Any]:
        """
        列出文档中的所有批注

        Args:
            path: 文档路径

        Returns:
            {success, comments, comment_count, message}
        """
        try:
            from lxml import etree

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            doc = Document(path)

            # 查找 comments part
            comments_rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
            comments_part = None

            for rel in doc.part.rels.values():
                if rel.reltype == comments_rel_type:
                    comments_part = rel.target_part
                    break

            if comments_part is None:
                return {
                    "success": True,
                    "comments": [],
                    "comment_count": 0,
                    "message": "文档中没有批注"
                }

            # 解析 comments XML
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            w_tag = '{%s}' % w_ns

            comments_elem = etree.fromstring(comments_part.blob)
            comments_list = []

            for comment in comments_elem.findall(w_tag + 'comment'):
                comment_id = comment.get(w_tag + 'id')
                comment_author = comment.get(w_tag + 'author', '')
                comment_date = comment.get(w_tag + 'date', '')

                # 提取批注文本
                texts = []
                for p in comment.findall(w_tag + 'p'):
                    for r in p.findall(w_tag + 'r'):
                        for t in r.findall(w_tag + 't'):
                            if t.text:
                                texts.append(t.text)

                comments_list.append({
                    "id": int(comment_id) if comment_id else None,
                    "author": comment_author,
                    "date": comment_date,
                    "text": ''.join(texts)
                })

            return {
                "success": True,
                "comments": comments_list,
                "comment_count": len(comments_list),
                "message": f"共 {len(comments_list)} 条批注"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}
