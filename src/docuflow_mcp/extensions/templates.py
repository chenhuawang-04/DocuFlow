"""
DocuFlow MCP - Template Management

Provides template creation and management functionality
"""

import os
import json
from typing import Dict, Any, Optional
from pathlib import Path

from docx import Document
from docuflow_mcp.document import parse_size, get_alignment
from docuflow_mcp.utils.formatters import apply_style_to_paragraph
from docuflow_mcp.core.registry import register_tool


class TemplateManager:
    """Template management operations"""

    @staticmethod
    def _get_presets_path():
        """Get the path to presets.json"""
        current_dir = Path(__file__).parent.parent.parent.parent
        return current_dir / "templates" / "presets.json"

    @staticmethod
    def load_presets() -> Dict[str, Any]:
        """Load preset templates from JSON"""
        presets_path = TemplateManager._get_presets_path()
        if not presets_path.exists():
            return {}

        with open(presets_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    @register_tool("template_list_presets", required_params=[], optional_params=[])
    @staticmethod
    def list_presets() -> Dict[str, Any]:
        """List all available preset templates"""
        presets = TemplateManager.load_presets()
        result = []
        for key, preset in presets.items():
            result.append({
                "id": key,
                "name": preset.get("name", key),
                "description": preset.get("description", "")
            })

        return {
            "success": True,
            "presets": result,
            "count": len(result)
        }

    @register_tool("template_create_from_preset", required_params=["preset_name", "output_path"], optional_params=["title"])
    @staticmethod
    def create_from_preset(preset_name: str, output_path: str, title: Optional[str] = None) -> Dict[str, Any]:
        """Create a document from a preset template"""
        presets = TemplateManager.load_presets()

        if preset_name not in presets:
            return {
                "success": False,
                "error": f"预设模板不存在: {preset_name}",
                "available_presets": list(presets.keys())
            }

        preset = presets[preset_name]
        doc = Document()

        # Apply page setup
        if 'page_setup' in preset:
            TemplateManager._apply_page_setup(doc, preset['page_setup'])

        # Apply styles
        if 'styles' in preset:
            TemplateManager._apply_styles(doc, preset['styles'])

        # Add title if provided
        if title:
            doc.add_heading(title, level=0)

        # Apply header/footer
        if 'header' in preset:
            TemplateManager._apply_header(doc, preset['header'])
        if 'footer' in preset:
            TemplateManager._apply_footer(doc, preset['footer'])

        # Ensure directory exists
        dir_path = os.path.dirname(output_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        doc.save(output_path)

        return {
            "success": True,
            "message": f"已从预设模板 '{preset.get('name', preset_name)}' 创建文档",
            "path": output_path,
            "preset": preset_name
        }

    @staticmethod
    def _apply_page_setup(doc: Document, page_setup: Dict[str, Any]):
        """Apply page setup to document"""
        section = doc.sections[0]

        # Apply margins
        if 'margins' in page_setup:
            margins = page_setup['margins']
            if 'top' in margins:
                section.top_margin = parse_size(margins['top'])
            if 'bottom' in margins:
                section.bottom_margin = parse_size(margins['bottom'])
            if 'left' in margins:
                section.left_margin = parse_size(margins['left'])
            if 'right' in margins:
                section.right_margin = parse_size(margins['right'])

        # Apply page size
        if 'size' in page_setup:
            size = page_setup['size']
            if 'width' in size:
                section.page_width = parse_size(size['width'])
            if 'height' in size:
                section.page_height = parse_size(size['height'])
            if 'orientation' in size:
                from docx.enum.section import WD_ORIENT
                if size['orientation'] == 'landscape':
                    section.orientation = WD_ORIENT.LANDSCAPE
                    new_width = section.page_height
                    new_height = section.page_width
                    section.page_width = new_width
                    section.page_height = new_height

    @staticmethod
    def _apply_styles(doc: Document, styles: Dict[str, Any]):
        """Apply styles to document"""
        for style_name, style_def in styles.items():
            try:
                style = doc.styles[style_name]

                # Apply font formatting
                if 'font' in style_def:
                    font_def = style_def['font']
                    if hasattr(style, 'font'):
                        if 'name' in font_def:
                            style.font.name = font_def['name']
                        if 'size' in font_def:
                            style.font.size = parse_size(font_def['size'])
                        if 'bold' in font_def:
                            style.font.bold = font_def['bold']
                        if 'italic' in font_def:
                            style.font.italic = font_def['italic']
                        if 'color' in font_def:
                            from docuflow_mcp.document import parse_color
                            color = parse_color(font_def['color'])
                            if color:
                                style.font.color.rgb = color

                # Apply paragraph formatting
                if 'paragraph' in style_def and hasattr(style, 'paragraph_format'):
                    para_def = style_def['paragraph']
                    pf = style.paragraph_format

                    if 'alignment' in para_def:
                        pf.alignment = get_alignment(para_def['alignment'])
                    if 'line_spacing' in para_def:
                        pf.line_spacing = para_def['line_spacing']
                    if 'space_before' in para_def:
                        pf.space_before = parse_size(para_def['space_before'])
                    if 'space_after' in para_def:
                        pf.space_after = parse_size(para_def['space_after'])
                    if 'first_line_indent' in para_def:
                        pf.first_line_indent = parse_size(para_def['first_line_indent'])
            except KeyError:
                pass  # Style doesn't exist

    @staticmethod
    def _apply_header(doc: Document, header_def: Dict[str, Any]):
        """Apply header to document"""
        section = doc.sections[0]
        header = section.header

        if 'text' in header_def and header_def['text']:
            if header.paragraphs:
                para = header.paragraphs[0]
            else:
                para = header.add_paragraph()

            para.text = header_def['text']

            if 'alignment' in header_def:
                para.alignment = get_alignment(header_def['alignment'])

    @staticmethod
    def _apply_footer(doc: Document, footer_def: Dict[str, Any]):
        """Apply footer to document"""
        section = doc.sections[0]
        footer = section.footer

        if footer_def.get('page_number'):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            if footer.paragraphs:
                para = footer.paragraphs[0]
                para.clear()
            else:
                para = footer.add_paragraph()

            para.alignment = get_alignment(footer_def.get('alignment', 'center'))

            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
        elif 'text' in footer_def and footer_def['text']:
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.text = footer_def['text']

            if 'alignment' in footer_def:
                para.alignment = get_alignment(footer_def['alignment'])

    @register_tool("template_apply_styles", required_params=["path", "preset_name"], optional_params=[])
    @staticmethod
    def apply_template_styles(path: str, preset_name: str) -> Dict[str, Any]:
        """Apply template styles to an existing document"""
        presets = TemplateManager.load_presets()

        if preset_name not in presets:
            return {
                "success": False,
                "error": f"预设模板不存在: {preset_name}"
            }

        preset = presets[preset_name]
        doc = Document(path)

        # Apply styles
        if 'styles' in preset:
            TemplateManager._apply_styles(doc, preset['styles'])

        doc.save(path)

        return {
            "success": True,
            "message": f"已应用模板样式: {preset.get('name', preset_name)}"
        }

    @register_tool("template_save", required_params=["path", "template_name"], optional_params=["description"])
    @staticmethod
    def save_as_template(path: str, template_name: str, description: Optional[str] = None) -> Dict[str, Any]:
        """Save current document as a custom template"""
        doc = Document(path)

        # Extract document properties
        template_def = {
            "name": template_name,
            "description": description or f"自定义模板: {template_name}",
            "created_from": path,
            "page_setup": {},
            "styles": {}
        }

        # Extract page setup from first section
        if doc.sections:
            section = doc.sections[0]
            template_def["page_setup"] = {
                "margins": {
                    "top": f"{section.top_margin.cm:.2f}cm" if section.top_margin else "2.54cm",
                    "bottom": f"{section.bottom_margin.cm:.2f}cm" if section.bottom_margin else "2.54cm",
                    "left": f"{section.left_margin.cm:.2f}cm" if section.left_margin else "2.54cm",
                    "right": f"{section.right_margin.cm:.2f}cm" if section.right_margin else "2.54cm"
                },
                "size": {
                    "width": f"{section.page_width.cm:.2f}cm" if section.page_width else "21cm",
                    "height": f"{section.page_height.cm:.2f}cm" if section.page_height else "29.7cm",
                }
            }

            # Extract header/footer
            if section.header.paragraphs and section.header.paragraphs[0].text:
                template_def["header"] = {
                    "text": section.header.paragraphs[0].text,
                    "alignment": str(section.header.paragraphs[0].alignment) if section.header.paragraphs[0].alignment else "left"
                }

            if section.footer.paragraphs and section.footer.paragraphs[0].text:
                template_def["footer"] = {
                    "text": section.footer.paragraphs[0].text,
                    "alignment": str(section.footer.paragraphs[0].alignment) if section.footer.paragraphs[0].alignment else "center"
                }

        # Extract styles
        for style in doc.styles:
            if not style.builtin or style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                style_def = {}

                # Extract font properties
                if hasattr(style, 'font') and style.font:
                    font_def = {}
                    if style.font.name:
                        font_def['name'] = style.font.name
                    if style.font.size:
                        font_def['size'] = f"{style.font.size.pt}pt"
                    if style.font.bold is not None:
                        font_def['bold'] = style.font.bold
                    if style.font.italic is not None:
                        font_def['italic'] = style.font.italic
                    if style.font.color and style.font.color.rgb:
                        font_def['color'] = f"#{style.font.color.rgb}"

                    if font_def:
                        style_def['font'] = font_def

                # Extract paragraph properties
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    para_def = {}
                    pf = style.paragraph_format

                    if pf.alignment:
                        para_def['alignment'] = str(pf.alignment)
                    if pf.line_spacing:
                        para_def['line_spacing'] = pf.line_spacing
                    if pf.space_before:
                        para_def['space_before'] = f"{pf.space_before.pt}pt"
                    if pf.space_after:
                        para_def['space_after'] = f"{pf.space_after.pt}pt"
                    if pf.first_line_indent:
                        para_def['first_line_indent'] = f"{pf.first_line_indent.pt}pt"

                    if para_def:
                        style_def['paragraph'] = para_def

                if style_def:
                    template_def['styles'][style.name] = style_def

        # Save template to custom templates directory
        custom_dir = Path(__file__).parent.parent.parent.parent / "templates" / "custom"
        custom_dir.mkdir(parents=True, exist_ok=True)

        template_file = custom_dir / f"{template_name}.json"
        with open(template_file, 'w', encoding='utf-8') as f:
            json.dump(template_def, f, indent=2, ensure_ascii=False)

        return {
            "success": True,
            "message": f"模板已保存: {template_name}",
            "template_file": str(template_file),
            "template_def": template_def
        }

    @register_tool("template_load", required_params=["template_name", "output_path"], optional_params=["title"])
    @staticmethod
    def load_custom_template(template_name: str, output_path: str, title: Optional[str] = None) -> Dict[str, Any]:
        """Load and create document from custom template"""
        # Try custom templates first
        custom_dir = Path(__file__).parent.parent.parent.parent / "templates" / "custom"
        template_file = custom_dir / f"{template_name}.json"

        if not template_file.exists():
            # Try presets
            return TemplateManager.create_from_preset(template_name, output_path, title)

        # Load custom template
        with open(template_file, 'r', encoding='utf-8') as f:
            template_def = json.load(f)

        doc = Document()

        # Apply page setup
        if 'page_setup' in template_def:
            TemplateManager._apply_page_setup(doc, template_def['page_setup'])

        # Apply styles
        if 'styles' in template_def:
            TemplateManager._apply_styles(doc, template_def['styles'])

        # Add title if provided
        if title:
            doc.add_heading(title, level=0)

        # Apply header/footer
        if 'header' in template_def:
            TemplateManager._apply_header(doc, template_def['header'])
        if 'footer' in template_def:
            TemplateManager._apply_footer(doc, template_def['footer'])

        # Ensure directory exists
        dir_path = os.path.dirname(output_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        doc.save(output_path)

        return {
            "success": True,
            "message": f"已从自定义模板 '{template_name}' 创建文档",
            "path": output_path,
            "template": template_name
        }

    @register_tool("template_delete", required_params=["template_name"], optional_params=[])
    @staticmethod
    def delete_custom_template(template_name: str) -> Dict[str, Any]:
        """Delete a custom template"""
        custom_dir = Path(__file__).parent.parent.parent.parent / "templates" / "custom"
        template_file = custom_dir / f"{template_name}.json"

        if not template_file.exists():
            return {
                "success": False,
                "error": f"自定义模板不存在: {template_name}"
            }

        template_file.unlink()

        return {
            "success": True,
            "message": f"已删除自定义模板: {template_name}"
        }

    @register_tool("template_list_custom", required_params=[], optional_params=[])
    @staticmethod
    def list_custom_templates() -> Dict[str, Any]:
        """List all custom templates"""
        custom_dir = Path(__file__).parent.parent.parent.parent / "templates" / "custom"

        if not custom_dir.exists():
            return {
                "success": True,
                "templates": [],
                "count": 0
            }

        templates = []
        for template_file in custom_dir.glob("*.json"):
            try:
                with open(template_file, 'r', encoding='utf-8') as f:
                    template_def = json.load(f)

                templates.append({
                    "id": template_file.stem,
                    "name": template_def.get("name", template_file.stem),
                    "description": template_def.get("description", ""),
                    "created_from": template_def.get("created_from", "")
                })
            except (json.JSONDecodeError, OSError, KeyError):
                pass

        return {
            "success": True,
            "templates": templates,
            "count": len(templates)
        }

