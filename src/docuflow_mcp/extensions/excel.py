"""
DocuFlow Excel - Excel表格处理模块

支持：
- 工作簿/工作表操作
- 单元格读写与格式化
- 行列操作
- 公式支持
- 与Word集成
"""
import os
from pathlib import Path
from typing import Optional, List, Dict, Any, Union

from ..core.registry import register_tool
from ..utils.deps import check_import


def _col_letter_to_index(col: Union[int, str]) -> int:
    """将列字母转换为索引（1-based）"""
    if isinstance(col, int):
        return col
    col = col.upper()
    result = 0
    for char in col:
        result = result * 26 + (ord(char) - ord('A') + 1)
    return result


def _col_index_to_letter(index: int) -> str:
    """将列索引转换为字母"""
    result = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        result = chr(65 + remainder) + result
    return result


class ExcelOperations:
    """Excel表格操作"""

    # ========== 工作簿操作 ==========

    @register_tool("excel_create",
                   required_params=['path'],
                   optional_params=['sheets', 'title'])
    @staticmethod
    def create(path: str,
               sheets: Optional[List[str]] = None,
               title: Optional[str] = None) -> Dict[str, Any]:
        """
        创建新Excel文件

        Args:
            path: 文件保存路径（.xlsx）
            sheets: 工作表名称列表，默认["Sheet1"]
            title: 文档标题（元数据）

        Returns:
            {success, path, sheets, message}
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import Workbook

            if not path.endswith('.xlsx'):
                return {"success": False, "error": "文件路径必须以.xlsx结尾"}

            # 创建目录
            dir_path = os.path.dirname(path)
            if dir_path and not os.path.exists(dir_path):
                os.makedirs(dir_path)

            wb = Workbook()

            # 设置工作表
            if sheets:
                # 重命名默认的Sheet
                ws = wb.active
                ws.title = sheets[0]
                # 添加其他工作表
                for sheet_name in sheets[1:]:
                    wb.create_sheet(title=sheet_name)
            else:
                sheets = ["Sheet1"]

            # 设置文档属性
            if title:
                wb.properties.title = title

            wb.save(path)

            return {
                "success": True,
                "path": path,
                "sheets": sheets,
                "message": f"Excel文件创建成功: {path}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("excel_read",
                   required_params=['path'],
                   optional_params=['sheet', 'range', 'include_formatting'])
    @staticmethod
    def read(path: str,
             sheet: Optional[str] = None,
             range: Optional[str] = None,
             include_formatting: bool = False) -> Dict[str, Any]:
        """
        读取Excel内容

        Args:
            path: Excel文件路径
            sheet: 工作表名称（默认活动表）
            range: 读取范围如"A1:D10"（默认全部有数据区域）
            include_formatting: 是否包含格式信息

        Returns:
            {success, data, rows, cols, message}
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path, data_only=True)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 确定读取范围
            if range:
                # 解析范围如"A1:D10"
                try:
                    cells = ws[range]
                except Exception:
                    return {"success": False, "error": f"无效的范围: {range}"}
            else:
                # 读取全部有数据的区域
                cells = ws.iter_rows(min_row=1, max_row=ws.max_row,
                                     min_col=1, max_col=ws.max_column)

            # 读取数据
            data = []
            for row in cells:
                if isinstance(row, tuple):
                    row_data = []
                    for cell in row:
                        if include_formatting:
                            cell_info = {
                                "value": cell.value,
                                "coordinate": cell.coordinate
                            }
                            if cell.font:
                                cell_info["font"] = {
                                    "name": cell.font.name,
                                    "size": cell.font.size,
                                    "bold": cell.font.bold,
                                    "italic": cell.font.italic
                                }
                            row_data.append(cell_info)
                        else:
                            row_data.append(cell.value)
                    data.append(row_data)
                else:
                    # 单个单元格
                    if include_formatting:
                        data.append({
                            "value": row.value,
                            "coordinate": row.coordinate
                        })
                    else:
                        data.append(row.value)

            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "data": data,
                "rows": len(data),
                "cols": len(data[0]) if data and isinstance(data[0], list) else 1,
                "message": f"读取成功: {len(data)} 行数据"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("excel_info",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def get_info(path: str) -> Dict[str, Any]:
        """
        获取工作簿信息

        Returns:
            {success, sheets, active_sheet, properties, statistics}
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path, read_only=True)

            # 收集工作表信息
            sheets_info = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheets_info.append({
                    "name": sheet_name,
                    "max_row": ws.max_row,
                    "max_column": ws.max_column
                })

            # 文档属性
            props = wb.properties
            properties = {
                "title": props.title,
                "creator": props.creator,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "subject": props.subject,
                "keywords": props.keywords
            }

            wb.close()

            return {
                "success": True,
                "path": path,
                "sheets": sheets_info,
                "sheet_count": len(sheets_info),
                "active_sheet": wb.active.title if wb.active else None,
                "properties": properties,
                "message": f"工作簿包含 {len(sheets_info)} 个工作表"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("excel_save_as",
                   required_params=['path', 'output_path'],
                   optional_params=['format'])
    @staticmethod
    def save_as(path: str,
                output_path: str,
                format: Optional[str] = None) -> Dict[str, Any]:
        """
        另存为

        Args:
            path: 源文件路径
            output_path: 目标路径
            format: 目标格式（xlsx/csv）
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            # 检测输出格式
            if format:
                out_format = format.lower()
            else:
                out_format = Path(output_path).suffix.lower().lstrip('.')

            # 创建输出目录
            dir_path = os.path.dirname(output_path)
            if dir_path and not os.path.exists(dir_path):
                os.makedirs(dir_path)

            wb = load_workbook(path)

            if out_format == 'xlsx':
                wb.save(output_path)
            elif out_format == 'csv':
                import csv
                ws = wb.active
                with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    for row in ws.iter_rows(values_only=True):
                        writer.writerow(row)
            else:
                return {"success": False, "error": f"不支持的格式: {out_format}"}

            wb.close()

            return {
                "success": True,
                "input_path": path,
                "output_path": output_path,
                "format": out_format,
                "message": f"文件已保存为: {output_path}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 工作表操作 ==========

    @register_tool("sheet_list",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def list_sheets(path: str) -> Dict[str, Any]:
        """列出所有工作表"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path, read_only=True)
            sheets = wb.sheetnames
            active = wb.active.title if wb.active else None
            wb.close()

            return {
                "success": True,
                "sheets": sheets,
                "count": len(sheets),
                "active_sheet": active,
                "message": f"共 {len(sheets)} 个工作表"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("sheet_add",
                   required_params=['path', 'name'],
                   optional_params=['position'])
    @staticmethod
    def add_sheet(path: str,
                  name: str,
                  position: Optional[int] = None) -> Dict[str, Any]:
        """添加新工作表"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            if name in wb.sheetnames:
                return {"success": False, "error": f"工作表已存在: {name}"}

            ws = wb.create_sheet(title=name, index=position)
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": name,
                "position": position if position is not None else len(wb.sheetnames) - 1,
                "message": f"工作表 '{name}' 已添加"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("sheet_delete",
                   required_params=['path', 'name'],
                   optional_params=[])
    @staticmethod
    def delete_sheet(path: str, name: str) -> Dict[str, Any]:
        """删除工作表"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            if name not in wb.sheetnames:
                return {"success": False, "error": f"工作表不存在: {name}"}

            if len(wb.sheetnames) == 1:
                return {"success": False, "error": "不能删除唯一的工作表"}

            del wb[name]
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "deleted_sheet": name,
                "message": f"工作表 '{name}' 已删除"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("sheet_rename",
                   required_params=['path', 'old_name', 'new_name'],
                   optional_params=[])
    @staticmethod
    def rename_sheet(path: str,
                     old_name: str,
                     new_name: str) -> Dict[str, Any]:
        """重命名工作表"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            if old_name not in wb.sheetnames:
                return {"success": False, "error": f"工作表不存在: {old_name}"}

            if new_name in wb.sheetnames:
                return {"success": False, "error": f"工作表名称已存在: {new_name}"}

            ws = wb[old_name]
            ws.title = new_name
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "old_name": old_name,
                "new_name": new_name,
                "message": f"工作表已重命名: '{old_name}' -> '{new_name}'"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("sheet_copy",
                   required_params=['path', 'source_name', 'target_name'],
                   optional_params=[])
    @staticmethod
    def copy_sheet(path: str,
                   source_name: str,
                   target_name: str) -> Dict[str, Any]:
        """复制工作表"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            if source_name not in wb.sheetnames:
                return {"success": False, "error": f"源工作表不存在: {source_name}"}

            if target_name in wb.sheetnames:
                return {"success": False, "error": f"目标工作表已存在: {target_name}"}

            source = wb[source_name]
            target = wb.copy_worksheet(source)
            target.title = target_name
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "source": source_name,
                "target": target_name,
                "message": f"工作表已复制: '{source_name}' -> '{target_name}'"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 单元格操作 ==========

    @register_tool("cell_read",
                   required_params=['path'],
                   optional_params=['sheet', 'cell', 'range'])
    @staticmethod
    def read_cell(path: str,
                  sheet: Optional[str] = None,
                  cell: Optional[str] = None,
                  range: Optional[str] = None) -> Dict[str, Any]:
        """
        读取单元格

        Args:
            path: Excel文件路径
            sheet: 工作表名称
            cell: 单个单元格如"A1"
            range: 区域如"A1:D10"
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            if not cell and not range:
                return {"success": False, "error": "必须指定cell或range参数"}

            wb = load_workbook(path, data_only=True)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            if cell:
                # 读取单个单元格
                value = ws[cell].value
                wb.close()
                return {
                    "success": True,
                    "sheet": sheet,
                    "cell": cell,
                    "value": value,
                    "message": f"单元格 {cell} = {value}"
                }
            else:
                # 读取范围
                data = []
                for row in ws[range]:
                    row_data = [c.value for c in row]
                    data.append(row_data)
                wb.close()
                return {
                    "success": True,
                    "sheet": sheet,
                    "range": range,
                    "data": data,
                    "rows": len(data),
                    "message": f"读取 {len(data)} 行数据"
                }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("cell_write",
                   required_params=['path'],
                   optional_params=['sheet', 'cell', 'value', 'range', 'data'])
    @staticmethod
    def write_cell(path: str,
                   sheet: Optional[str] = None,
                   cell: Optional[str] = None,
                   value: Optional[Any] = None,
                   range: Optional[str] = None,
                   data: Optional[List[List]] = None) -> Dict[str, Any]:
        """
        写入单元格

        Args:
            path: Excel文件路径
            sheet: 工作表名称
            cell: 单个单元格如"A1"
            value: 单个值
            range: 区域起点如"A1"（与data配合使用）
            data: 二维数组数据
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            if cell and value is not None:
                # 写入单个单元格
                ws[cell] = value
                wb.save(path)
                wb.close()
                return {
                    "success": True,
                    "sheet": sheet,
                    "cell": cell,
                    "value": value,
                    "message": f"已写入 {cell} = {value}"
                }

            elif range and data:
                # 写入数据区域
                col_letter, row_num = coordinate_from_string(range)
                start_col = column_index_from_string(col_letter)
                start_row = row_num

                for i, row_data in enumerate(data):
                    for j, val in enumerate(row_data):
                        ws.cell(row=start_row + i, column=start_col + j, value=val)

                wb.save(path)
                wb.close()
                return {
                    "success": True,
                    "sheet": sheet,
                    "start": range,
                    "rows_written": len(data),
                    "cols_written": len(data[0]) if data else 0,
                    "message": f"已写入 {len(data)} 行数据"
                }

            else:
                return {"success": False, "error": "需要指定 (cell + value) 或 (range + data)"}

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("cell_format",
                   required_params=['path', 'range'],
                   optional_params=['sheet', 'font_name', 'font_size', 'bold',
                                   'italic', 'font_color', 'bg_color',
                                   'border', 'alignment', 'number_format'])
    @staticmethod
    def format_cell(path: str,
                    range: str,
                    sheet: Optional[str] = None,
                    font_name: Optional[str] = None,
                    font_size: Optional[int] = None,
                    bold: Optional[bool] = None,
                    italic: Optional[bool] = None,
                    font_color: Optional[str] = None,
                    bg_color: Optional[str] = None,
                    border: Optional[str] = None,
                    alignment: Optional[str] = None,
                    number_format: Optional[str] = None) -> Dict[str, Any]:
        """
        设置单元格格式

        Args:
            range: 格式化范围如"A1:D10"或"A1"
            font_name: 字体名称
            font_size: 字号
            bold: 加粗
            italic: 斜体
            font_color: 字体颜色（十六进制如"FF0000"）
            bg_color: 背景颜色
            border: 边框样式（thin/medium/thick）
            alignment: 对齐（left/center/right）
            number_format: 数字格式（如"0.00%", "#,##0"）
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 准备样式
            font_kwargs = {}
            if font_name:
                font_kwargs['name'] = font_name
            if font_size:
                font_kwargs['size'] = font_size
            if bold is not None:
                font_kwargs['bold'] = bold
            if italic is not None:
                font_kwargs['italic'] = italic
            if font_color:
                font_kwargs['color'] = font_color.replace('#', '')

            font = Font(**font_kwargs) if font_kwargs else None

            fill = None
            if bg_color:
                fill = PatternFill(start_color=bg_color.replace('#', ''),
                                   end_color=bg_color.replace('#', ''),
                                   fill_type='solid')

            border_obj = None
            if border:
                side = Side(style=border)
                border_obj = Border(left=side, right=side, top=side, bottom=side)

            align = None
            if alignment:
                align_map = {'left': 'left', 'center': 'center', 'right': 'right'}
                align = Alignment(horizontal=align_map.get(alignment, 'left'))

            # 应用样式
            cells = ws[range]
            if hasattr(cells, '__iter__') and not isinstance(cells, str):
                # 多行范围
                for row in cells:
                    if hasattr(row, '__iter__'):
                        for cell in row:
                            if font:
                                cell.font = font
                            if fill:
                                cell.fill = fill
                            if border_obj:
                                cell.border = border_obj
                            if align:
                                cell.alignment = align
                            if number_format:
                                cell.number_format = number_format
                    else:
                        # 单行
                        if font:
                            row.font = font
                        if fill:
                            row.fill = fill
                        if border_obj:
                            row.border = border_obj
                        if align:
                            row.alignment = align
                        if number_format:
                            row.number_format = number_format
            else:
                # 单个单元格
                if font:
                    cells.font = font
                if fill:
                    cells.fill = fill
                if border_obj:
                    cells.border = border_obj
                if align:
                    cells.alignment = align
                if number_format:
                    cells.number_format = number_format

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "message": f"格式已应用到 {range}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("cell_merge",
                   required_params=['path', 'range'],
                   optional_params=['sheet', 'unmerge'])
    @staticmethod
    def merge_cell(path: str,
                   range: str,
                   sheet: Optional[str] = None,
                   unmerge: bool = False) -> Dict[str, Any]:
        """合并/取消合并单元格"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            if unmerge:
                ws.unmerge_cells(range)
                action = "取消合并"
            else:
                ws.merge_cells(range)
                action = "合并"

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "action": action,
                "message": f"已{action}单元格: {range}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("cell_formula",
                   required_params=['path', 'cell', 'formula'],
                   optional_params=['sheet'])
    @staticmethod
    def set_formula(path: str,
                    cell: str,
                    formula: str,
                    sheet: Optional[str] = None) -> Dict[str, Any]:
        """
        设置公式

        Args:
            cell: 单元格位置如"E1"
            formula: 公式如"=SUM(A1:D1)"
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 确保公式以=开头
            if not formula.startswith('='):
                formula = '=' + formula

            ws[cell] = formula
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "cell": cell,
                "formula": formula,
                "message": f"公式已设置: {cell} = {formula}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 行列操作 ==========

    @register_tool("row_insert",
                   required_params=['path', 'row'],
                   optional_params=['sheet', 'count'])
    @staticmethod
    def insert_row(path: str,
                   row: int,
                   sheet: Optional[str] = None,
                   count: int = 1) -> Dict[str, Any]:
        """插入行"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            ws.insert_rows(row, amount=count)
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "row": row,
                "count": count,
                "message": f"已在第 {row} 行插入 {count} 行"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("row_delete",
                   required_params=['path', 'row'],
                   optional_params=['sheet', 'count'])
    @staticmethod
    def delete_row(path: str,
                   row: int,
                   sheet: Optional[str] = None,
                   count: int = 1) -> Dict[str, Any]:
        """删除行"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            ws.delete_rows(row, amount=count)
            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "row": row,
                "count": count,
                "message": f"已删除第 {row} 行起的 {count} 行"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("col_insert",
                   required_params=['path', 'col'],
                   optional_params=['sheet', 'count'])
    @staticmethod
    def insert_col(path: str,
                   col: Union[int, str],
                   sheet: Optional[str] = None,
                   count: int = 1) -> Dict[str, Any]:
        """插入列（col可以是数字1或字母"A"）"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            col_idx = _col_letter_to_index(col)
            ws.insert_cols(col_idx, amount=count)
            wb.save(path)
            wb.close()

            col_letter = _col_index_to_letter(col_idx)
            return {
                "success": True,
                "sheet": sheet,
                "column": col_letter,
                "count": count,
                "message": f"已在第 {col_letter} 列插入 {count} 列"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("col_delete",
                   required_params=['path', 'col'],
                   optional_params=['sheet', 'count'])
    @staticmethod
    def delete_col(path: str,
                   col: Union[int, str],
                   sheet: Optional[str] = None,
                   count: int = 1) -> Dict[str, Any]:
        """删除列"""
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            col_idx = _col_letter_to_index(col)
            ws.delete_cols(col_idx, amount=count)
            wb.save(path)
            wb.close()

            col_letter = _col_index_to_letter(col_idx)
            return {
                "success": True,
                "sheet": sheet,
                "column": col_letter,
                "count": count,
                "message": f"已删除第 {col_letter} 列起的 {count} 列"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 高级功能 ==========

    @register_tool("excel_to_word",
                   required_params=['excel_path', 'word_path'],
                   optional_params=['sheet', 'range', 'style'])
    @staticmethod
    def to_word(excel_path: str,
                word_path: str,
                sheet: Optional[str] = None,
                range: Optional[str] = None,
                style: Optional[str] = None) -> Dict[str, Any]:
        """
        将Excel表格插入Word文档

        Args:
            excel_path: Excel文件路径
            word_path: Word文档路径（已存在则追加，不存在则创建）
            sheet: 工作表名称
            range: 数据范围如"A1:D10"
            style: Word表格样式
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            try:
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                return {"success": False, "error": "需要安装python-docx"}

            from openpyxl import load_workbook

            if not os.path.exists(excel_path):
                return {"success": False, "error": f"Excel文件不存在: {excel_path}"}

            # 读取Excel数据
            wb = load_workbook(excel_path, data_only=True)

            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 获取数据
            if range:
                data = []
                for row in ws[range]:
                    data.append([cell.value for cell in row])
            else:
                data = []
                for row in ws.iter_rows(min_row=1, max_row=ws.max_row,
                                        min_col=1, max_col=ws.max_column, values_only=True):
                    data.append(list(row))

            wb.close()

            if not data:
                return {"success": False, "error": "没有数据可插入"}

            # 打开或创建Word文档
            if os.path.exists(word_path):
                doc = Document(word_path)
            else:
                doc = Document()
                # 创建目录
                dir_path = os.path.dirname(word_path)
                if dir_path and not os.path.exists(dir_path):
                    os.makedirs(dir_path)

            # 创建表格
            rows = len(data)
            cols = len(data[0]) if data else 0

            table = doc.add_table(rows=rows, cols=cols)
            if style:
                try:
                    table.style = style
                except Exception:
                    table.style = 'Table Grid'
            else:
                table.style = 'Table Grid'

            # 填充数据
            for i, row_data in enumerate(data):
                for j, value in enumerate(row_data):
                    table.cell(i, j).text = str(value) if value is not None else ''

            doc.save(word_path)

            return {
                "success": True,
                "excel_path": excel_path,
                "word_path": word_path,
                "sheet": sheet,
                "rows": rows,
                "cols": cols,
                "message": f"已将 {rows}x{cols} 表格插入Word文档"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("excel_status",
                   required_params=[],
                   optional_params=[])
    @staticmethod
    def get_status() -> Dict[str, Any]:
        """
        获取Excel模块状态

        Returns:
            {success, openpyxl_available, version, message}
        """
        try:
            available = check_import("openpyxl")
            version = None

            if available:
                import openpyxl
                version = openpyxl.__version__

            return {
                "success": True,
                "openpyxl_available": available,
                "version": version,
                "features": [
                    "工作簿操作（创建、读取、保存）",
                    "工作表管理（增删改查、复制）",
                    "单元格操作（读写、格式化、合并、公式）",
                    "行列操作（插入、删除）",
                    "与Word集成（表格导入）",
                    "高级功能：批量公式、数据排序、筛选、验证",
                    "高级功能：统计摘要、条件格式、命名范围",
                    "高级功能：图表创建与修改"
                ] if available else [],
                "message": f"Excel模块可用，openpyxl版本: {version}" if available else "需要安装openpyxl: pip install openpyxl"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 高级功能：公式增强 ==========

    @register_tool("formula_batch",
                   required_params=['path', 'range', 'formula'],
                   optional_params=['sheet'])
    @staticmethod
    def formula_batch(path: str,
                      range: str,
                      formula: str,
                      sheet: Optional[str] = None) -> Dict[str, Any]:
        """
        批量设置公式

        Args:
            path: Excel文件路径
            range: 目标范围如"E2:E100"
            formula: 公式模板，支持{row}占位符，如"=SUM(A{row}:D{row})"
            sheet: 工作表名称

        Example:
            formula_batch(path="data.xlsx", range="E2:E10", formula="=SUM(A{row}:D{row})")
            # E2 = =SUM(A2:D2)
            # E3 = =SUM(A3:D3)
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 解析范围
            cells = ws[range]
            count = 0

            if hasattr(cells, '__iter__') and not isinstance(cells, str):
                for row in cells:
                    if hasattr(row, '__iter__'):
                        for cell in row:
                            row_num = cell.row
                            cell_formula = formula.replace('{row}', str(row_num))
                            if not cell_formula.startswith('='):
                                cell_formula = '=' + cell_formula
                            cell.value = cell_formula
                            count += 1
                    else:
                        row_num = row.row
                        cell_formula = formula.replace('{row}', str(row_num))
                        if not cell_formula.startswith('='):
                            cell_formula = '=' + cell_formula
                        row.value = cell_formula
                        count += 1
            else:
                row_num = cells.row
                cell_formula = formula.replace('{row}', str(row_num))
                if not cell_formula.startswith('='):
                    cell_formula = '=' + cell_formula
                cells.value = cell_formula
                count = 1

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "formula_template": formula,
                "cells_updated": count,
                "message": f"已为 {count} 个单元格设置公式"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("formula_quick",
                   required_params=['path', 'data_range', 'function', 'output_cell'],
                   optional_params=['sheet'])
    @staticmethod
    def formula_quick(path: str,
                      data_range: str,
                      function: str,
                      output_cell: str,
                      sheet: Optional[str] = None) -> Dict[str, Any]:
        """
        快捷函数生成

        Args:
            path: Excel文件路径
            data_range: 数据范围如"A1:A100"
            function: 函数类型 (sum/average/max/min/count/counta/stdev/var/median)
            output_cell: 输出单元格如"B1"
            sheet: 工作表名称
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            # 函数映射
            func_map = {
                'sum': 'SUM',
                'average': 'AVERAGE',
                'avg': 'AVERAGE',
                'max': 'MAX',
                'min': 'MIN',
                'count': 'COUNT',
                'counta': 'COUNTA',
                'stdev': 'STDEV',
                'var': 'VAR',
                'median': 'MEDIAN'
            }

            func_name = func_map.get(function.lower())
            if not func_name:
                return {"success": False, "error": f"不支持的函数: {function}，支持: {list(func_map.keys())}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 生成公式
            formula = f"={func_name}({data_range})"
            ws[output_cell] = formula

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "data_range": data_range,
                "function": func_name,
                "output_cell": output_cell,
                "formula": formula,
                "message": f"已设置公式: {output_cell} = {formula}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 高级功能：数据操作 ==========

    @register_tool("data_sort",
                   required_params=['path', 'range'],
                   optional_params=['sheet', 'sort_by', 'has_header'])
    @staticmethod
    def data_sort(path: str,
                  range: str,
                  sheet: Optional[str] = None,
                  sort_by: Optional[List[Dict]] = None,
                  has_header: bool = True) -> Dict[str, Any]:
        """
        数据排序

        Args:
            path: Excel文件路径
            range: 排序范围如"A1:D100"
            sheet: 工作表名称
            sort_by: 排序规则列表 [{"col": "C", "order": "desc"}, {"col": "A", "order": "asc"}]
            has_header: 是否有标题行
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 读取数据
            data = []
            header = None
            for i, row in enumerate(ws[range]):
                row_data = [cell.value for cell in row]
                if i == 0 and has_header:
                    header = row_data
                else:
                    data.append(row_data)

            if not data:
                return {"success": False, "error": "没有数据可排序"}

            # 确定排序规则
            if not sort_by:
                sort_by = [{"col": "A", "order": "asc"}]

            # 获取范围的起始列
            start_cell = range.split(':')[0]
            start_col_letter, start_row = coordinate_from_string(start_cell)
            start_col_idx = column_index_from_string(start_col_letter)

            # 多级排序（从后向前处理排序规则）
            for rule in reversed(sort_by):
                col = rule.get('col', 'A')
                col_idx = _col_letter_to_index(col) - start_col_idx
                reverse = rule.get('order', 'asc').lower() == 'desc'

                def get_key(row, idx=col_idx):
                    if 0 <= idx < len(row):
                        val = row[idx]
                        return (0, val) if val is not None else (1, "")
                    return (1, "")

                data.sort(key=get_key, reverse=reverse)

            # 写回数据
            cells = list(ws[range])
            data_start = 1 if has_header else 0

            for i, row_data in enumerate(data):
                row_cells = cells[data_start + i] if data_start + i < len(cells) else None
                if row_cells:
                    if hasattr(row_cells, '__iter__'):
                        for j, cell in enumerate(row_cells):
                            if j < len(row_data):
                                cell.value = row_data[j]
                    else:
                        row_cells.value = row_data[0] if row_data else None

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "rows_sorted": len(data),
                "sort_by": sort_by,
                "message": f"已排序 {len(data)} 行数据"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("data_filter",
                   required_params=['path', 'range'],
                   optional_params=['sheet', 'clear'])
    @staticmethod
    def data_filter(path: str,
                    range: str,
                    sheet: Optional[str] = None,
                    clear: bool = False) -> Dict[str, Any]:
        """
        自动筛选

        Args:
            path: Excel文件路径
            range: 筛选范围如"A1:D100"
            sheet: 工作表名称
            clear: 是否清除筛选
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            if clear:
                ws.auto_filter.ref = None
                action = "已清除筛选"
            else:
                ws.auto_filter.ref = range
                action = f"已设置筛选范围: {range}"

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range if not clear else None,
                "action": "clear" if clear else "set",
                "message": action
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("data_validate",
                   required_params=['path', 'range', 'type'],
                   optional_params=['sheet', 'values', 'min_val', 'max_val', 'formula', 'error_message'])
    @staticmethod
    def data_validate(path: str,
                      range: str,
                      type: str,
                      sheet: Optional[str] = None,
                      values: Optional[List[str]] = None,
                      min_val: Optional[float] = None,
                      max_val: Optional[float] = None,
                      formula: Optional[str] = None,
                      error_message: Optional[str] = None) -> Dict[str, Any]:
        """
        数据验证

        Args:
            path: Excel文件路径
            range: 验证范围如"B2:B100"
            type: 验证类型 (list/whole/decimal/date/text_length/custom)
            sheet: 工作表名称
            values: 下拉列表值（type=list时）
            min_val/max_val: 数值范围
            formula: 自定义公式（type=custom时）
            error_message: 错误提示信息
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.worksheet.datavalidation import DataValidation

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 创建数据验证
            if type == 'list':
                if not values:
                    return {"success": False, "error": "下拉列表需要提供values参数"}
                formula1 = '"' + ','.join(values) + '"'
                dv = DataValidation(type="list", formula1=formula1, allow_blank=True)
            elif type in ['whole', 'decimal']:
                dv = DataValidation(type=type, operator="between",
                                    formula1=str(min_val) if min_val is not None else "0",
                                    formula2=str(max_val) if max_val is not None else "999999999")
            elif type == 'date':
                dv = DataValidation(type="date", operator="between")
            elif type == 'text_length':
                dv = DataValidation(type="textLength", operator="between",
                                    formula1=str(int(min_val)) if min_val is not None else "0",
                                    formula2=str(int(max_val)) if max_val is not None else "255")
            elif type == 'custom':
                if not formula:
                    return {"success": False, "error": "自定义验证需要提供formula参数"}
                dv = DataValidation(type="custom", formula1=formula)
            else:
                return {"success": False, "error": f"不支持的验证类型: {type}"}

            if error_message:
                dv.error = error_message
                dv.errorTitle = "输入错误"

            dv.add(range)
            ws.add_data_validation(dv)

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "type": type,
                "message": f"已为 {range} 设置数据验证"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("data_deduplicate",
                   required_params=['path', 'range'],
                   optional_params=['sheet', 'columns', 'keep'])
    @staticmethod
    def data_deduplicate(path: str,
                         range: str,
                         sheet: Optional[str] = None,
                         columns: Optional[List[str]] = None,
                         keep: str = 'first') -> Dict[str, Any]:
        """
        去除重复行

        Args:
            path: Excel文件路径
            range: 数据范围如"A1:D100"
            sheet: 工作表名称
            columns: 用于判断重复的列，如["A", "B"]，None表示全部列
            keep: 保留策略 (first/last)
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 读取数据
            data = []
            row_refs = []
            for row in ws[range]:
                row_data = [cell.value for cell in row]
                data.append(row_data)
                row_refs.append(row[0].row if hasattr(row, '__iter__') else row.row)

            if not data:
                return {"success": False, "error": "没有数据"}

            # 获取起始列索引
            start_cell = range.split(':')[0]
            start_col_letter, _ = coordinate_from_string(start_cell)
            start_col_idx = column_index_from_string(start_col_letter)

            # 确定用于判断重复的列索引
            if columns:
                check_indices = [_col_letter_to_index(c) - start_col_idx for c in columns]
            else:
                check_indices = list(range(len(data[0]))) if data else []

            # 去重
            seen = {}
            unique_data = []
            duplicates_removed = 0

            if keep == 'last':
                data = list(reversed(data))
                row_refs = list(reversed(row_refs))

            for i, row in enumerate(data):
                key = tuple(row[j] if j < len(row) else None for j in check_indices)
                if key not in seen:
                    seen[key] = True
                    unique_data.append((row_refs[i], row))
                else:
                    duplicates_removed += 1

            if keep == 'last':
                unique_data = list(reversed(unique_data))

            # 清空原范围并写入去重后的数据
            cells = list(ws[range])
            for row_cells in cells:
                if hasattr(row_cells, '__iter__'):
                    for cell in row_cells:
                        cell.value = None
                else:
                    row_cells.value = None

            # 写入去重后的数据
            for i, (_, row_data) in enumerate(unique_data):
                if i < len(cells):
                    row_cells = cells[i]
                    if hasattr(row_cells, '__iter__'):
                        for j, cell in enumerate(row_cells):
                            if j < len(row_data):
                                cell.value = row_data[j]
                    else:
                        row_cells.value = row_data[0] if row_data else None

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "original_rows": len(data),
                "unique_rows": len(unique_data),
                "duplicates_removed": duplicates_removed,
                "message": f"已去除 {duplicates_removed} 行重复数据，保留 {len(unique_data)} 行"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("data_fill",
                   required_params=['path', 'range', 'type'],
                   optional_params=['sheet', 'start', 'step'])
    @staticmethod
    def data_fill(path: str,
                  range: str,
                  type: str,
                  sheet: Optional[str] = None,
                  start: Optional[float] = None,
                  step: Optional[float] = None) -> Dict[str, Any]:
        """
        序列填充

        Args:
            path: Excel文件路径
            range: 填充范围如"A1:A10"
            type: 填充类型 (linear/growth/date)
            sheet: 工作表名称
            start: 起始值
            step: 步长（等差）或比率（等比）
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from datetime import datetime, timedelta

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 设置默认值
            if start is None:
                start = 1
            if step is None:
                step = 1

            # 获取单元格列表
            cells = []
            for row in ws[range]:
                if hasattr(row, '__iter__'):
                    for cell in row:
                        cells.append(cell)
                else:
                    cells.append(row)

            # 填充数据
            if type == 'linear':
                # 等差序列
                for i, cell in enumerate(cells):
                    cell.value = start + i * step
            elif type == 'growth':
                # 等比序列
                value = start
                for i, cell in enumerate(cells):
                    cell.value = value
                    value *= step
            elif type == 'date':
                # 日期序列
                if isinstance(start, (int, float)):
                    current_date = datetime.now()
                else:
                    current_date = start if isinstance(start, datetime) else datetime.now()

                for i, cell in enumerate(cells):
                    cell.value = current_date + timedelta(days=int(i * step))
            else:
                return {"success": False, "error": f"不支持的填充类型: {type}"}

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "type": type,
                "start": start,
                "step": step,
                "cells_filled": len(cells),
                "message": f"已填充 {len(cells)} 个单元格"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 高级功能：统计与格式 ==========

    @register_tool("stats_summary",
                   required_params=['path', 'data_range'],
                   optional_params=['sheet', 'output_cell', 'metrics'])
    @staticmethod
    def stats_summary(path: str,
                      data_range: str,
                      sheet: Optional[str] = None,
                      output_cell: Optional[str] = None,
                      metrics: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        统计摘要

        Args:
            path: Excel文件路径
            data_range: 数据范围
            sheet: 工作表名称
            output_cell: 输出起始单元格（可选，不指定则只返回结果）
            metrics: 统计指标 ["sum", "average", "max", "min", "count", "stdev", "var", "median"]
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string
            import statistics

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path, data_only=True)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 读取数值数据
            values = []
            for row in ws[data_range]:
                if hasattr(row, '__iter__'):
                    for cell in row:
                        if isinstance(cell.value, (int, float)):
                            values.append(cell.value)
                else:
                    if isinstance(row.value, (int, float)):
                        values.append(row.value)

            if not values:
                return {"success": False, "error": "没有数值数据"}

            # 默认统计指标
            if not metrics:
                metrics = ["sum", "average", "max", "min", "count"]

            # 计算统计量
            stats = {}
            metric_labels = {}

            if "sum" in metrics:
                stats["sum"] = sum(values)
                metric_labels["sum"] = "总和"
            if "average" in metrics or "avg" in metrics:
                stats["average"] = sum(values) / len(values)
                metric_labels["average"] = "平均值"
            if "max" in metrics:
                stats["max"] = max(values)
                metric_labels["max"] = "最大值"
            if "min" in metrics:
                stats["min"] = min(values)
                metric_labels["min"] = "最小值"
            if "count" in metrics:
                stats["count"] = len(values)
                metric_labels["count"] = "计数"
            if "stdev" in metrics and len(values) > 1:
                stats["stdev"] = statistics.stdev(values)
                metric_labels["stdev"] = "标准差"
            if "var" in metrics and len(values) > 1:
                stats["var"] = statistics.variance(values)
                metric_labels["var"] = "方差"
            if "median" in metrics:
                stats["median"] = statistics.median(values)
                metric_labels["median"] = "中位数"

            # 如果指定了输出单元格，写入结果
            if output_cell:
                wb.close()
                wb = load_workbook(path)
                ws = wb[sheet]

                col_letter, row_num = coordinate_from_string(output_cell)
                col_idx = column_index_from_string(col_letter)

                for i, (key, value) in enumerate(stats.items()):
                    label = metric_labels.get(key, key)
                    ws.cell(row=row_num + i, column=col_idx).value = label
                    ws.cell(row=row_num + i, column=col_idx + 1).value = value

                wb.save(path)

            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "data_range": data_range,
                "statistics": stats,
                "output_cell": output_cell,
                "message": f"统计完成: {len(values)} 个数值"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("conditional_format",
                   required_params=['path', 'range', 'rule'],
                   optional_params=['sheet', 'value', 'value2', 'format', 'color_scale', 'data_bar'])
    @staticmethod
    def conditional_format(path: str,
                           range: str,
                           rule: str,
                           sheet: Optional[str] = None,
                           value: Optional[Any] = None,
                           value2: Optional[Any] = None,
                           format: Optional[Dict] = None,
                           color_scale: Optional[Dict] = None,
                           data_bar: Optional[Dict] = None) -> Dict[str, Any]:
        """
        条件格式

        Args:
            path: Excel文件路径
            range: 格式化范围
            rule: 规则类型
                - greater_than / less_than / equal / between
                - color_scale / data_bar
            sheet: 工作表名称
            value: 比较值
            value2: 第二个比较值（between时使用）
            format: 格式设置 {"bg_color": "FF0000", "font_color": "FFFFFF", "bold": True}
            color_scale: 色阶设置 {"min_color": "F8696B", "max_color": "63BE7B"}
            data_bar: 数据条设置 {"color": "638EC6"}
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.formatting.rule import (
                CellIsRule, ColorScaleRule, DataBarRule
            )
            from openpyxl.styles import PatternFill, Font

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 创建条件格式规则
            if rule == 'color_scale':
                if not color_scale:
                    color_scale = {"min_color": "F8696B", "max_color": "63BE7B"}

                mid_color = color_scale.get("mid_color")
                if mid_color:
                    cf_rule = ColorScaleRule(
                        start_type='min', start_color=color_scale.get("min_color", "F8696B"),
                        mid_type='percentile', mid_value=50, mid_color=mid_color,
                        end_type='max', end_color=color_scale.get("max_color", "63BE7B")
                    )
                else:
                    cf_rule = ColorScaleRule(
                        start_type='min', start_color=color_scale.get("min_color", "F8696B"),
                        end_type='max', end_color=color_scale.get("max_color", "63BE7B")
                    )

            elif rule == 'data_bar':
                if not data_bar:
                    data_bar = {"color": "638EC6"}
                cf_rule = DataBarRule(
                    start_type='min', end_type='max',
                    color=data_bar.get("color", "638EC6")
                )

            elif rule in ['greater_than', 'less_than', 'equal', 'between', 'not_between']:
                # 准备填充和字体
                fill = None
                font = None

                if format:
                    if format.get("bg_color"):
                        fill = PatternFill(start_color=format["bg_color"].replace('#', ''),
                                          end_color=format["bg_color"].replace('#', ''),
                                          fill_type='solid')
                    if format.get("font_color") or format.get("bold"):
                        font = Font(
                            color=format.get("font_color", "000000").replace('#', ''),
                            bold=format.get("bold", False)
                        )

                # 映射规则类型
                op_map = {
                    'greater_than': 'greaterThan',
                    'less_than': 'lessThan',
                    'equal': 'equal',
                    'between': 'between',
                    'not_between': 'notBetween'
                }

                if rule == 'between' or rule == 'not_between':
                    cf_rule = CellIsRule(
                        operator=op_map[rule],
                        formula=[str(value), str(value2)],
                        fill=fill, font=font
                    )
                else:
                    cf_rule = CellIsRule(
                        operator=op_map[rule],
                        formula=[str(value)],
                        fill=fill, font=font
                    )
            else:
                return {"success": False, "error": f"不支持的规则类型: {rule}"}

            ws.conditional_formatting.add(range, cf_rule)

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "range": range,
                "rule": rule,
                "message": f"已为 {range} 添加条件格式"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("named_range",
                   required_params=['path', 'action'],
                   optional_params=['name', 'range', 'sheet'])
    @staticmethod
    def named_range(path: str,
                    action: str,
                    name: Optional[str] = None,
                    range: Optional[str] = None,
                    sheet: Optional[str] = None) -> Dict[str, Any]:
        """
        命名范围操作

        Args:
            path: Excel文件路径
            action: 操作类型 (create/list/delete)
            name: 范围名称
            range: 单元格范围（创建时需要）
            sheet: 工作表名称（创建时需要）
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.workbook.defined_name import DefinedName

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            if action == 'create':
                if not name or not range:
                    return {"success": False, "error": "创建命名范围需要name和range参数"}

                # 确定工作表
                if sheet:
                    if sheet not in wb.sheetnames:
                        return {"success": False, "error": f"工作表不存在: {sheet}"}
                else:
                    sheet = wb.active.title

                # 创建完整引用
                ref = f"'{sheet}'!{range}"

                # 添加命名范围
                defn = DefinedName(name, attr_text=ref)
                wb.defined_names[name] = defn

                wb.save(path)
                wb.close()

                return {
                    "success": True,
                    "action": "create",
                    "name": name,
                    "range": range,
                    "sheet": sheet,
                    "message": f"已创建命名范围: {name} = {ref}"
                }

            elif action == 'list':
                names = []
                for defn in wb.defined_names.values():
                    names.append({
                        "name": defn.name,
                        "value": defn.attr_text
                    })
                wb.close()

                return {
                    "success": True,
                    "action": "list",
                    "named_ranges": names,
                    "count": len(names),
                    "message": f"共 {len(names)} 个命名范围"
                }

            elif action == 'delete':
                if not name:
                    return {"success": False, "error": "删除命名范围需要name参数"}

                if name in wb.defined_names:
                    del wb.defined_names[name]
                    wb.save(path)
                    wb.close()
                    return {
                        "success": True,
                        "action": "delete",
                        "name": name,
                        "message": f"已删除命名范围: {name}"
                    }
                else:
                    wb.close()
                    return {"success": False, "error": f"命名范围不存在: {name}"}

            else:
                wb.close()
                return {"success": False, "error": f"不支持的操作: {action}"}

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 高级功能：图表 ==========

    @register_tool("chart_create",
                   required_params=['path', 'type', 'data_range'],
                   optional_params=['sheet', 'position', 'title', 'x_title', 'y_title', 'style'])
    @staticmethod
    def chart_create(path: str,
                     type: str,
                     data_range: str,
                     sheet: Optional[str] = None,
                     position: Optional[str] = None,
                     title: Optional[str] = None,
                     x_title: Optional[str] = None,
                     y_title: Optional[str] = None,
                     style: Optional[int] = None) -> Dict[str, Any]:
        """
        创建图表

        Args:
            path: Excel文件路径
            type: 图表类型 (bar/column/line/pie/scatter/area/doughnut/radar)
            data_range: 数据范围如"A1:B10"
            sheet: 工作表名称
            position: 图表位置如"E1"
            title: 图表标题
            x_title: X轴标题
            y_title: Y轴标题
            style: 图表样式编号(1-48)
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook
            from openpyxl.chart import (
                BarChart, LineChart, PieChart, ScatterChart,
                AreaChart, DoughnutChart, RadarChart, Reference
            )
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 图表类型映射
            chart_map = {
                'bar': BarChart,
                'column': BarChart,
                'line': LineChart,
                'pie': PieChart,
                'scatter': ScatterChart,
                'area': AreaChart,
                'doughnut': DoughnutChart,
                'radar': RadarChart
            }

            chart_class = chart_map.get(type.lower())
            if not chart_class:
                return {"success": False, "error": f"不支持的图表类型: {type}"}

            # 创建图表
            chart = chart_class()

            if type.lower() == 'bar':
                chart.type = "bar"
            elif type.lower() == 'column':
                chart.type = "col"

            # 设置标题
            if title:
                chart.title = title
            if x_title:
                chart.x_axis.title = x_title
            if y_title:
                chart.y_axis.title = y_title
            if style:
                chart.style = style

            # 解析数据范围
            range_parts = data_range.split(':')
            start_cell = range_parts[0]
            end_cell = range_parts[1] if len(range_parts) > 1 else start_cell

            start_col, start_row = coordinate_from_string(start_cell)
            end_col, end_row = coordinate_from_string(end_cell)

            min_col = column_index_from_string(start_col)
            max_col = column_index_from_string(end_col)
            min_row = start_row
            max_row = end_row

            # 创建数据引用
            data = Reference(ws, min_col=min_col + 1, min_row=min_row,
                            max_col=max_col, max_row=max_row)
            categories = Reference(ws, min_col=min_col, min_row=min_row + 1,
                                   max_row=max_row)

            chart.add_data(data, titles_from_data=True)

            if type.lower() not in ['pie', 'doughnut']:
                chart.set_categories(categories)

            # 设置图表位置
            if position:
                chart.anchor = position
            else:
                chart.anchor = "E1"

            ws.add_chart(chart)

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "type": type,
                "data_range": data_range,
                "position": position or "E1",
                "title": title,
                "message": f"已创建 {type} 图表"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("excel_chart_modify",
                   required_params=['path', 'chart_index'],
                   optional_params=['sheet', 'title', 'x_title', 'y_title', 'style', 'width', 'height'])
    @staticmethod
    def chart_modify(path: str,
                     chart_index: int,
                     sheet: Optional[str] = None,
                     title: Optional[str] = None,
                     x_title: Optional[str] = None,
                     y_title: Optional[str] = None,
                     style: Optional[int] = None,
                     width: Optional[float] = None,
                     height: Optional[float] = None) -> Dict[str, Any]:
        """
        修改图表

        Args:
            path: Excel文件路径
            chart_index: 图表索引（从0开始）
            sheet: 工作表名称
            title: 新标题
            x_title: X轴标题
            y_title: Y轴标题
            style: 新样式
            width/height: 尺寸（厘米）
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from openpyxl import load_workbook

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 获取图表
            charts = ws._charts
            if chart_index >= len(charts):
                return {"success": False, "error": f"图表索引超出范围，共 {len(charts)} 个图表"}

            chart = charts[chart_index]

            # 修改属性
            modified = []
            if title is not None:
                chart.title = title
                modified.append("title")
            if x_title is not None:
                chart.x_axis.title = x_title
                modified.append("x_title")
            if y_title is not None:
                chart.y_axis.title = y_title
                modified.append("y_title")
            if style is not None:
                chart.style = style
                modified.append("style")
            if width is not None:
                chart.width = width
                modified.append("width")
            if height is not None:
                chart.height = height
                modified.append("height")

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "sheet": sheet,
                "chart_index": chart_index,
                "modified": modified,
                "message": f"已修改图表属性: {', '.join(modified)}"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ========== 数据透视汇总 ==========

    @register_tool("pivot_create",
                   required_params=['path', 'source_range', 'target_cell', 'rows', 'values'],
                   optional_params=['sheet', 'agg_func', 'target_sheet', 'include_totals'])
    @staticmethod
    def pivot_create(path: str,
                     source_range: str,
                     target_cell: str,
                     rows: List[str],
                     values: List[str],
                     sheet: Optional[str] = None,
                     agg_func: str = 'sum',
                     target_sheet: Optional[str] = None,
                     include_totals: bool = True) -> Dict[str, Any]:
        """
        创建数据透视汇总表

        Args:
            path: Excel文件路径
            source_range: 数据源范围如"A1:D100"（第一行为表头）
            target_cell: 输出起始单元格如"F1"
            rows: 分组行字段名列表（表头名称）
            values: 聚合值字段名列表（表头名称）
            sheet: 源数据工作表名称
            agg_func: 聚合函数 (sum/average/count/max/min)
            target_sheet: 输出工作表名称（默认同源）
            include_totals: 是否包含总计行

        Returns:
            {success, message}
        """
        try:
            if not check_import("openpyxl"):
                return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}

            from collections import defaultdict
            from openpyxl import load_workbook
            from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
            from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

            if not os.path.exists(path):
                return {"success": False, "error": f"文件不存在: {path}"}

            wb = load_workbook(path)

            # 选择源工作表
            if sheet:
                if sheet not in wb.sheetnames:
                    return {"success": False, "error": f"工作表不存在: {sheet}"}
                ws = wb[sheet]
            else:
                ws = wb.active
                sheet = ws.title

            # 读取源数据
            data_rows = []
            header = None
            for i, row in enumerate(ws[source_range]):
                row_values = [cell.value for cell in row]
                if i == 0:
                    header = row_values
                else:
                    data_rows.append(row_values)

            if not header:
                return {"success": False, "error": "源范围无数据"}

            # 验证字段名
            for field in rows + values:
                if field not in header:
                    return {"success": False, "error": f"字段 '{field}' 不在表头中。可用字段: {header}"}

            # 获取字段索引
            row_indices = [header.index(f) for f in rows]
            val_indices = [header.index(f) for f in values]

            # 按行字段分组聚合
            groups = defaultdict(lambda: defaultdict(list))
            for row_data in data_rows:
                key = tuple(row_data[i] for i in row_indices)
                for vi in val_indices:
                    val = row_data[vi]
                    if isinstance(val, (int, float)):
                        groups[key][vi].append(val)

            # 聚合函数
            agg_funcs = {
                'sum': lambda vals: sum(vals),
                'average': lambda vals: sum(vals) / len(vals) if vals else 0,
                'count': lambda vals: len(vals),
                'max': lambda vals: max(vals) if vals else 0,
                'min': lambda vals: min(vals) if vals else 0,
            }

            if agg_func.lower() not in agg_funcs:
                return {"success": False, "error": f"不支持的聚合函数: {agg_func}。支持: {list(agg_funcs.keys())}"}

            agg_fn = agg_funcs[agg_func.lower()]

            # 选择目标工作表
            if target_sheet:
                if target_sheet not in wb.sheetnames:
                    wb.create_sheet(title=target_sheet)
                target_ws = wb[target_sheet]
            else:
                target_ws = ws

            # 解析目标起始位置
            col_letter, start_row = coordinate_from_string(target_cell)
            start_col = column_index_from_string(col_letter)

            # 样式
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            total_font = Font(bold=True)
            total_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")

            # 写表头
            pivot_headers = rows + [f"{v} ({agg_func})" for v in values]
            for j, h in enumerate(pivot_headers):
                cell = target_ws.cell(row=start_row, column=start_col + j, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center')

            # 写数据行
            sorted_keys = sorted(groups.keys(), key=lambda k: tuple(str(x) for x in k))
            totals = defaultdict(float)
            total_counts = defaultdict(int)

            for i, key in enumerate(sorted_keys):
                current_row = start_row + 1 + i
                # 行字段
                for j, val in enumerate(key):
                    cell = target_ws.cell(row=current_row, column=start_col + j, value=val)
                    cell.border = thin_border
                # 值字段
                for j, vi in enumerate(val_indices):
                    agg_val = agg_fn(groups[key][vi]) if groups[key][vi] else 0
                    cell = target_ws.cell(row=current_row, column=start_col + len(rows) + j, value=agg_val)
                    cell.border = thin_border
                    cell.number_format = '#,##0.00' if isinstance(agg_val, float) else '#,##0'
                    # 累计总计
                    if agg_func.lower() in ('sum', 'count'):
                        totals[j] += agg_val
                    elif agg_func.lower() == 'max':
                        totals[j] = max(totals.get(j, float('-inf')), agg_val)
                    elif agg_func.lower() == 'min':
                        if j not in totals:
                            totals[j] = agg_val
                        else:
                            totals[j] = min(totals[j], agg_val)
                    elif agg_func.lower() == 'average':
                        totals[j] += agg_val
                        total_counts[j] += 1

            # 写总计行
            if include_totals and sorted_keys:
                total_row = start_row + 1 + len(sorted_keys)
                cell = target_ws.cell(row=total_row, column=start_col, value="总计")
                cell.font = total_font
                cell.fill = total_fill
                cell.border = thin_border

                for j in range(1, len(rows)):
                    cell = target_ws.cell(row=total_row, column=start_col + j)
                    cell.fill = total_fill
                    cell.border = thin_border

                for j in range(len(val_indices)):
                    if agg_func.lower() == 'average' and total_counts.get(j, 0) > 0:
                        total_val = totals[j] / total_counts[j]
                    else:
                        total_val = totals.get(j, 0)
                    cell = target_ws.cell(row=total_row, column=start_col + len(rows) + j, value=total_val)
                    cell.font = total_font
                    cell.fill = total_fill
                    cell.border = thin_border
                    cell.number_format = '#,##0.00' if isinstance(total_val, float) else '#,##0'

            wb.save(path)
            wb.close()

            return {
                "success": True,
                "path": path,
                "source_range": source_range,
                "target_cell": target_cell,
                "rows": rows,
                "values": values,
                "agg_func": agg_func,
                "group_count": len(sorted_keys),
                "message": f"数据透视汇总完成: {len(sorted_keys)} 组数据"
            }

        except Exception as e:
            return {"success": False, "error": str(e)}
