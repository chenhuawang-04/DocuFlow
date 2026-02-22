"""
DocuFlow MCP - Batch Operations

Provides batch formatting and style application for multiple paragraphs
"""

from typing import Dict, Any, Optional, List
from docx import Document
from docuflow_mcp.core.registry import register_tool
from docuflow_mcp.utils.formatters import apply_font_format, apply_paragraph_format


class BatchOperations:
    """Batch operations for efficient multi-paragraph formatting"""

    @register_tool("batch_format_range",
                   required_params=['path', 'start_index', 'end_index'],
                   optional_params=['font_name', 'font_size', 'font_color', 'bold', 'italic',
                                    'underline', 'alignment', 'line_spacing', 'space_before',
                                    'space_after', 'first_line_indent'])
    @staticmethod
    def format_range(path: str, start_index: int, end_index: int,
                     font_name: Optional[str] = None,
                     font_size: Optional[str] = None,
                     font_color: Optional[str] = None,
                     bold: Optional[bool] = None,
                     italic: Optional[bool] = None,
                     underline: Optional[bool] = None,
                     alignment: Optional[str] = None,
                     line_spacing: Optional[float] = None,
                     space_before: Optional[str] = None,
                     space_after: Optional[str] = None,
                     first_line_indent: Optional[str] = None) -> Dict[str, Any]:
        """
        Format a range of paragraphs with the same formatting.

        This is much more efficient than calling paragraph_modify multiple times.

        Args:
            path: Document path
            start_index: Starting paragraph index (0-based)
            end_index: Ending paragraph index (inclusive)
            font_name: Font name (e.g., '宋体', 'Arial')
            font_size: Font size (e.g., '12pt', '14')
            font_color: Font color (e.g., '#FF0000', 'red')
            bold: Bold formatting
            italic: Italic formatting
            underline: Underline formatting
            alignment: Paragraph alignment ('left', 'center', 'right', 'justify')
            line_spacing: Line spacing multiplier (e.g., 1.5, 2.0)
            space_before: Space before paragraph (e.g., '12pt')
            space_after: Space after paragraph (e.g., '12pt')
            first_line_indent: First line indent (e.g., '2em')

        Returns:
            Dict with success status and formatted paragraph count
        """
        doc = Document(path)
        paragraphs = doc.paragraphs

        # Validate indices
        if start_index < 0 or end_index >= len(paragraphs):
            return {
                "success": False,
                "error": f"索引超出范围: 文档共有 {len(paragraphs)} 个段落，请求范围 [{start_index}, {end_index}]"
            }

        if start_index > end_index:
            return {
                "success": False,
                "error": f"起始索引({start_index})不能大于结束索引({end_index})"
            }

        formatted_count = 0

        # Apply formatting to each paragraph in range
        for idx in range(start_index, end_index + 1):
            para = paragraphs[idx]

            # Apply paragraph formatting
            apply_paragraph_format(
                para,
                alignment=alignment,
                line_spacing=line_spacing,
                space_before=space_before,
                space_after=space_after,
                first_line_indent=first_line_indent
            )

            # Apply font formatting to all runs in the paragraph
            for run in para.runs:
                apply_font_format(
                    run,
                    font_name=font_name,
                    font_size=font_size,
                    font_color=font_color,
                    bold=bold,
                    italic=italic,
                    underline=underline
                )

            formatted_count += 1

        doc.save(path)

        return {
            "success": True,
            "message": f"已格式化 {formatted_count} 个段落 (索引 {start_index} 到 {end_index})",
            "formatted_count": formatted_count,
            "start_index": start_index,
            "end_index": end_index
        }

    @register_tool("batch_apply_style",
                   required_params=['path', 'paragraph_indices', 'style_name'],
                   optional_params=[])
    @staticmethod
    def apply_style(path: str, paragraph_indices: List[int], style_name: str) -> Dict[str, Any]:
        """
        Apply a style to multiple paragraphs.

        Args:
            path: Document path
            paragraph_indices: List of paragraph indices (e.g., [0, 2, 5, 7])
            style_name: Style name to apply (e.g., 'Heading 1', 'Normal')

        Returns:
            Dict with success status and applied count
        """
        doc = Document(path)
        paragraphs = doc.paragraphs

        # Validate style exists
        try:
            doc.styles[style_name]
        except KeyError:
            return {
                "success": False,
                "error": f"样式不存在: {style_name}"
            }

        # Validate all indices
        invalid_indices = [idx for idx in paragraph_indices if idx < 0 or idx >= len(paragraphs)]
        if invalid_indices:
            return {
                "success": False,
                "error": f"以下索引超出范围: {invalid_indices}，文档共有 {len(paragraphs)} 个段落"
            }

        applied_count = 0

        # Apply style to each paragraph
        for idx in paragraph_indices:
            para = paragraphs[idx]
            para.style = style_name
            applied_count += 1

        doc.save(path)

        return {
            "success": True,
            "message": f"已将样式 '{style_name}' 应用到 {applied_count} 个段落",
            "applied_count": applied_count,
            "style_name": style_name,
            "paragraph_indices": paragraph_indices
        }

    @register_tool("batch_copy_format",
                   required_params=['path', 'source_index', 'target_indices'],
                   optional_params=[])
    @staticmethod
    def copy_format(path: str, source_index: int, target_indices: List[int]) -> Dict[str, Any]:
        """
        Copy formatting from one paragraph to multiple paragraphs (format painter).

        This is the equivalent of the "format painter" tool in Word.

        Args:
            path: Document path
            source_index: Source paragraph index to copy format from
            target_indices: List of target paragraph indices to apply format to

        Returns:
            Dict with success status and copied count
        """
        doc = Document(path)
        paragraphs = doc.paragraphs

        # Validate source index
        if source_index < 0 or source_index >= len(paragraphs):
            return {
                "success": False,
                "error": f"源段落索引超出范围: {source_index}，文档共有 {len(paragraphs)} 个段落"
            }

        # Validate target indices
        invalid_indices = [idx for idx in target_indices if idx < 0 or idx >= len(paragraphs)]
        if invalid_indices:
            return {
                "success": False,
                "error": f"以下目标索引超出范围: {invalid_indices}，文档共有 {len(paragraphs)} 个段落"
            }

        source_para = paragraphs[source_index]
        source_pf = source_para.paragraph_format

        # Extract source paragraph formatting
        source_format = {
            'alignment': source_pf.alignment,
            'line_spacing': source_pf.line_spacing,
            'space_before': source_pf.space_before,
            'space_after': source_pf.space_after,
            'first_line_indent': source_pf.first_line_indent,
            'left_indent': source_pf.left_indent,
            'right_indent': source_pf.right_indent
        }

        # Extract source font formatting from first run
        source_font = {}
        if source_para.runs:
            first_run = source_para.runs[0]
            source_font = {
                'name': first_run.font.name,
                'size': first_run.font.size,
                'color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None,
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline
            }

        copied_count = 0

        # Apply formatting to each target paragraph
        for idx in target_indices:
            target_para = paragraphs[idx]
            target_pf = target_para.paragraph_format

            # Copy paragraph formatting
            if source_format['alignment']:
                target_pf.alignment = source_format['alignment']
            if source_format['line_spacing']:
                target_pf.line_spacing = source_format['line_spacing']
            if source_format['space_before']:
                target_pf.space_before = source_format['space_before']
            if source_format['space_after']:
                target_pf.space_after = source_format['space_after']
            if source_format['first_line_indent']:
                target_pf.first_line_indent = source_format['first_line_indent']
            if source_format['left_indent']:
                target_pf.left_indent = source_format['left_indent']
            if source_format['right_indent']:
                target_pf.right_indent = source_format['right_indent']

            # Copy font formatting to all runs
            for run in target_para.runs:
                if source_font.get('name'):
                    run.font.name = source_font['name']
                if source_font.get('size'):
                    run.font.size = source_font['size']
                if source_font.get('color'):
                    run.font.color.rgb = source_font['color']
                if source_font.get('bold') is not None:
                    run.font.bold = source_font['bold']
                if source_font.get('italic') is not None:
                    run.font.italic = source_font['italic']
                if source_font.get('underline') is not None:
                    run.font.underline = source_font['underline']

            copied_count += 1

        doc.save(path)

        return {
            "success": True,
            "message": f"已将段落 {source_index} 的格式复制到 {copied_count} 个段落",
            "copied_count": copied_count,
            "source_index": source_index,
            "target_indices": target_indices
        }

    @register_tool("batch_replace_format",
                   required_params=['path', 'find_style', 'replace_options'],
                   optional_params=[])
    @staticmethod
    def replace_format(path: str, find_style: str, replace_options: Dict[str, Any]) -> Dict[str, Any]:
        """
        Find all paragraphs with a specific style and apply new formatting.

        Args:
            path: Document path
            find_style: Style name to search for (e.g., 'Normal', 'Heading 1')
            replace_options: Dictionary of formatting options to apply
                Example: {
                    "font_name": "宋体",
                    "font_size": "14pt",
                    "bold": true,
                    "alignment": "center"
                }

        Returns:
            Dict with success status and replaced count
        """
        doc = Document(path)
        paragraphs = doc.paragraphs

        # Validate style exists
        try:
            doc.styles[find_style]
        except KeyError:
            return {
                "success": False,
                "error": f"样式不存在: {find_style}"
            }

        # Find all paragraphs with the specified style
        matched_indices = []
        for idx, para in enumerate(paragraphs):
            if para.style.name == find_style:
                matched_indices.append(idx)

        if not matched_indices:
            return {
                "success": True,
                "message": f"未找到使用样式 '{find_style}' 的段落",
                "replaced_count": 0,
                "matched_indices": []
            }

        replaced_count = 0

        # Apply new formatting to matched paragraphs
        for idx in matched_indices:
            para = paragraphs[idx]

            # Apply paragraph formatting
            apply_paragraph_format(
                para,
                alignment=replace_options.get('alignment'),
                line_spacing=replace_options.get('line_spacing'),
                space_before=replace_options.get('space_before'),
                space_after=replace_options.get('space_after'),
                first_line_indent=replace_options.get('first_line_indent')
            )

            # Apply font formatting to all runs
            for run in para.runs:
                apply_font_format(
                    run,
                    font_name=replace_options.get('font_name'),
                    font_size=replace_options.get('font_size'),
                    font_color=replace_options.get('font_color'),
                    bold=replace_options.get('bold'),
                    italic=replace_options.get('italic'),
                    underline=replace_options.get('underline')
                )

            replaced_count += 1

        doc.save(path)

        return {
            "success": True,
            "message": f"已将 {replaced_count} 个使用样式 '{find_style}' 的段落重新格式化",
            "replaced_count": replaced_count,
            "find_style": find_style,
            "matched_indices": matched_indices
        }
