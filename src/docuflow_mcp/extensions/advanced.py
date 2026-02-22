"""
DocuFlow MCP - Advanced Operations Module

Provides advanced document operations including:
- Document comparison
- Statistical analysis
- Metadata extraction
- Enhanced export functions
"""

from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
import difflib
from collections import Counter
from ..core.registry import register_tool


class AdvancedOperations:
    """Advanced document operations"""

    @register_tool("doc_compare",
                   required_params=['path1', 'path2'],
                   optional_params=['compare_format', 'output_path'])
    @staticmethod
    def compare_documents(path1: str,
                          path2: str,
                          compare_format: bool = False,
                          output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Compare two documents and return differences

        Args:
            path1: Path to first document
            path2: Path to second document
            compare_format: Whether to compare formatting (default: False)
            output_path: Optional path to save comparison report

        Returns:
            Comparison results with added/removed/modified content
        """
        try:
            doc1 = Document(path1)
            doc2 = Document(path2)

            # Extract text from paragraphs
            text1 = [para.text for para in doc1.paragraphs]
            text2 = [para.text for para in doc2.paragraphs]

            # Use difflib to find differences
            differ = difflib.Differ()
            diff = list(differ.compare(text1, text2))

            # Categorize differences
            added = []
            removed = []
            modified = []

            i = 0
            while i < len(diff):
                line = diff[i]
                if line.startswith('+ '):
                    added.append(line[2:])
                elif line.startswith('- '):
                    removed.append(line[2:])
                elif line.startswith('? '):
                    # This indicates a modification
                    pass
                i += 1

            # Find modified paragraphs (both removed and added nearby)
            for rem in removed[:]:
                for add in added[:]:
                    # If similarity > 0.6, consider it a modification
                    similarity = difflib.SequenceMatcher(None, rem, add).ratio()
                    if similarity > 0.6:
                        modified.append({
                            'original': rem,
                            'modified': add,
                            'similarity': round(similarity, 2)
                        })
                        removed.remove(rem)
                        added.remove(add)
                        break

            # Compare document statistics
            stats_comparison = {
                'doc1': {
                    'paragraphs': len(doc1.paragraphs),
                    'tables': len(doc1.tables),
                    'sections': len(doc1.sections)
                },
                'doc2': {
                    'paragraphs': len(doc2.paragraphs),
                    'tables': len(doc2.tables),
                    'sections': len(doc2.sections)
                }
            }

            # Format comparison if requested
            format_diff = None
            if compare_format:
                format_diff = AdvancedOperations._compare_formats(doc1, doc2)

            result = {
                'success': True,
                'added_count': len(added),
                'removed_count': len(removed),
                'modified_count': len(modified),
                'added': added,
                'removed': removed,
                'modified': modified,
                'statistics': stats_comparison,
                'format_differences': format_diff
            }

            # Generate and save report if requested
            if output_path:
                report = AdvancedOperations._generate_comparison_report(result, path1, path2)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(report)
                result['report_path'] = output_path

            return result

        except Exception as e:
            return {
                'success': False,
                'error': f'文档对比失败: {str(e)}'
            }

    @staticmethod
    def _compare_formats(doc1: Document, doc2: Document) -> Dict[str, Any]:
        """Compare formatting between two documents"""
        # Compare styles
        styles1 = set(style.name for style in doc1.styles)
        styles2 = set(style.name for style in doc2.styles)

        # Compare page setup
        section1 = doc1.sections[0] if doc1.sections else None
        section2 = doc2.sections[0] if doc2.sections else None

        page_diff = {}
        if section1 and section2:
            if section1.page_width != section2.page_width:
                page_diff['width'] = {
                    'doc1': str(section1.page_width),
                    'doc2': str(section2.page_width)
                }
            if section1.page_height != section2.page_height:
                page_diff['height'] = {
                    'doc1': str(section1.page_height),
                    'doc2': str(section2.page_height)
                }

        return {
            'styles_only_in_doc1': list(styles1 - styles2),
            'styles_only_in_doc2': list(styles2 - styles1),
            'page_setup_differences': page_diff
        }

    @staticmethod
    def _generate_comparison_report(result: Dict[str, Any], path1: str, path2: str) -> str:
        """Generate a detailed comparison report"""
        lines = []
        lines.append("=" * 60)
        lines.append("  文档对比报告")
        lines.append("=" * 60)
        lines.append("")
        lines.append(f"文档1: {path1}")
        lines.append(f"文档2: {path2}")
        lines.append("")
        lines.append("-" * 60)
        lines.append("总结")
        lines.append("-" * 60)
        lines.append(f"新增内容: {result['added_count']} 段")
        lines.append(f"删除内容: {result['removed_count']} 段")
        lines.append(f"修改内容: {result['modified_count']} 段")
        lines.append("")

        if result['added']:
            lines.append("-" * 60)
            lines.append("新增内容")
            lines.append("-" * 60)
            for i, text in enumerate(result['added'][:10], 1):
                lines.append(f"{i}. {text[:100]}...")
            if len(result['added']) > 10:
                lines.append(f"... 还有 {len(result['added']) - 10} 段")
            lines.append("")

        if result['removed']:
            lines.append("-" * 60)
            lines.append("删除内容")
            lines.append("-" * 60)
            for i, text in enumerate(result['removed'][:10], 1):
                lines.append(f"{i}. {text[:100]}...")
            if len(result['removed']) > 10:
                lines.append(f"... 还有 {len(result['removed']) - 10} 段")
            lines.append("")

        if result['modified']:
            lines.append("-" * 60)
            lines.append("修改内容")
            lines.append("-" * 60)
            for i, mod in enumerate(result['modified'][:10], 1):
                lines.append(f"{i}. 相似度: {mod['similarity']}")
                lines.append(f"   原文: {mod['original'][:80]}...")
                lines.append(f"   新文: {mod['modified'][:80]}...")
            if len(result['modified']) > 10:
                lines.append(f"... 还有 {len(result['modified']) - 10} 段")
            lines.append("")

        lines.append("=" * 60)
        return "\n".join(lines)

    @register_tool("doc_analyze_statistics",
                   required_params=['path'],
                   optional_params=['detailed'])
    @staticmethod
    def analyze_statistics(path: str, detailed: bool = False) -> Dict[str, Any]:
        """
        Analyze document statistics

        Args:
            path: Path to document
            detailed: Whether to include detailed analysis (default: False)

        Returns:
            Statistical analysis of the document
        """
        try:
            doc = Document(path)

            # Basic counts
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            total_sections = len(doc.sections)

            # Text analysis
            all_text = []
            non_empty_paragraphs = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    non_empty_paragraphs += 1
                    all_text.append(para.text)

            full_text = ' '.join(all_text)
            total_characters = len(full_text)
            total_characters_no_spaces = len(full_text.replace(' ', ''))

            # Word count (split by spaces)
            words = full_text.split()
            total_words = len(words)

            # Average words per paragraph
            avg_words_per_para = total_words / non_empty_paragraphs if non_empty_paragraphs > 0 else 0

            # Font analysis
            font_usage = Counter()
            font_size_usage = Counter()

            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        font_usage[run.font.name] += 1
                    if run.font.size:
                        size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size
                        font_size_usage[f"{size_pt}pt"] += 1

            # Style usage
            style_usage = Counter()
            for para in doc.paragraphs:
                if para.style:
                    style_usage[para.style.name] += 1

            # Heading analysis
            headings = []
            for i, para in enumerate(doc.paragraphs):
                if para.style and 'Heading' in para.style.name:
                    headings.append({
                        'index': i,
                        'level': para.style.name,
                        'text': para.text[:50]
                    })

            # Table analysis
            table_stats = []
            for i, table in enumerate(doc.tables):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                cells = rows * cols
                table_stats.append({
                    'index': i,
                    'rows': rows,
                    'columns': cols,
                    'cells': cells
                })

            result = {
                'success': True,
                'counts': {
                    'paragraphs': total_paragraphs,
                    'non_empty_paragraphs': non_empty_paragraphs,
                    'tables': total_tables,
                    'sections': total_sections,
                    'headings': len(headings),
                    'words': total_words,
                    'characters': total_characters,
                    'characters_no_spaces': total_characters_no_spaces
                },
                'averages': {
                    'words_per_paragraph': round(avg_words_per_para, 2),
                    'characters_per_word': round(total_characters / total_words, 2) if total_words > 0 else 0
                },
                'fonts': {
                    'most_used': font_usage.most_common(5),
                    'total_unique': len(font_usage)
                },
                'font_sizes': {
                    'most_used': font_size_usage.most_common(5),
                    'total_unique': len(font_size_usage)
                },
                'styles': {
                    'most_used': style_usage.most_common(5),
                    'total_unique': len(style_usage)
                }
            }

            # Add detailed information if requested
            if detailed:
                result['detailed'] = {
                    'headings': headings,
                    'tables': table_stats,
                    'all_fonts': dict(font_usage),
                    'all_font_sizes': dict(font_size_usage),
                    'all_styles': dict(style_usage)
                }

            return result

        except Exception as e:
            return {
                'success': False,
                'error': f'统计分析失败: {str(e)}'
            }

    @register_tool("doc_get_metadata",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def get_metadata(path: str) -> Dict[str, Any]:
        """
        Extract document metadata

        Args:
            path: Path to document

        Returns:
            Document metadata including core properties
        """
        try:
            doc = Document(path)
            core_props = doc.core_properties

            metadata = {
                'success': True,
                'core_properties': {
                    'title': core_props.title or '',
                    'subject': core_props.subject or '',
                    'author': core_props.author or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'category': core_props.category or '',
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None,
                    'last_modified_by': core_props.last_modified_by or '',
                    'revision': core_props.revision
                }
            }

            return metadata

        except Exception as e:
            return {
                'success': False,
                'error': f'获取元数据失败: {str(e)}'
            }

    @register_tool("doc_set_metadata",
                   required_params=['path'],
                   optional_params=['title', 'subject', 'author', 'keywords',
                                    'comments', 'category'])
    @staticmethod
    def set_metadata(path: str,
                     title: Optional[str] = None,
                     subject: Optional[str] = None,
                     author: Optional[str] = None,
                     keywords: Optional[str] = None,
                     comments: Optional[str] = None,
                     category: Optional[str] = None) -> Dict[str, Any]:
        """
        Set document metadata

        Args:
            path: Path to document
            title: Document title
            subject: Document subject
            author: Author name
            keywords: Keywords
            comments: Comments
            category: Category

        Returns:
            Success status
        """
        try:
            doc = Document(path)
            core_props = doc.core_properties

            if title is not None:
                core_props.title = title
            if subject is not None:
                core_props.subject = subject
            if author is not None:
                core_props.author = author
            if keywords is not None:
                core_props.keywords = keywords
            if comments is not None:
                core_props.comments = comments
            if category is not None:
                core_props.category = category

            doc.save(path)

            return {
                'success': True,
                'message': '元数据已更新'
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'设置元数据失败: {str(e)}'
            }

    @register_tool("doc_extract_links",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def extract_links(path: str) -> Dict[str, Any]:
        """
        Extract all hyperlinks from document

        Args:
            path: Path to document

        Returns:
            List of all hyperlinks found in the document
        """
        try:
            doc = Document(path)

            links = []
            for para in doc.paragraphs:
                # Check for hyperlinks in paragraph
                if para._element.xpath('.//w:hyperlink'):
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        # Get the relationship ID
                        rid = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rid:
                            try:
                                # Get the actual URL from relationships
                                rel = doc.part.rels[rid]
                                url = rel.target_ref

                                # Get the display text
                                text_elements = hyperlink.xpath('.//w:t', namespaces={
                                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                })
                                display_text = ''.join(t.text for t in text_elements if t.text)

                                links.append({
                                    'text': display_text,
                                    'url': url
                                })
                            except Exception:
                                pass

            return {
                'success': True,
                'link_count': len(links),
                'links': links
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'提取链接失败: {str(e)}'
            }

    @register_tool("doc_word_frequency",
                   required_params=['path'],
                   optional_params=['top_n', 'min_length'])
    @staticmethod
    def word_frequency(path: str,
                       top_n: int = 20,
                       min_length: int = 2) -> Dict[str, Any]:
        """
        Analyze word frequency in document

        Args:
            path: Path to document
            top_n: Number of top words to return (default: 20)
            min_length: Minimum word length to include (default: 2)

        Returns:
            Word frequency analysis
        """
        try:
            doc = Document(path)

            # Extract all text
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)

            full_text = ' '.join(all_text)

            # Split into words and count
            words = full_text.split()

            # Filter by length
            filtered_words = [w for w in words if len(w) >= min_length]

            # Count frequency
            word_count = Counter(filtered_words)

            # Get top N
            top_words = word_count.most_common(top_n)

            return {
                'success': True,
                'total_words': len(words),
                'unique_words': len(word_count),
                'filtered_words': len(filtered_words),
                'top_words': [{'word': word, 'count': count} for word, count in top_words]
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'词频分析失败: {str(e)}'
            }
