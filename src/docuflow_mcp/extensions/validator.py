"""
DocuFlow MCP - Format Validation

Provides format validation and auto-fix functionality
"""

import json
import re
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docuflow_mcp.core.registry import register_tool

# 度量单位解析正则
_MEASURE_RE = re.compile(r'^([0-9]*\.?[0-9]+)\s*(cm|in|pt|mm|emu)?$', re.IGNORECASE)


def _parse_measurement(value: str) -> int:
    """将度量字符串转换为 EMU（docx 内部单位）。

    支持: '2.54cm', '1in', '72pt', '25.4mm', '914400emu', '2.54'(默认cm)
    """
    m = _MEASURE_RE.match(value.strip())
    if not m:
        raise ValueError(f"无法解析度量值: {value!r}")
    num = float(m.group(1))
    unit = (m.group(2) or 'cm').lower()
    if unit == 'cm':
        return Cm(num)
    elif unit == 'in':
        return Inches(num)
    elif unit == 'pt':
        return Pt(num)
    elif unit == 'mm':
        return Cm(num / 10)
    elif unit == 'emu':
        return int(num)
    raise ValueError(f"不支持的单位: {unit}")


class FormatValidator:
    """Format validation and compliance checking"""

    @staticmethod
    def _get_validation_rules_path():
        """Get path to validation rules directory"""
        current_dir = Path(__file__).parent.parent.parent.parent
        return current_dir / "validation_rules"

    @staticmethod
    def _load_preset_rules(preset_name: str) -> Optional[Dict[str, Any]]:
        """Load preset validation rules from JSON"""
        rules_dir = FormatValidator._get_validation_rules_path()
        rules_file = rules_dir / f"{preset_name}.json"

        if not rules_file.exists():
            return None

        with open(rules_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    @staticmethod
    def _check_page_margins(doc: Document, expected: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Check page margins against expected values"""
        issues = []
        section = doc.sections[0]

        # Parse tolerance (default 0.1cm)
        tolerance_str = expected.get('tolerance', '0.1cm')
        tolerance = _parse_measurement(tolerance_str)

        margins = {
            'top': (section.top_margin, expected.get('top')),
            'bottom': (section.bottom_margin, expected.get('bottom')),
            'left': (section.left_margin, expected.get('left')),
            'right': (section.right_margin, expected.get('right'))
        }

        for margin_name, (actual, expected_str) in margins.items():
            if expected_str:
                expected_val = _parse_measurement(expected_str)
                diff = abs(actual - expected_val)

                if diff > tolerance:
                    issues.append({
                        'type': 'page_margin',
                        'margin': margin_name,
                        'expected': expected_str,
                        'actual': f"{actual.cm:.2f}cm",
                        'severity': 'warning'
                    })

        return issues

    @staticmethod
    def _check_page_size(doc: Document, expected: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Check page size against expected values"""
        issues = []
        section = doc.sections[0]

        if 'width' in expected:
            expected_width = _parse_measurement(expected['width'])
            if abs(section.page_width - expected_width) > Cm(0.1):
                issues.append({
                    'type': 'page_width',
                    'expected': expected['width'],
                    'actual': f"{section.page_width.cm:.2f}cm",
                    'severity': 'warning'
                })

        if 'height' in expected:
            expected_height = _parse_measurement(expected['height'])
            if abs(section.page_height - expected_height) > Cm(0.1):
                issues.append({
                    'type': 'page_height',
                    'expected': expected['height'],
                    'actual': f"{section.page_height.cm:.2f}cm",
                    'severity': 'warning'
                })

        return issues

    @staticmethod
    def _check_style_format(doc: Document, style_name: str, expected: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Check if a style matches expected formatting"""
        issues = []

        try:
            style = doc.styles[style_name]
        except KeyError:
            issues.append({
                'type': 'missing_style',
                'style': style_name,
                'severity': 'error'
            })
            return issues

        # Check font properties
        if 'font' in expected and hasattr(style, 'font'):
            font_expected = expected['font']

            if 'name' in font_expected:
                if style.font.name != font_expected['name']:
                    issues.append({
                        'type': 'style_font_name',
                        'style': style_name,
                        'property': 'font_name',
                        'expected': font_expected['name'],
                        'actual': style.font.name or 'None',
                        'severity': 'warning'
                    })

            if 'size' in font_expected:
                expected_size = Pt(float(font_expected['size'].replace('pt', '')))
                if style.font.size and abs(style.font.size - expected_size) > Pt(0.5):
                    issues.append({
                        'type': 'style_font_size',
                        'style': style_name,
                        'property': 'font_size',
                        'expected': font_expected['size'],
                        'actual': f"{style.font.size.pt}pt" if style.font.size else 'None',
                        'severity': 'warning'
                    })

        # Check paragraph properties
        if 'paragraph' in expected and hasattr(style, 'paragraph_format'):
            para_expected = expected['paragraph']
            pf = style.paragraph_format

            if 'line_spacing' in para_expected:
                if pf.line_spacing != para_expected['line_spacing']:
                    issues.append({
                        'type': 'style_line_spacing',
                        'style': style_name,
                        'property': 'line_spacing',
                        'expected': para_expected['line_spacing'],
                        'actual': pf.line_spacing or 'None',
                        'severity': 'info'
                    })

        return issues

    @register_tool("validate_format",
                   required_params=['path'],
                   optional_params=['rules', 'preset_rules'])
    @staticmethod
    def validate_format(path: str,
                        rules: Optional[Dict[str, Any]] = None,
                        preset_rules: Optional[str] = None) -> Dict[str, Any]:
        """
        Validate document format against specified rules.

        Args:
            path: Document path
            rules: Custom validation rules dictionary
            preset_rules: Name of preset rules (e.g., 'mba_thesis')

        Returns:
            Dict with validation results including all issues found
        """
        doc = Document(path)

        # Load rules
        if preset_rules:
            rules = FormatValidator._load_preset_rules(preset_rules)
            if not rules:
                return {
                    "success": False,
                    "error": f"预设规则不存在: {preset_rules}"
                }
        elif not rules:
            return {
                "success": False,
                "error": "必须提供 rules 或 preset_rules 参数"
            }

        all_issues = []

        # Check page setup
        if 'page_setup' in rules:
            page_setup = rules['page_setup']

            if 'margins' in page_setup:
                issues = FormatValidator._check_page_margins(doc, page_setup['margins'])
                all_issues.extend(issues)

            if 'size' in page_setup:
                issues = FormatValidator._check_page_size(doc, page_setup['size'])
                all_issues.extend(issues)

        # Check styles
        if 'styles' in rules:
            for style_name, expected_format in rules['styles'].items():
                issues = FormatValidator._check_style_format(doc, style_name, expected_format)
                all_issues.extend(issues)

        # Categorize issues by severity
        errors = [i for i in all_issues if i.get('severity') == 'error']
        warnings = [i for i in all_issues if i.get('severity') == 'warning']
        info = [i for i in all_issues if i.get('severity') == 'info']

        return {
            "success": True,
            "compliant": len(all_issues) == 0,
            "total_issues": len(all_issues),
            "errors": len(errors),
            "warnings": len(warnings),
            "info": len(info),
            "issues": all_issues
        }

    @register_tool("validate_auto_fix",
                   required_params=['path'],
                   optional_params=['rules', 'preset_rules'])
    @staticmethod
    def auto_fix(path: str,
                 rules: Optional[Dict[str, Any]] = None,
                 preset_rules: Optional[str] = None) -> Dict[str, Any]:
        """
        Automatically fix format issues based on validation rules.

        Args:
            path: Document path
            rules: Custom validation rules dictionary
            preset_rules: Name of preset rules

        Returns:
            Dict with fix results
        """
        doc = Document(path)

        # Load rules
        if preset_rules:
            rules = FormatValidator._load_preset_rules(preset_rules)
            if not rules:
                return {
                    "success": False,
                    "error": f"预设规则不存在: {preset_rules}"
                }
        elif not rules:
            return {
                "success": False,
                "error": "必须提供 rules 或 preset_rules 参数"
            }

        fixed_count = 0
        fixes_applied = []

        # Fix page setup
        if 'page_setup' in rules:
            page_setup = rules['page_setup']
            section = doc.sections[0]

            # Fix margins
            if 'margins' in page_setup:
                margins = page_setup['margins']
                if 'top' in margins:
                    section.top_margin = Cm(float(margins['top'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页边距-上')
                if 'bottom' in margins:
                    section.bottom_margin = Cm(float(margins['bottom'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页边距-下')
                if 'left' in margins:
                    section.left_margin = Cm(float(margins['left'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页边距-左')
                if 'right' in margins:
                    section.right_margin = Cm(float(margins['right'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页边距-右')

            # Fix page size
            if 'size' in page_setup:
                size = page_setup['size']
                if 'width' in size:
                    section.page_width = Cm(float(size['width'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页面宽度')
                if 'height' in size:
                    section.page_height = Cm(float(size['height'].replace('cm', '')))
                    fixed_count += 1
                    fixes_applied.append('页面高度')

        # Fix styles
        if 'styles' in rules:
            for style_name, expected_format in rules['styles'].items():
                try:
                    style = doc.styles[style_name]

                    # Fix font properties
                    if 'font' in expected_format and hasattr(style, 'font'):
                        font_expected = expected_format['font']

                        if 'name' in font_expected:
                            style.font.name = font_expected['name']
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-字体名称')

                        if 'size' in font_expected:
                            style.font.size = Pt(float(font_expected['size'].replace('pt', '')))
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-字号')

                        if 'bold' in font_expected:
                            style.font.bold = font_expected['bold']
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-加粗')

                    # Fix paragraph properties
                    if 'paragraph' in expected_format and hasattr(style, 'paragraph_format'):
                        para_expected = expected_format['paragraph']
                        pf = style.paragraph_format

                        if 'alignment' in para_expected:
                            from docuflow_mcp.document import get_alignment
                            pf.alignment = get_alignment(para_expected['alignment'])
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-对齐方式')

                        if 'line_spacing' in para_expected:
                            pf.line_spacing = para_expected['line_spacing']
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-行距')

                        if 'first_line_indent' in para_expected:
                            from docuflow_mcp.document import parse_size
                            pf.first_line_indent = parse_size(para_expected['first_line_indent'])
                            fixed_count += 1
                            fixes_applied.append(f'{style_name}-首行缩进')

                except KeyError:
                    pass  # Style doesn't exist, skip

        doc.save(path)

        return {
            "success": True,
            "message": f"已自动修正 {fixed_count} 个格式问题",
            "fixed_count": fixed_count,
            "fixes_applied": fixes_applied
        }

    @register_tool("validate_generate_report",
                   required_params=['path'],
                   optional_params=['rules', 'preset_rules', 'output_path'])
    @staticmethod
    def generate_report(path: str,
                        rules: Optional[Dict[str, Any]] = None,
                        preset_rules: Optional[str] = None,
                        output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Generate a detailed validation report.

        Args:
            path: Document path
            rules: Custom validation rules dictionary
            preset_rules: Name of preset rules
            output_path: Optional path to save report text file

        Returns:
            Dict with report content
        """
        # Run validation
        validation_result = FormatValidator.validate_format(path, rules, preset_rules)

        if not validation_result.get("success"):
            return validation_result

        # Generate report text
        report_lines = []
        report_lines.append("="*60)
        report_lines.append("  DocuFlow 文档格式验证报告")
        report_lines.append("="*60)
        report_lines.append(f"\n文档: {path}")
        report_lines.append(f"验证规则: {preset_rules or '自定义规则'}\n")

        # Summary
        report_lines.append("-"*60)
        report_lines.append("总结")
        report_lines.append("-"*60)

        if validation_result['compliant']:
            report_lines.append("✓ 文档完全符合格式规范")
        else:
            report_lines.append(f"✗ 发现 {validation_result['total_issues']} 个格式问题")
            report_lines.append(f"  - 错误: {validation_result['errors']}")
            report_lines.append(f"  - 警告: {validation_result['warnings']}")
            report_lines.append(f"  - 信息: {validation_result['info']}")

        # Detailed issues
        if validation_result['issues']:
            report_lines.append("\n" + "-"*60)
            report_lines.append("详细问题")
            report_lines.append("-"*60)

            for i, issue in enumerate(validation_result['issues'], 1):
                severity_icon = {
                    'error': '✗',
                    'warning': '⚠',
                    'info': 'ℹ'
                }.get(issue.get('severity', 'info'), '•')

                report_lines.append(f"\n{i}. [{severity_icon}] {issue.get('type', 'unknown')}")

                if 'style' in issue:
                    report_lines.append(f"   样式: {issue['style']}")
                if 'property' in issue:
                    report_lines.append(f"   属性: {issue['property']}")
                if 'expected' in issue:
                    report_lines.append(f"   期望值: {issue['expected']}")
                if 'actual' in issue:
                    report_lines.append(f"   实际值: {issue['actual']}")

        report_lines.append("\n" + "="*60)

        report_text = "\n".join(report_lines)

        # Save to file if output_path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_text)

            return {
                "success": True,
                "report": report_text,
                "output_path": output_path,
                "compliant": validation_result['compliant'],
                "total_issues": validation_result['total_issues']
            }
        else:
            return {
                "success": True,
                "report": report_text,
                "compliant": validation_result['compliant'],
                "total_issues": validation_result['total_issues']
            }

    @register_tool("validate_check_consistency",
                   required_params=['path'],
                   optional_params=[])
    @staticmethod
    def check_consistency(path: str) -> Dict[str, Any]:
        """
        Check format consistency within the document.

        Analyzes all paragraphs to find formatting inconsistencies.

        Args:
            path: Document path

        Returns:
            Dict with consistency analysis results
        """
        doc = Document(path)
        paragraphs = doc.paragraphs

        # Analyze font usage
        font_usage = {}
        font_size_usage = {}
        alignment_usage = {}

        for para in paragraphs:
            if not para.runs:
                continue

            # Check font name
            for run in para.runs:
                font_name = run.font.name or 'Default'
                font_usage[font_name] = font_usage.get(font_name, 0) + 1

                if run.font.size:
                    size_str = f"{run.font.size.pt}pt"
                    font_size_usage[size_str] = font_size_usage.get(size_str, 0) + 1

            # Check alignment
            if para.paragraph_format.alignment:
                align_str = str(para.paragraph_format.alignment)
                alignment_usage[align_str] = alignment_usage.get(align_str, 0) + 1

        # Find inconsistencies (anything used less than 10% of the time)
        total_runs = sum(font_usage.values())
        total_alignments = sum(alignment_usage.values())

        inconsistent_fonts = {k: v for k, v in font_usage.items() if v < total_runs * 0.1 and v > 0}
        inconsistent_sizes = {k: v for k, v in font_size_usage.items() if v < total_runs * 0.1 and v > 0}
        inconsistent_alignments = {k: v for k, v in alignment_usage.items() if v < total_alignments * 0.1 and v > 0}

        issues = []

        if inconsistent_fonts:
            issues.append({
                'type': 'inconsistent_fonts',
                'message': f'发现 {len(inconsistent_fonts)} 种不常用字体',
                'details': inconsistent_fonts,
                'severity': 'info'
            })

        if inconsistent_sizes:
            issues.append({
                'type': 'inconsistent_font_sizes',
                'message': f'发现 {len(inconsistent_sizes)} 种不常用字号',
                'details': inconsistent_sizes,
                'severity': 'info'
            })

        if inconsistent_alignments:
            issues.append({
                'type': 'inconsistent_alignments',
                'message': f'发现 {len(inconsistent_alignments)} 种不常用对齐方式',
                'details': inconsistent_alignments,
                'severity': 'info'
            })

        return {
            "success": True,
            "consistent": len(issues) == 0,
            "total_issues": len(issues),
            "issues": issues,
            "statistics": {
                "font_usage": font_usage,
                "font_size_usage": font_size_usage,
                "alignment_usage": alignment_usage
            }
        }
