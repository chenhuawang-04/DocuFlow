"""
DocuFlow MCP - Style Management

Provides style creation and management functionality
"""

import json
from typing import Dict, Any, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docuflow_mcp.document import parse_size, parse_color, get_alignment
from docuflow_mcp.core.registry import register_tool


class StyleManager:
    """Style management operations"""

    @register_tool("style_create", required_params=["path", "style_name"], optional_params=["style_type", "base_style", "font_config", "paragraph_config"])
    @staticmethod
    def create_style(path: str, style_name: str, style_type: str = "paragraph",
                     base_style: Optional[str] = None, font_config: Optional[Dict] = None,
                     paragraph_config: Optional[Dict] = None) -> Dict[str, Any]:
        """Create a custom style"""
        doc = Document(path)

        # Map style type
        type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE
        }

        if style_type not in type_map:
            return {"success": False, "error": f"无效的样式类型: {style_type}"}

        try:
            # Check if style already exists
            if style_name in [s.name for s in doc.styles]:
                return {"success": False, "error": f"样式已存在: {style_name}"}

            # Create new style
            style = doc.styles.add_style(style_name, type_map[style_type])

            # Set base style
            if base_style:
                try:
                    style.base_style = doc.styles[base_style]
                except KeyError:
                    pass

            # Apply font configuration
            if font_config and hasattr(style, 'font'):
                if 'name' in font_config:
                    style.font.name = font_config['name']
                if 'size' in font_config:
                    style.font.size = parse_size(font_config['size'])
                if 'bold' in font_config:
                    style.font.bold = font_config['bold']
                if 'italic' in font_config:
                    style.font.italic = font_config['italic']
                if 'color' in font_config:
                    color = parse_color(font_config['color'])
                    if color:
                        style.font.color.rgb = color

            # Apply paragraph configuration
            if paragraph_config and hasattr(style, 'paragraph_format'):
                pf = style.paragraph_format
                if 'alignment' in paragraph_config:
                    pf.alignment = get_alignment(paragraph_config['alignment'])
                if 'line_spacing' in paragraph_config:
                    pf.line_spacing = paragraph_config['line_spacing']
                if 'space_before' in paragraph_config:
                    pf.space_before = parse_size(paragraph_config['space_before'])
                if 'space_after' in paragraph_config:
                    pf.space_after = parse_size(paragraph_config['space_after'])
                if 'first_line_indent' in paragraph_config:
                    pf.first_line_indent = parse_size(paragraph_config['first_line_indent'])

            doc.save(path)

            return {
                "success": True,
                "message": f"样式已创建: {style_name}",
                "style_name": style_name
            }

        except Exception as e:
            return {"success": False, "error": f"创建样式失败: {str(e)}"}

    @register_tool("style_modify", required_params=["path", "style_name"], optional_params=["font_config", "paragraph_config"])
    @staticmethod
    def modify_style(path: str, style_name: str, font_config: Optional[Dict] = None,
                     paragraph_config: Optional[Dict] = None) -> Dict[str, Any]:
        """Modify an existing style"""
        doc = Document(path)

        try:
            style = doc.styles[style_name]
        except KeyError:
            return {"success": False, "error": f"样式不存在: {style_name}"}

        # Apply font configuration
        if font_config and hasattr(style, 'font'):
            if 'name' in font_config:
                style.font.name = font_config['name']
            if 'size' in font_config:
                style.font.size = parse_size(font_config['size'])
            if 'bold' in font_config:
                style.font.bold = font_config['bold']
            if 'italic' in font_config:
                style.font.italic = font_config['italic']
            if 'color' in font_config:
                color = parse_color(font_config['color'])
                if color:
                    style.font.color.rgb = color

        # Apply paragraph configuration
        if paragraph_config and hasattr(style, 'paragraph_format'):
            pf = style.paragraph_format
            if 'alignment' in paragraph_config:
                pf.alignment = get_alignment(paragraph_config['alignment'])
            if 'line_spacing' in paragraph_config:
                pf.line_spacing = paragraph_config['line_spacing']
            if 'space_before' in paragraph_config:
                pf.space_before = parse_size(paragraph_config['space_before'])
            if 'space_after' in paragraph_config:
                pf.space_after = parse_size(paragraph_config['space_after'])
            if 'first_line_indent' in paragraph_config:
                pf.first_line_indent = parse_size(paragraph_config['first_line_indent'])

        doc.save(path)

        return {
            "success": True,
            "message": f"样式已修改: {style_name}"
        }

    @register_tool("style_export", required_params=["path"], optional_params=["output_path"])
    @staticmethod
    def export_styles(path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """Export styles to JSON"""
        doc = Document(path)

        styles_data = {}

        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_info = {
                    "type": "paragraph",
                    "builtin": style.builtin
                }

                # Export font properties
                if hasattr(style, 'font'):
                    font_info = {}
                    if style.font.name:
                        font_info['name'] = style.font.name
                    if style.font.size:
                        font_info['size'] = f"{style.font.size.pt}pt"
                    if style.font.bold is not None:
                        font_info['bold'] = style.font.bold
                    if style.font.italic is not None:
                        font_info['italic'] = style.font.italic
                    if font_info:
                        style_info['font'] = font_info

                # Export paragraph properties
                if hasattr(style, 'paragraph_format'):
                    pf = style.paragraph_format
                    para_info = {}
                    if pf.alignment:
                        para_info['alignment'] = str(pf.alignment).split('.')[-1].lower()
                    if pf.line_spacing:
                        para_info['line_spacing'] = pf.line_spacing
                    if pf.space_before:
                        para_info['space_before'] = f"{pf.space_before.pt}pt"
                    if pf.space_after:
                        para_info['space_after'] = f"{pf.space_after.pt}pt"
                    if pf.first_line_indent:
                        para_info['first_line_indent'] = f"{pf.first_line_indent.pt}pt"
                    if para_info:
                        style_info['paragraph'] = para_info

                styles_data[style.name] = style_info

        json_data = json.dumps(styles_data, ensure_ascii=False, indent=2)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_data)
            return {
                "success": True,
                "message": f"样式已导出到: {output_path}",
                "output_path": output_path,
                "style_count": len(styles_data)
            }
        else:
            return {
                "success": True,
                "styles": styles_data,
                "style_count": len(styles_data)
            }

    @register_tool("style_import", required_params=["path", "styles_json"], optional_params=[])
    @staticmethod
    def import_styles(path: str, styles_json: str) -> Dict[str, Any]:
        """Import styles from JSON"""
        doc = Document(path)

        try:
            styles_data = json.loads(styles_json)
        except json.JSONDecodeError as e:
            return {"success": False, "error": f"JSON解析失败: {str(e)}"}

        created_count = 0
        updated_count = 0

        for style_name, style_def in styles_data.items():
            # Check if style exists
            style_exists = style_name in [s.name for s in doc.styles]

            if not style_exists:
                # Create new style
                try:
                    style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    created_count += 1
                except:
                    continue
            else:
                style = doc.styles[style_name]
                updated_count += 1

            # Apply font configuration
            if 'font' in style_def and hasattr(style, 'font'):
                font_def = style_def['font']
                if 'name' in font_def:
                    style.font.name = font_def['name']
                if 'size' in font_def:
                    style.font.size = parse_size(font_def['size'])
                if 'bold' in font_def:
                    style.font.bold = font_def['bold']
                if 'italic' in font_def:
                    style.font.italic = font_def['italic']

            # Apply paragraph configuration
            if 'paragraph' in style_def and hasattr(style, 'paragraph_format'):
                para_def = style_def['paragraph']
                pf = style.paragraph_format
                if 'alignment' in para_def:
                    pf.alignment = get_alignment(para_def['alignment'])
                if 'line_spacing' in para_def:
                    pf.line_spacing = para_def['line_spacing']
                if 'space_before' in para_def:
                    pf.space_before = parse_size(para_def['space_before'])
                if 'space_after' in para_def:
                    pf.space_after = parse_size(para_def['space_after'])
                if 'first_line_indent' in para_def:
                    pf.first_line_indent = parse_size(para_def['first_line_indent'])

        doc.save(path)

        return {
            "success": True,
            "message": f"样式已导入: 创建 {created_count} 个，更新 {updated_count} 个",
            "created": created_count,
            "updated": updated_count
        }

    @register_tool("style_copy", required_params=["path", "source_style", "new_style_name"], optional_params=[])
    @staticmethod
    def copy_style(path: str, source_style: str, new_style_name: str) -> Dict[str, Any]:
        """Copy an existing style to create a new style"""
        doc = Document(path)

        try:
            source = doc.styles[source_style]
        except KeyError:
            return {"success": False, "error": f"源样式不存在: {source_style}"}

        # Check if target style already exists
        if new_style_name in [s.name for s in doc.styles]:
            return {"success": False, "error": f"目标样式已存在: {new_style_name}"}

        try:
            # Create new style with same type
            new_style = doc.styles.add_style(new_style_name, source.type)

            # Copy font properties
            if hasattr(source, 'font') and hasattr(new_style, 'font'):
                if source.font.name:
                    new_style.font.name = source.font.name
                if source.font.size:
                    new_style.font.size = source.font.size
                if source.font.bold is not None:
                    new_style.font.bold = source.font.bold
                if source.font.italic is not None:
                    new_style.font.italic = source.font.italic
                if source.font.underline is not None:
                    new_style.font.underline = source.font.underline
                if source.font.color and source.font.color.rgb:
                    new_style.font.color.rgb = source.font.color.rgb

            # Copy paragraph properties
            if hasattr(source, 'paragraph_format') and hasattr(new_style, 'paragraph_format'):
                src_pf = source.paragraph_format
                dest_pf = new_style.paragraph_format

                if src_pf.alignment:
                    dest_pf.alignment = src_pf.alignment
                if src_pf.line_spacing:
                    dest_pf.line_spacing = src_pf.line_spacing
                if src_pf.space_before:
                    dest_pf.space_before = src_pf.space_before
                if src_pf.space_after:
                    dest_pf.space_after = src_pf.space_after
                if src_pf.first_line_indent:
                    dest_pf.first_line_indent = src_pf.first_line_indent
                if src_pf.left_indent:
                    dest_pf.left_indent = src_pf.left_indent
                if src_pf.right_indent:
                    dest_pf.right_indent = src_pf.right_indent

            doc.save(path)

            return {
                "success": True,
                "message": f"样式已复制: {source_style} -> {new_style_name}",
                "source_style": source_style,
                "new_style": new_style_name
            }

        except Exception as e:
            return {"success": False, "error": f"复制样式失败: {str(e)}"}

    @register_tool("style_delete", required_params=["path", "style_name"], optional_params=[])
    @staticmethod
    def delete_style(path: str, style_name: str) -> Dict[str, Any]:
        """Delete a custom style"""
        doc = Document(path)

        try:
            style = doc.styles[style_name]
        except KeyError:
            return {"success": False, "error": f"样式不存在: {style_name}"}

        # Check if style is builtin
        if style.builtin:
            return {"success": False, "error": f"无法删除内置样式: {style_name}"}

        try:
            # Delete style by removing from document styles
            # Note: python-docx doesn't have direct delete method
            # We need to use the underlying XML element
            style_element = style.element
            style_element.getparent().remove(style_element)

            doc.save(path)

            return {
                "success": True,
                "message": f"样式已删除: {style_name}"
            }

        except Exception as e:
            return {"success": False, "error": f"删除样式失败: {str(e)}"}

    @register_tool("style_get_info", required_params=["path", "style_name"], optional_params=[])
    @staticmethod
    def get_style_info(path: str, style_name: str) -> Dict[str, Any]:
        """Get detailed information about a style"""
        doc = Document(path)

        try:
            style = doc.styles[style_name]
        except KeyError:
            return {"success": False, "error": f"样式不存在: {style_name}"}

        info = {
            "name": style.name,
            "type": str(style.type).split('.')[-1],
            "builtin": style.builtin,
            "hidden": style.hidden if hasattr(style, 'hidden') else None,
            "priority": style.priority if hasattr(style, 'priority') else None
        }

        # Get font properties
        if hasattr(style, 'font'):
            font_info = {}
            if style.font.name:
                font_info['name'] = style.font.name
            if style.font.size:
                font_info['size'] = f"{style.font.size.pt}pt"
            if style.font.bold is not None:
                font_info['bold'] = style.font.bold
            if style.font.italic is not None:
                font_info['italic'] = style.font.italic
            if style.font.underline is not None:
                font_info['underline'] = True
            if style.font.color and style.font.color.rgb:
                font_info['color'] = f"#{style.font.color.rgb}"

            if font_info:
                info['font'] = font_info

        # Get paragraph properties
        if hasattr(style, 'paragraph_format'):
            pf = style.paragraph_format
            para_info = {}

            if pf.alignment:
                para_info['alignment'] = str(pf.alignment).split('.')[-1].lower()
            if pf.line_spacing:
                para_info['line_spacing'] = pf.line_spacing
            if pf.space_before:
                para_info['space_before'] = f"{pf.space_before.pt}pt"
            if pf.space_after:
                para_info['space_after'] = f"{pf.space_after.pt}pt"
            if pf.first_line_indent:
                para_info['first_line_indent'] = f"{pf.first_line_indent.pt}pt"
            if pf.left_indent:
                para_info['left_indent'] = f"{pf.left_indent.pt}pt"
            if pf.right_indent:
                para_info['right_indent'] = f"{pf.right_indent.pt}pt"

            if para_info:
                info['paragraph'] = para_info

        return {
            "success": True,
            "style_info": info
        }

