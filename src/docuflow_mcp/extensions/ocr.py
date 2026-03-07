"""
DocuFlow OCR - ????????

???
- Tesseract OCR???????
- completion Vision?AI????????
- PDF?????
- ?????Word
"""
import json
import os
import base64
import ssl
import tempfile
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass

from ..core.registry import register_tool
from ..utils.deps import check_import, check_command
from .image_gen import (
    DEFAULT_API_URL as DEFAULT_COMPLETION_API_URL,
    DEFAULT_MODEL as DEFAULT_COMPLETION_MODEL,
    DEFAULT_TIMEOUT as DEFAULT_COMPLETION_TIMEOUT,
    _request_chat_completion,
)


CONFIG_FILE = Path(__file__).resolve().parent.parent.parent.parent / "ocr_config.json"


@dataclass
class OCRResult:
    """OCR????"""

    text: str
    confidence: float  # 0-1
    engine: str  # 'tesseract' or 'completion'
    page: int = 1

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "confidence": self.confidence,
            "engine": self.engine,
            "page": self.page,
        }


class OCROperations:
    """OCR????"""

    IMAGE_FORMATS = [".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp"]

    LANG_MAP = {
        "chinese": "chi_sim+chi_tra",
        "chinese_simplified": "chi_sim",
        "chinese_traditional": "chi_tra",
        "english": "eng",
        "japanese": "jpn",
        "korean": "kor",
        "french": "fra",
        "german": "deu",
        "spanish": "spa",
        "russian": "rus",
        "auto": "chi_sim+eng",
    }

    COMPLETION_ENGINE_ALIASES = {"completion", "claude"}

    @staticmethod
    def _load_config() -> Dict[str, Any]:
        """?? OCR ?????"""
        if CONFIG_FILE.is_file():
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    return data
            except (json.JSONDecodeError, OSError):
                pass
        return {}

    @staticmethod
    def _get_completion_api_key(config: Dict[str, Any], api_key: Optional[str] = None) -> str:
        """Get the completion API key from call arguments or local config."""
        if api_key:
            return api_key

        config_api_key = config.get("api_key")
        if isinstance(config_api_key, str) and config_api_key.strip():
            return config_api_key.strip()

        raise ValueError(
            "???? completion API Key???? api_key??? ocr_config.json ??? api_key?"
        )

    @staticmethod
    def _has_completion_credentials(config: Dict[str, Any], api_key: Optional[str] = None) -> bool:
        try:
            OCROperations._get_completion_api_key(config, api_key)
            return True
        except ValueError:
            return False

    @staticmethod
    def _extract_completion_text(content: Any) -> str:
        """? completion ??????????"""
        if isinstance(content, str):
            return content.strip()
        if isinstance(content, list):
            parts = []
            for item in content:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, dict):
                    text = item.get("text") or item.get("output_text")
                    if isinstance(text, str):
                        parts.append(text)
            return "\n".join(part.strip() for part in parts if part and part.strip()).strip()
        if isinstance(content, dict):
            text = content.get("text") or content.get("output_text")
            if isinstance(text, str):
                return text.strip()
        return ""

    @staticmethod
    def _extract_completion_response_text(response: Dict[str, Any]) -> str:
        """? completion ?? JSON ??? OCR ???"""
        choices = response.get("choices")
        if isinstance(choices, list):
            for choice in choices:
                if not isinstance(choice, dict):
                    continue
                message = choice.get("message")
                if isinstance(message, dict):
                    text = OCROperations._extract_completion_text(message.get("content"))
                    if text:
                        return text
                text = choice.get("text")
                if isinstance(text, str) and text.strip():
                    return text.strip()

        output = response.get("output")
        if isinstance(output, list):
            for item in output:
                if not isinstance(item, dict):
                    continue
                text = OCROperations._extract_completion_text(item.get("content"))
                if text:
                    return text

        text = response.get("output_text")
        if isinstance(text, str) and text.strip():
            return text.strip()

        raise RuntimeError("completion API ??????? OCR ??")

    @staticmethod
    def _image_to_base64(image_path: str) -> str:
        """????? base64?"""
        with open(image_path, "rb") as f:
            return base64.standard_b64encode(f.read()).decode("utf-8")

    @staticmethod
    def _get_image_media_type(image_path: str) -> str:
        """????? media type?"""
        ext = Path(image_path).suffix.lower()
        media_types = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
            ".bmp": "image/bmp",
            ".tiff": "image/tiff",
            ".tif": "image/tiff",
        }
        return media_types.get(ext, "image/png")

    @staticmethod
    def _pdf_to_images(pdf_path: str, dpi: int = 200) -> List[str]:
        """? PDF ????????"""
        if not check_import("pdf2image"):
            raise ImportError("????pdf2image: pip install pdf2image")

        from pdf2image import convert_from_path

        temp_dir = tempfile.mkdtemp(prefix="docuflow_ocr_")
        images = convert_from_path(pdf_path, dpi=dpi)

        image_paths = []
        for i, image in enumerate(images):
            image_path = os.path.join(temp_dir, f"page_{i + 1}.png")
            image.save(image_path, "PNG")
            image_paths.append(image_path)

        return image_paths

    @staticmethod
    def _ocr_with_tesseract(image_path: str, lang: str = "auto") -> OCRResult:
        """?? Tesseract ?? OCR?"""
        if not check_command("tesseract"):
            raise RuntimeError("Tesseract???????")

        tess_lang = OCROperations.LANG_MAP.get(lang, lang)

        try:
            result = subprocess.run(
                ["tesseract", image_path, "stdout", "-l", tess_lang],
                capture_output=True,
                text=True,
                timeout=60,
                encoding="utf-8",
            )
            if result.returncode != 0:
                raise RuntimeError(f"Tesseract??: {result.stderr}")

            text = result.stdout.strip()
            confidence = OCROperations._estimate_confidence(text)
            return OCRResult(text=text, confidence=confidence, engine="tesseract")
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("Tesseract????") from exc

    @staticmethod
    def _estimate_confidence(text: str) -> float:
        """?? OCR ???????"""
        if not text:
            return 0.0

        total_chars = len(text)
        if total_chars == 0:
            return 0.0

        valid_chars = sum(
            1
            for c in text
            if c.isalnum() or "一" <= c <= "鿿" or c in "???????\"'??"
        )
        valid_ratio = valid_chars / total_chars

        import re

        garbage_pattern = re.compile(r"[^\w\s\u4e00-\u9fff???????\"'??]{3,}")
        garbage_matches = garbage_pattern.findall(text)
        garbage_ratio = sum(len(m) for m in garbage_matches) / total_chars if garbage_matches else 0

        confidence = valid_ratio * 0.7 + (1 - garbage_ratio) * 0.3
        return max(0.0, min(1.0, confidence))

    @staticmethod
    def _ocr_with_completion(
        image_path: str,
        api_key: Optional[str] = None,
        prompt: Optional[str] = None,
        model: Optional[str] = None,
        api_url: Optional[str] = None,
        timeout: Optional[int] = None,
    ) -> OCRResult:
        """?? completion ???? OCR?"""
        config = OCROperations._load_config()
        key = OCROperations._get_completion_api_key(config, api_key)
        image_data = OCROperations._image_to_base64(image_path)
        media_type = OCROperations._get_image_media_type(image_path)

        default_prompt = """Task: perform OCR on this image.

Rules:
1. Output only the text visible in the image.
2. Do not answer questions, summarize, explain, translate, or comment.
3. Do not add headings, bullet points, markdown fences, labels, or confidence notes.
4. Preserve line breaks and reading order as closely as possible.
5. If the image contains a table, output plain text rows in reading order.
6. Do not repeat lines unless they are visibly repeated in the image.
7. If some characters are unclear, make a best-effort transcription and still output text only.
8. If no text is visible, return an empty string.
"""

        payload = {
            "model": model or config.get("model") or DEFAULT_COMPLETION_MODEL,
            "max_tokens": 4096,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a strict OCR engine. Return only the extracted text from the image, with no explanation or extra content.",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt or default_prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{media_type};base64,{image_data}"},
                        },
                    ],
                }
            ],
        }
        final_api_url = api_url or config.get("api_url") or DEFAULT_COMPLETION_API_URL
        final_timeout = int(timeout or config.get("timeout") or DEFAULT_COMPLETION_TIMEOUT)
        response, _ = _request_chat_completion(
            final_api_url,
            key,
            payload,
            final_timeout,
            ssl.create_default_context(),
        )
        text = OCROperations._extract_completion_response_text(response)
        return OCRResult(text=text, confidence=0.95, engine="completion")

    @register_tool(
        "ocr_image",
        required_params=["image_path"],
        optional_params=["lang", "engine", "api_key", "prompt", "model", "api_url", "timeout"],
    )
    @staticmethod
    def ocr_image(
        image_path: str,
        lang: str = "auto",
        engine: str = "auto",
        api_key: Optional[str] = None,
        prompt: Optional[str] = None,
        model: Optional[str] = None,
        api_url: Optional[str] = None,
        timeout: Optional[int] = None,
    ) -> Dict[str, Any]:
        """OCR???????"""
        try:
            path = Path(image_path)
            config = OCROperations._load_config()
            if not path.exists():
                return {"success": False, "error": f"?????: {image_path}"}
            if path.suffix.lower() not in OCROperations.IMAGE_FORMATS:
                return {"success": False, "error": f"????????: {path.suffix}"}

            if engine == "auto":
                if check_command("tesseract"):
                    result = OCROperations._ocr_with_tesseract(image_path, lang)
                    if result.confidence < 0.6 and OCROperations._has_completion_credentials(config, api_key):
                        try:
                            result = OCROperations._ocr_with_completion(
                                image_path,
                                api_key,
                                prompt,
                                model,
                                api_url,
                                timeout,
                            )
                        except (RuntimeError, ValueError, OSError):
                            pass
                elif OCROperations._has_completion_credentials(config, api_key):
                    result = OCROperations._ocr_with_completion(
                        image_path,
                        api_key,
                        prompt,
                        model,
                        api_url,
                        timeout,
                    )
                else:
                    return {
                        "success": False,
                        "error": "?????OCR?????Tesseract?completion API?",
                    }
            elif engine == "tesseract":
                if not check_command("tesseract"):
                    return {"success": False, "error": "Tesseract???"}
                result = OCROperations._ocr_with_tesseract(image_path, lang)
            elif engine in OCROperations.COMPLETION_ENGINE_ALIASES:
                result = OCROperations._ocr_with_completion(
                    image_path,
                    api_key,
                    prompt,
                    model,
                    api_url,
                    timeout,
                )
            else:
                return {"success": False, "error": f"????: {engine}"}

            return {
                "success": True,
                "text": result.text,
                "confidence": result.confidence,
                "engine": result.engine,
                "message": f"???????{result.engine}??",
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool(
        "ocr_pdf",
        required_params=["pdf_path"],
        optional_params=["pages", "lang", "engine", "dpi", "api_key", "prompt", "model", "api_url", "timeout"],
    )
    @staticmethod
    def ocr_pdf(
        pdf_path: str,
        pages: Optional[List[int]] = None,
        lang: str = "auto",
        engine: str = "auto",
        dpi: int = 200,
        api_key: Optional[str] = None,
        prompt: Optional[str] = None,
        model: Optional[str] = None,
        api_url: Optional[str] = None,
        timeout: Optional[int] = None,
    ) -> Dict[str, Any]:
        """OCR?? PDF ??????????"""
        try:
            path = Path(pdf_path)
            if not path.exists():
                return {"success": False, "error": f"PDF???: {pdf_path}"}
            if path.suffix.lower() != ".pdf":
                return {"success": False, "error": "????PDF??"}
            if not check_import("pdf2image"):
                return {"success": False, "error": "????pdf2image: pip install pdf2image"}

            image_paths = OCROperations._pdf_to_images(pdf_path, dpi)
            if not image_paths:
                return {"success": False, "error": "PDF?????????????"}

            temp_dir = os.path.dirname(image_paths[0])
            total_pages = len(image_paths)

            try:
                if pages:
                    invalid_pages = [p for p in pages if p < 1 or p > total_pages]
                    if invalid_pages:
                        return {
                            "success": False,
                            "error": f"????: {invalid_pages}?????? 1-{total_pages}",
                        }
                    target_pages = pages
                else:
                    target_pages = list(range(1, total_pages + 1))

                results = []
                full_text_parts = []
                for page_num in target_pages:
                    image_path = image_paths[page_num - 1]
                    ocr_result = OCROperations.ocr_image(
                        image_path=image_path,
                        lang=lang,
                        engine=engine,
                        api_key=api_key,
                        prompt=prompt,
                        model=model,
                        api_url=api_url,
                        timeout=timeout,
                    )
                    if ocr_result.get("success"):
                        results.append(
                            {
                                "page": page_num,
                                "text": ocr_result["text"],
                                "confidence": ocr_result["confidence"],
                                "engine": ocr_result["engine"],
                            }
                        )
                        full_text_parts.append(f"=== ? {page_num} ? ===\n{ocr_result['text']}")
                    else:
                        results.append({"page": page_num, "error": ocr_result.get("error", "????")})
            finally:
                import shutil

                shutil.rmtree(temp_dir, ignore_errors=True)

            valid_results = [r for r in results if "confidence" in r]
            avg_confidence = (
                sum(r["confidence"] for r in valid_results) / len(valid_results) if valid_results else 0
            )
            return {
                "success": True,
                "total_pages": total_pages,
                "processed_pages": len(target_pages),
                "average_confidence": avg_confidence,
                "results": results,
                "full_text": "\n\n".join(full_text_parts),
                "message": f"???? {len(valid_results)}/{len(target_pages)} ?",
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool(
        "ocr_to_docx",
        required_params=["source", "output_path"],
        optional_params=["lang", "engine", "dpi", "api_key", "prompt", "title", "model", "api_url", "timeout"],
    )
    @staticmethod
    def ocr_to_docx(
        source: str,
        output_path: str,
        lang: str = "auto",
        engine: str = "auto",
        dpi: int = 200,
        api_key: Optional[str] = None,
        prompt: Optional[str] = None,
        title: Optional[str] = None,
        model: Optional[str] = None,
        api_url: Optional[str] = None,
        timeout: Optional[int] = None,
    ) -> Dict[str, Any]:
        """OCR??????? Word ???"""
        try:
            source_path = Path(source)
            if not source_path.exists():
                return {"success": False, "error": f"??????: {source}"}

            ext = source_path.suffix.lower()
            if ext == ".pdf":
                ocr_result = OCROperations.ocr_pdf(
                    pdf_path=source,
                    lang=lang,
                    engine=engine,
                    dpi=dpi,
                    api_key=api_key,
                    prompt=prompt,
                    model=model,
                    api_url=api_url,
                    timeout=timeout,
                )
                if not ocr_result.get("success"):
                    return ocr_result
                full_text = ocr_result.get("full_text", "")
                pages = ocr_result.get("processed_pages", 1)
            elif ext in OCROperations.IMAGE_FORMATS:
                ocr_result = OCROperations.ocr_image(
                    image_path=source,
                    lang=lang,
                    engine=engine,
                    api_key=api_key,
                    prompt=prompt,
                    model=model,
                    api_url=api_url,
                    timeout=timeout,
                )
                if not ocr_result.get("success"):
                    return ocr_result
                full_text = ocr_result.get("text", "")
                pages = 1
            else:
                return {"success": False, "error": f"????????: {ext}"}

            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                return {"success": False, "error": "????python-docx"}

            doc = Document()
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in full_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("=== ?") and line.endswith("==="):
                    doc.add_page_break()
                    doc.add_heading(line.strip("= "), level=1)
                else:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.first_line_indent = Pt(21)

            output_dir = os.path.dirname(output_path)
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
            doc.save(output_path)
            return {
                "success": True,
                "output_path": output_path,
                "pages": pages,
                "message": f"OCR????????Word??: {output_path}",
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    @register_tool("ocr_status", required_params=[], optional_params=[])
    @staticmethod
    def get_status() -> Dict[str, Any]:
        """?? OCR ??????????"""
        config = OCROperations._load_config()
        engines = {
            "tesseract": {
                "available": check_command("tesseract"),
                "description": "??OCR??????????",
                "best_for": "?????????",
            },
            "completion": {
                "available": True,
                "api_key_set": OCROperations._has_completion_credentials(config),
                "description": "Remote OCR via completion API",
                "best_for": "complex layouts, tables, scanned pages",
            },
        }
        dependencies = {
            "pdf2image": {
                "available": check_import("pdf2image"),
                "purpose": "PDF????OCR PDF???",
            },
            "PIL": {
                "available": check_import("PIL"),
                "purpose": "????",
            },
        }
        available_engines = [name for name, info in engines.items() if info.get("available")]
        return {
            "success": True,
            "config_file": str(CONFIG_FILE),
            "api_url": config.get("api_url", DEFAULT_COMPLETION_API_URL),
            "model": config.get("model", DEFAULT_COMPLETION_MODEL),
            "timeout": config.get("timeout", DEFAULT_COMPLETION_TIMEOUT),
            "engines": engines,
            "dependencies": dependencies,
            "available_engines": available_engines,
            "supported_languages": list(OCROperations.LANG_MAP.keys()),
            "supported_image_formats": OCROperations.IMAGE_FORMATS,
            "engine_aliases": {"claude": "completion"},
            "message": f"????: {', '.join(available_engines) if available_engines else '?'}",
        }
